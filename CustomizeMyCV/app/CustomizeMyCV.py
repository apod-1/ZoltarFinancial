# -*- coding: utf-8 -*-
"""
Created on Mon Jan 13 20:33:27 2025
This app will create a custom version of your resume to ensure higher visibility by the employers through:
    1. identification of pieces that need to be changed using the identification Agent (1),
    2. sending out pieces for work to respective SME Agents (unlimited) using Prompt Writer Agent (2),
    3. sending it out for verification to a checker Agent (3),
    4. combining it back together with another 'compiler' Agent (4)
    5. sending result to the checker Agent (3) together with instructions sent from results of identification Agent (1)

requirements:
    python -m venv streamlit_env
    streamlit_env\Scripts\activate
    pip install streamlit
    import streamlit as st
    
  **  To Launch:  **
    activate myenv
    streamlit_env\Scripts\activate
    cd C:\ Users/apod7/CustomizeMyCV    
    streamlit run CustomizeMyCV.py
@author: andrew
"""
from docx import Document
from docx.shared import Pt, RGBColor
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
# from email.mime.image import MIMEImage
from email.mime.base import MIMEBase
from email import encoders
import smtplib



from datetime import datetime
import os
import openai
import streamlit as st
# Local imports
# import sys
# import base64
# from dotenv import load_dotenv
# import json

OPENAI_API=None
# Load environment variables
# load_dotenv()
# openai.api_key = os.getenv('API_KEY')
# openai.api_key = OPENAI_API
# Set up OpenAI API key
# OPENAI_API=openai.api_key



try:
    favicon = "https://github.com/apod-1/ZoltarFinancial/raw/main/CustomizeMyCV/media/customizecv_logo48x48.png"
except (KeyError, FileNotFoundError):
    favicon = st.secrets["browser"]["favicon"]

st.set_page_config(page_title="Multi-Agent Resume Customization", page_icon=favicon, layout="wide")

hide_streamlit_style = """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    </style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True) 

   
# removed 1.16 to sort
# def read_resume_sections(directory):
#     resume_sections = {}
#     for filename in os.listdir(directory):
#         if filename.startswith("OG_resume_") and filename.endswith(".txt"):
#             try:
#                 with open(os.path.join(directory, filename), 'r', encoding='utf-8') as file:
#                     resume_sections[filename] = file.read()
#             except UnicodeDecodeError:
#                 st.warning(f"Could not decode {filename}. Please check the file encoding.")
#             except Exception as e:
#                 st.warning(f"An error occurred while reading {filename}: {e}")
#     return resume_sections
def read_resume_sections(directory):
    resume_sections = {}
    for filename in sorted(os.listdir(directory)):
        if filename.startswith("OG_resume_") and filename.endswith(".txt"):
            try:
                with open(os.path.join(directory, filename), 'r', encoding='utf-8') as file:
                    resume_sections[filename] = file.read()
            except UnicodeDecodeError:
                st.warning(f"Could not decode {filename}. Please check the file encoding.")
            except Exception as e:
                st.warning(f"An error occurred while reading {filename}: {e}")
    return dict(sorted(resume_sections.items()))

# Function to interact with OpenAI API (unchanged)
def query_openai(prompt):
    with st.spinner('Generating response...'):
        messages = [{"role": "user", "content": prompt}]
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=messages
        )
        return response.choices[0].message['content'].strip()

pre_prompt = "Without taking away from the overall structure of the resume, examine the below description to make slight changes in the resume to stand out as the perfect candidate without overstepping my own abilities. Here is the job desciption:"

# Agent functions (unchanged)
def identification_agent(resume_sections, user_query):
    prompt = f"Identify parts of the following resume sections that need to be changed based on this request, and Name them sequentially Section X: {user_query}\n\n"
    for section, content in resume_sections.items():
        prompt += f"{section}:\n{content}\n\n"
    return query_openai(prompt)


def prompt_writer_agent(identified_changes, section_content):
    prompt = f"Create a prompt for SME agents to modify this resume section based on the following identified changes:\n{identified_changes}\n\nHere is the original section:\n{section_content}\n"
    response = query_openai(prompt)
    return response  # Return the generated prompt directly

def sme_agent(section_content, section_prompt):
    prompt = f"You are improving this section in a resume to make it more tailored to job descriptions and stand out to recruiters:\n{section_content}\n\nHere is the actual request to work on: {section_prompt}"
    return query_openai(prompt)

def checker_agent(modified_content, original_instructions, resume_sections):
    prompt = f"Verify if the following modified content meets the original instructions to improve the original content :\n\nOriginal Instructions:\n{original_instructions}\n\nModified Content:\n{modified_content}. This is my original resume that only needs mimimal changes to make it stand out to the recruiter:"
    for section, content in resume_sections.items():
        prompt += f"{section}:\n{content}\n\n"
    return query_openai(prompt)

def compiler_agent(modified_sections, resume_sections):
    prompt = f"Combine the following modified resume sections into a cohesive resume:\n\n{modified_sections}. This is my original resume that only needs mimimal changes to make it stand out to the recruiter:"
    for section, content in resume_sections.items():
        prompt += f"{section}:\n{content}\n\n"
    return query_openai(prompt)

def create_output_directory(base_path, today):
    # Create a folder with today's date
    # today = datetime.now().strftime("%Y-%m-%d")
    output_directory = os.path.join(base_path, today)
    os.makedirs(output_directory, exist_ok=True)
    return output_directory


# 1.15.25
# def save_responses_to_txt(output_directory, responses, modified_sections, section_prompts, today):
#     timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
#     for section, response in responses.items():
#         filename = f"{section.replace(' ', '_')}_response_{timestamp}.txt"
#         with open(os.path.join(output_directory, filename), 'w', encoding='utf-8') as file:
#             file.write(response)
    
#     # Save section_prompts
#     filename = f"section_prompts_{timestamp}.txt"
#     with open(os.path.join(output_directory, filename), 'w', encoding='utf-8') as file:
#         json.dump(section_prompts, file, indent=2)

def add_horizontal_line(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def save_modified_responses_to_doc(modified_sections, output_directory, timestamp):
    doc = Document()
    doc.add_heading('Modified Responses', level=1)
    
    for section, content in modified_sections.items():
        doc.add_heading(f'Section: {section}', level=2)
        
        # Split content into lines
        lines = content.split('\n')
        
        for line in lines:
            if line.startswith('###'):
                # Subheading (H3)
                p = doc.add_paragraph()
                run = p.add_run(line.strip('# '))
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 139)  # Dark Blue
            elif line.startswith('**') and line.endswith('**'):
                # Bold text
                p = doc.add_paragraph()
                run = p.add_run(line.strip('*'))
                run.bold = True
            elif line.startswith('---'):
                # Horizontal line
                p = doc.add_paragraph()
                add_horizontal_line(p)
            elif line.strip() == '':
                # Empty line
                doc.add_paragraph()
            else:
                # Regular paragraph
                doc.add_paragraph(line)
        
        doc.add_page_break()
    
    filename = f"Modified_Responses_{timestamp}.docx"
    doc.save(os.path.join(output_directory, filename))



def save_sme_prompts_to_doc(section_prompts, output_directory, timestamp):
    doc = Document()
    doc.add_heading('SME Prompts', level=1)
    
    for section, prompt in section_prompts.items():
        doc.add_heading(f'Section: {section}', level=2)
        
        lines = prompt.split('\n')
        
        for line in lines:
            if line.startswith('###'):
                # Subheading (H3)
                p = doc.add_paragraph()
                run = p.add_run(line.strip('# '))
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 139)  # Dark Blue
            elif line.startswith('**') and line.endswith('**'):
                # Bold text
                p = doc.add_paragraph()
                run = p.add_run(line.strip('*'))
                run.bold = True
            elif line.startswith('1.') or line.startswith('2.') or line.startswith('3.'):
                # Numbered list
                p = doc.add_paragraph(line, style='List Number')
            elif line.startswith('- '):
                # Bullet point
                p = doc.add_paragraph(line.strip('- '), style='List Bullet')
            elif line.strip() == '':
                # Empty line
                doc.add_paragraph()
            else:
                # Regular paragraph
                doc.add_paragraph(line)
        
        doc.add_page_break()
    
    filename = f"SME_Prompts_{timestamp}.docx"
    doc.save(os.path.join(output_directory, filename))

def save_final_check_to_doc(final_check, output_directory, timestamp):
    doc = Document()
    doc.add_heading('Final Check', level=1)
    
    lines = final_check.split('\n')
    
    for line in lines:
        if line.startswith('###'):
            # Section heading (H3)
            p = doc.add_paragraph()
            run = p.add_run(line.strip('# '))
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 139)  # Dark Blue
        elif line.startswith('- '):
            # Bullet point
            p = doc.add_paragraph(line.strip('- '), style='List Bullet')
        elif line.strip() == '':
            # Empty line
            doc.add_paragraph()
        else:
            # Regular paragraph
            doc.add_paragraph(line)
    
    filename = f"Final_Check_{timestamp}.docx"
    doc.save(os.path.join(output_directory, filename))    
    
def save_compiled_resume_to_doc(compiled_resume, output_directory, timestamp):
    doc = Document()
    doc.add_heading('Customized Resume', level=1)
    
    sections = compiled_resume.split('\n\n')
    for section in sections:
        lines = section.split('\n')
        for line in lines:
            if line.startswith('###'):
                p = doc.add_paragraph()
                run = p.add_run(line.strip('# '))
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 139)
            elif '**' in line:
                p = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:  # Odd-indexed parts are between ** and should be bold
                        run.bold = True
            elif line.startswith('- '):
                p = doc.add_paragraph(line.strip('- '), style='List Bullet')
            elif line.startswith('---'):
                continue
            elif line.strip() == '':
                doc.add_paragraph()
            else:
                doc.add_paragraph(line)
    
    filename = f"Customized_Resume_{timestamp}.docx"
    doc.save(os.path.join(output_directory, filename))
    
def send_email_with_attachments(recipient_email, output_directory, timestamp):
    try:
        sender_email = st.secrets["GMAIL"]["GMAIL_ACCT"]
        sender_password = st.secrets["GMAIL"]["GMAIL_PASS"]
    except:
        # If Streamlit secrets are not available, use environment variables
        sender_email = os.getenv('GMAIL_ACCT')
        sender_password = os.getenv('GMAIL_PASS') 
        st.error("Gmail credentials not found in secrets. Please check your configuration.")
        return
    
    if not sender_email or not sender_password:
        st.error("Gmail credentials not found. Please check your configuration.")
        return False

    subject = "Your Customized Resume"
    
    # HTML content
    html_content = f"""
    <html>
        <body>
            <p>Please find attached your customized resume documents.</p>
            <p>These documents were generated on: {timestamp}</p>
            <p>Thank you for using our Multi-Agent Resume Customization App!</p>
        </body>
    </html>
    """

    message = MIMEMultipart()
    message['From'] = f"Resume Customization App <{sender_email}>"
    message['To'] = recipient_email
    message['Subject'] = subject
    message.attach(MIMEText(html_content, 'html'))

    # Attach files
    for filename in os.listdir(output_directory):
        filepath = os.path.join(output_directory, filename)
        if os.path.isfile(filepath):
            with open(filepath, 'rb') as attachment:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename="{filename}"')
            message.attach(part)
            st.write(f"Attached file: {filename}")  # Debug information

    try:
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(sender_email, sender_password)
            server.send_message(message)
        st.success("Email sent successfully!")
        return True
    except Exception as e:
        st.error(f"Failed to send email: {str(e)}")
        return False

    
# 1.15.25 instead of above 
def main():

    try:
        if OPENAI_API:
            openai.api_key = OPENAI_API       
        else: 
            openai.api_key = st.secrets["openai"]["api_key"] 
    except KeyError:
        st.error("OpenAI API key not found in secrets. Please clear cache and reboot app.")
        st.stop()     
        ### need it
        # openai.api_key=OPENAI_API
    if 'resume_customized' not in st.session_state:
        st.session_state.resume_customized = False
    if 'start_over' not in st.session_state:
        st.session_state.start_over = False
    if 'output_directory' not in st.session_state:
        st.session_state.output_directory = None
        
    st.title("Multi-Agent Resume Customization App")
    st.markdown(" ",unsafe_allow_html=True,help="Hi there! I built this to help with my own search but made it available to everyone to increase chances of getting hired!  Improvements under way! -Andrew")
    today = datetime.now().strftime("%Y%m%d_%H%M%S")

    # st.markdown(
    #     """
    #     <style>
    #         .top-frame {
    #             position: relative;
    #             width: 100vw;
    #             height: 30vh;  /* Video height */
    #             overflow: visible;  /* Key change: from hidden to visible */
    #             margin-bottom: -50px;
    #         }
    #         .top-frame video {
    #             position: absolute;
    #             top: 0;
    #             left: 0;
    #             width: 100%;
    #             height: 100%;
    #             object-fit: cover;
    #         }
    #         .image-container {
    #             position: absolute;
    #             top: -5px;  /* Adjusted to move image further up */
    #             left: 50%;
    #             transform: translateX(-50%);
    #             z-index: 1;
    #         }
    #         .image-container img {
    #             width: 180px;  /* Back to original size */
    #             height: 180px;
    #             border-radius: 50%;
    #             object-fit: cover;
    #             border: 4px solid white;
    #             box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    #         }
    #         .divider {
    #             display: none;
    #         }
    #     </style>
    #     <div class="top-frame">
    #         <video autoplay loop muted>
    #             <source src="https://github.com/apod-1/ZoltarFinancial/raw/main/docs/wave_vid.mp4" type="video/mp4">
    #         </video>
    #         <div class="image-container">
    #             <img src="https://github.com/apod-1/ZoltarFinancial/raw/main/CustomizeMyCV/media/customizecv_logo300x300.png" alt="Sprinkle Job Description on your Resume Image">
    #         </div>
    #     </div>
    #     <div class="divider"></div>
    #     """,
    #     unsafe_allow_html=True
    # )

            # .image-container {
            #     position: absolute;
            #     top: 50%;
            #     left: 50%;
            #     transform: translate(-50%, -50%);
            #     z-index: 1;
            # }
            # .image-container img {
            #     width: 200px;  /* Slightly smaller logo */
            #     height: 200px;
            #     border-radius: 50%;
            #     object-fit: cover;
            #     border: 4px solid white;
            #     box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            # }

        # @keyframes scroll {
        #     0% { transform: translateX(100%); opacity: 0; }
        #     10% { transform: translateX(80%); opacity: 1; }
        #     45% { transform: translateX(0%); }
        #     55% { transform: translateX(-50%); }
        #     90% { transform: translateX(-80%); opacity: 1; }
        #     100% { transform: translateX(-100%); opacity: 0; }
        # }

# cubic-bezier(0.25, 0.1, 0.25, 1) infinite;
        # .step:nth-child(2) { animation-delay: 20s; }
        # .step:nth-child(3) { animation-delay: 40s; }        

        # @keyframes scroll {
        #     0% { 
        #         transform: translateX(400%);
        #         opacity: 0; 
        #     }
        #     20% { 
        #         transform: translateX(150%);
        #         opacity: 0.2; 
        #     }
        #     40% { transform: translateX(0%); 
        #          opacity: 0.8; 
        #          }
        #     70% { transform: translateX(-100%); 
        #          opacity: 1; 
        #          }
        #     95% { 
        #         transform: translateX(-100%);
        #         opacity: 1; 
        #     }
        #     100% { 
        #         transform: translateX(-2000%);
        #         opacity: 0; 
        #     }
        # }

    st.markdown(
        """
        <style>
        .ticker-wrapper {
            width: 100%;
            height: 50px;
            overflow: hidden;
            background: black;
            position: relative;
            color: white;
        }
        @keyframes scroll {
            0% { 
                transform: translateX(400%);  /* Even further to the right */
                opacity: 0; 
            }
            10% { 
                transform: translateX(250%);  /* More gradual entry */
                opacity: 1; 
            }
            45% { transform: translateX(0%); }
            55% { transform: translateX(-50%); }
            80% { 
                transform: translateX(-300%);  /* Quicker exit to left */
                opacity: 0.2; 
            }
            100% { 
                transform: translateX(-400%);  /* Further off screen */
                opacity: 0; 
            }
        }

        .step {
            position: absolute;
            white-space: nowrap;
            font-size: 1.2rem;
            animation: scroll 40s linear;
            animation-fill-mode: both;
        }
        .top-frame {
            position: relative;
            height: 33vh;
            overflow: hidden;
            width: 100%;
            margin: 0 auto;
        }
        .image-container {
            position: absolute;
            top: 30%;
            left: 50%;
            transform: translate(-50%, -50%);
            z-index: 2;
            width: 9.5vw;
            height: 9.5vw;
            border-radius: 50%;
            overflow: hidden;
            box-shadow: 0 0 10px rgba(0,0,0,0.5);
        }
        .image-container img {
            width: 100%;
            height: 100%;
            object-fit: cover;
        }
        .top-frame video {
            position: absolute;
            top: 0%;
            bottom: -30%;
            left: 0;
            width: 100%;
            height: 166.67%;
            object-fit: cover;
            object-position: center center;
            z-index: 1;
        }
        .divider {
            border-top: 3px solid black;
            margin-top: 20px;
            margin-bottom: 20px;
        }
        @media (orientation: portrait) {
            .top-frame {
                height: 25vh;
            }
            .top-frame video {
                top: -37.5%;
                bottom: -37.5%;
                height: 175%;
                object-position: center center;
            }
            .image-container {
                width: 19vw;
                height: 19vw;
            }
        }
        </style>
    
        <div class="ticker-wrapper">
            <div class="step" style="animation-delay: 0s;">Welcome to the Multi-Agent Resume Customization App!</div>
            <div class="step" style="animation-delay: 8s;">Step 1: Choose to use Andrew's resume or enter your own.</div>
            <div class="step" style="animation-delay: 16s;">Step 2: Enter the job description you're applying for.</div>
            <div class="step" style="animation-delay: 24s;">Step 3: Click 'Customize Resume' to start the process.</div>
            <div class="step" style="animation-delay: 32s;">Step 4: Review the customized sections of your resume.</div>
            <div class="step" style="animation-delay: 40s;">Step 5: Enter your email to receive the customized resume.</div>
            <div class="step" style="animation-delay: 48s;">Step 6: Click 'Send Resume via Email' to get your tailored resume.</div>
            <div class="step" style="animation-delay: 56s;">Tip: Use 'Start Over' to begin a new customization process.</div>
        </div>
        
        <div class="top-frame">
            <video autoplay loop muted>
                <source src="https://github.com/apod-1/ZoltarFinancial/raw/main/docs/wave_vid.mp4" type="video/mp4">
            </video>
            <div class="image-container">
                <img src="https://github.com/apod-1/ZoltarFinancial/raw/main/CustomizeMyCV/media/customizecv_logo300x300.png" alt="Sprinkle Job Description on your Resume Image">
            </div>
        </div>
        <div class="divider"></div>
        
        <script>
        document.addEventListener('DOMContentLoaded', (event) => {
            const steps = document.querySelectorAll('.step');
            const stepDuration = 7.5; // 7.5 seconds per step animation (halved from 15)
            const totalDuration = steps.length * stepDuration;
            const cycleDelay = 2.5; // 2.5 seconds delay between cycles (halved from 5)
            
            function restartAnimation() {
                steps.forEach((step, index) => {
                    step.style.animationDelay = `${index * stepDuration}s`;
                    step.style.animationIterationCount = '10';
                });
            }
        
            function startCycle() {
                restartAnimation();
                setTimeout(startCycle, totalDuration * 1000 + cycleDelay * 1000);
            }
        
            startCycle();
        });
        </script>
        """,
        unsafe_allow_html=True
    )


    # Initialize resume_sections as an empty dictionary
    resume_sections = {}
    resume_source = st.radio("Choose resume source:", ["Use Andrew's Resume", "Manual Entry"])

    # Add a button to use Andrew's resume
    # if st.button("Use Andrew's Resume"):
    # if resume_source == "Use Andrew's Resume":
    #     if os.path.exists(r'C:\Users\apod7\CustomizeMyCV\docs'):
    #         # Cloud environment
    #         resume_directory = r'C:\Users\apod7\CustomizeMyCV\docs'
    #     else:
    #         # Local environment
    #         resume_directory = '/mount/src/zoltarfinancial/CustomizeMyCV/docs'

    #     # resume_directory = r"C:\Users\apod7\CustomizeMyCV\docs"
    #     if os.path.exists(resume_directory):
    #         resume_sections = read_resume_sections(resume_directory)
    #         if not resume_sections:
    #             st.warning("No resume sections found in the directory.")
    #     else:
    #         st.error("Resume directory not found.")


# alphabetic - in order

    if resume_source == "Use Andrew's Resume":
        if os.path.exists(r'C:\Users\apod7\CustomizeMyCV\docs'):
            # Cloud environment
            resume_directory = r'C:\Users\apod7\CustomizeMyCV\docs'
        else:
            # Local environment
            resume_directory = '/mount/src/zoltarfinancial/CustomizeMyCV/docs'
    
        if os.path.exists(resume_directory):
            resume_sections = read_resume_sections(resume_directory)
            if resume_sections:
                st.subheader("Original Resume Sections")
                for section, content in resume_sections.items():
                    with st.expander(f"Section: {section}"):
                        st.text(content)
            else:
                st.warning("No resume sections found in the directory.")
        else:
            st.error("Resume directory not found.")

    else:
        st.write("Please enter your resume manually:")
        manual_resume = st.text_area("Enter your resume text here:")
        if manual_resume:
            resume_sections = {"manual_resume.txt": manual_resume}

    # if resume_sections:
    # Display original resume sections
    if resume_sections:
        # st.subheader("Original Resume Sections")
        # for section, content in resume_sections.items():
        #     with st.expander(f"Section: {section}"):
        #         st.text(content)

        # if not st.session_state.resume_customized:
        # Input for customization query
        user_query1 = st.text_input("Enter job desciption:")
        user_query = pre_prompt + user_query1

        if st.button("Customize Resume"):
            if user_query:
                # Step 1: Identification Agent (unchanged)
                identified_changes = identification_agent(resume_sections, user_query)
                st.subheader("Identified Changes")
                st.info(identified_changes)
    
                # Define the mapping between resume section filenames and modification keys
                if "manual_resume.txt" in resume_sections:
                    section_mapping = {"manual_resume.txt": "Manual Resume"}
                else:
                    section_mapping = {
                        "OG_resume_1_overview.txt": "Overview",
                        "OG_resume_2_toolkit.txt": "Toolkit",
                        "OG_resume_3_innovation.txt": "Analytics + AI Innovation",
                        "OG_resume_4_zoltar.txt": "Zoltar Financial, Inc.",
                        "OG_resume_5_citi.txt": "Citigroup",
                        "OG_resume_6_enova.txt": "Enova International",
                        "OG_resume_7_catalina.txt": "Catalina",
                        "OG_resume_8_leo.txt": "Leo Burnett",
                        "OG_resume_9a_tu.txt": "TransUnion",
                        "OG_resume_9b_KBMG.txt": "KBM",
                        "OG_resume_9c_AMS.txt": "AMS Direct"
                    }
    
                # Step 2 & 3: Process each section one at a time
                modified_sections = {}
                responses = {}  # To store all agent responses
                section_prompts = {}  # To store all section prompts
    
                for section_filename, section_content in resume_sections.items():
                    modification_key = section_mapping.get(section_filename)
                    if modification_key:
                        # Generate prompt for the current section based on identified changes
                        section_prompt = prompt_writer_agent(identified_changes, section_content)
                        section_prompts[section_filename] = section_prompt
                        st.subheader(f"SME Prompt for {section_filename}")
                        st.info(section_prompt)  # Display the generated prompt
    
                        # Send only the current section content to sme_agent
                        modified_content = sme_agent(section_content, section_prompt)
                        modified_sections[section_filename] = modified_content
                        responses[section_filename] = modified_content  # Store response
                        
                        st.subheader(f"Modified Section: {section_filename}")
                        st.write("Original Content:")
                        st.text(section_content)
                        st.write("Modified Content:")
                        st.success(modified_content)   
                # Step 4: Compiler Agent
                compiled_resume = compiler_agent(modified_sections, resume_sections)
                st.subheader("Compiled Resume")
                st.info(compiled_resume)
                
               # # Save compiled resume to .docx file
               #  output_file_path = save_compiled_resume_to_doc(compiled_resume)
               #  st.success(f"Resume saved successfully at: {output_file_path}")    
                
                # Step 5: Checker Agent
                final_check = checker_agent(compiled_resume, identified_changes, resume_sections)
                st.subheader("Final Check")
                st.success(final_check)

                # 1.15.25pm                
                # Create output directory and save files
                if os.path.exists(r'C:\Users\apod7\CustomizeMyCV\output'):
                    # Cloud environment
                    output_directory = create_output_directory(r"C:\Users\apod7\CustomizeMyCV\output", today)
                else:
                    # Local environment
                    output_directory = create_output_directory('/mount/src/zoltarfinancial/CustomizeMyCV/output', today)

                # output_directory = create_output_directory(r"C:\Users\apod7\CustomizeMyCV\output", today)
                # Save the output_directory to session state
                st.session_state.output_directory = output_directory                  
                # Save SME Prompts to .docx
                save_sme_prompts_to_doc(section_prompts, output_directory, today)
                
                # Save Modified Responses to .docx
                save_modified_responses_to_doc(modified_sections, output_directory, today)

                # Save compiled resume to .docx
                save_compiled_resume_to_doc(compiled_resume, output_directory, today)
                
                # Save Final Check to .docx
                save_final_check_to_doc(final_check, output_directory, today)
                
                st.success(f"All files saved successfully to: {output_directory}")                
                
                # st.success(f"All files saved successfully to: {output_directory}")


# 1.16.25 - new section to send results
                # Add email input and send button
                st.session_state.resume_customized = True
            else:
                    st.warning("Please enter a customization query.")
                    
        # recipient_email = st.text_input("Enter your email to receive the customized resume:", key='first_email')
        # send_button_clicked = st.button("Send Resume via Email",key = 'first')
        # start_over_clicked = st.button("Start Over", key='startover')
        if st.session_state.resume_customized:
            recipient_email = st.text_input("Enter your email to receive the customized resume:", key='first_email')
            if st.button("Send Resume via Email",key = 'first'):
                if recipient_email:
                    st.write(f"Attempting to send email to: {recipient_email}")
                    st.write(f"Using output directory: {st.session_state.output_directory}")
                    st.write(f"Files in output directory: {os.listdir(st.session_state.output_directory)}")
                    if send_email_with_attachments(recipient_email, st.session_state.output_directory, today):
                        st.success("Email sent successfully with attachments!")
                    else:
                        st.error("Failed to send email. Please check the error messages above.")
                else:
                    st.warning("Please enter a valid email address.")

        # Add a "Start Over" button
        if st.button("Start Over", key='startover'):
            st.session_state.resume_customized = False
            st.session_state.start_over = True
            st.session_state.output_directory= None
            st.rerun()

                # # After all processing is done:
                # st.rerun()
            # else:
            #         st.warning("Please enter a customization query.")
    # if st.session_state.resume_customized:
    #     # Display the email input and send button
    #     recipient_email = st.text_input("Enter your email to receive the customized resume:", key='second_email')
    #     if st.button("Send Resume via Email", key = 'second'):
    #         if recipient_email:
    #             if send_email_with_attachments(recipient_email, st.session_state.output_directory, today):
    #                 st.success("Email sent successfully!")
    #             else:
    #                 st.error("Failed to send email. Please try again.")
    #         else:
    #             st.warning("Please enter a valid email address.")

        # # Add a "Start Over" button
        # if st.button("Start Over", key='startover'):
        #     st.session_state.resume_customized = False
        #     st.session_state.start_over = True
        #     st.rerun()

# 1.16.25 end

        # else:
        #     st.warning("Resume loaded! Check Query and Proceed...")    

    else:
        st.warning("No resume content available. Please use Andrew's resume or enter manually.")    
        # else:
        #     st.warning("Please enter a customization query.")
if __name__ == "__main__":
    main()