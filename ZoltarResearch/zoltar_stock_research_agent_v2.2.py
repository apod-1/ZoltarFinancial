# -*- coding: utf-8 -*-
"""
Created on Thu Apr  3 14:46:53 2025

General description:
    This is a capstone project for genAI 5-day Intensive Google-led course
    
    Main goal: Show knowledge of concepts to establish a grounding database, vectorize it, and utilize available functions to be envoked by AI Agent at the appropriate time (as-needed)
    Secondary goal: Show knowledge of advanced concepts, including: 
            multi-agent system, 
            langgraph representation
            transparent chain-of-thought steps, 
            API calls/search in real-time data, 
            and source citation encoding.


Application: Helpful bot that can analyze stocks and present a comprehensive report of the analysis to the user, utilizing proprietary Zoltar Ranks as grounding database.

launch (at home)
    activate myenv
    streamlit_env\Scripts\activate
    cd C:\ Users\apod7\StockPicker\app\ZoltarFinancial\ZoltarResearch    
    streamlit run zoltar_stock_research_agent.py
@author: apod
"""
import os
import openai
import pandas as pd
import streamlit as st
import markdown2    
import json
import sqlite3
import textwrap
import asyncio
# import IPython
import io
import altair as alt
import seaborn as sns
import base64
import re
import matplotlib.pyplot as plt  # Import Matplotlib globally
import requests
# from IPython.display import display, Image, Markdown
from pprint import pprint
from pprint import pformat
from dotenv import load_dotenv
from google import genai 
from google.genai import types
from google.api_core import retry
from datetime import datetime
from io import BytesIO
from PIL import Image  # Now safe from namespace collision
from websockets.exceptions import ConnectionClosedError
# load_dotenv()
# # GOOGLE_API = os.getenv('GOOGLE_API_KEY')
# GMAIL_ACCT = os.getenv('GMAIL_ACCT')
# GMAIL_PASS = os.getenv('GMAIL_PASS')


# Load environment variables
try:
    GOOGLE_API = None #os.getenv('GOOGLE_API_KEY')
    if GOOGLE_API:
        GOOGLE_API_KEY = GOOGLE_API        
    else: 
        GOOGLE_API_KEY = st.secrets["google_api"]["api_key"]
except:
    print("Error")
# from langchain_openai import ChatOpenAI
# from langchain.globals import set_verbose, set_debug, set_llm_cache
# from langchain.cache import InMemoryCache  # For in-memory caching

# Set verbosity and debugging options
# set_verbose(True)  # Enables detailed logs globally
# set_debug(False)   # Disables deep debugging for now

# Set up in-memory LLM caching to avoid redundant API calls
# set_llm_cache(InMemoryCache())

try:
    favicon = "https://github.com/apod-1/ZoltarFinancial/raw/main/docs/ZoltarSurf_48x48.png"
except (KeyError, FileNotFoundError):
    favicon = st.secrets["browser"]["favicon"]

st.set_page_config(page_title="Zoltar Stock Research Agent", page_icon=favicon, layout="wide", initial_sidebar_state="collapsed")


# st.markdown("""
# <style>
#     [data-testid="collapsedControl"] {
#         display: none !important;
#     }
# </style>
# """, unsafe_allow_html=True)

hide_streamlit_style = """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    </style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)    

# These are the Python functions defined above.
# db_tools = [list_tables, describe_table, execute_query]

# instruction = """You are a helpful chatbot that can interact with an SQL database
# for a computer store. You will take the users questions and turn them into SQL
# queries using the tools available. Once you have the information you need, you will
# answer the user's question using the data returned.

# Use list_tables to see what tables are present, describe_table to understand the
# schema, and execute_query to issue an SQL SELECT query."""
col1, col2, col3 = st.columns([1, 5, 1])    


with col2:
# Streamlit UI for user input
    st.title("US Equities Zoltar Research Agent 🤖", help="I am here to help you make better decisions! Don't be shy - ask away...")


# 5.26.25 - initialize session state variables 
if 'final_agent_result' not in st.session_state:
    st.session_state.final_agent_result = ""
if 'image' not in st.session_state:
    st.session_state.image = None
if "temp_selected" not in st.session_state:
    st.session_state["temp_selected"] = "0.1 - Middle"  # or your desired default

if "top_p_selected" not in st.session_state:
    st.session_state["top_p_selected"] = "0.9 - Middle"   # or your desired default

# A little pre-work to set up what we need:

    # 1. Set up databases we'll need (5 total)

    # 2. context and metadata
    
    # 3. define functions and other tools available, including live API
    
    # 4. create sliders/selectors for tuning model parameters (to be expanded later)
    
    # 5. define and create interactive structure and guidelines / tool use instructions for agents


# Helper function to verify files and print their columns

def get_latest_file(data_dir=None, prefix=None):
    """
    Finds the latest file in a directory based on a given prefix.

    Args:
        data_dir (str): Directory to search for files. Defaults to Zoltar Financial's daily ranks directory.
        prefix (str): Prefix to filter files (e.g., "high_risk_PROD", "low_risk_PROD").

    Returns:
        str: Path to the latest file matching the prefix, or None if no valid files are found.
    """
    try:
        # Default directory setup
        if data_dir is None:
            # if os.path.exists("https://github.com/apod-1/ZoltarFinancial/main/daily_ranks"):
            #     data_dir = 'https://github.com/apod-1/ZoltarFinancial/main/daily_ranks/'
            # else:
            data_dir = '/mount/src/zoltarfinancial/daily_ranks'

        # Find files matching the prefix
        files = [f for f in os.listdir(data_dir) if f.startswith(prefix) and f.endswith(".pkl")]
        if not files:
            print(f"No valid files found for prefix '{prefix}' in directory '{data_dir}'.")
            return None

        # Find the latest file based on modification time
        latest_file = max(files, key=lambda x: os.path.getmtime(os.path.join(data_dir, x)))

        # Return the full path to the latest file
        return os.path.join(data_dir, latest_file)

    except FileNotFoundError:
        st.error("Unable to load the latest files. Please try again later.")
        return None

def get_latest_file_from_github(base_url, prefix):
    """
    Finds and downloads the latest .pkl file from a GitHub directory listing page,
    based on file name prefix, using raw file URLs.
    """
    try:
        response = requests.get(base_url)
        response.raise_for_status()
        html_content = response.text
        pattern_href = re.compile(r'href="([^"]*\.pkl)"')
        matches = pattern_href.findall(html_content)

        # --- FIXED MATCHING LOGIC ---
        pattern = re.compile(rf"^{re.escape(prefix)}(_|\..*$)")
        matching_files = [f.split('/')[-1] for f in matches if pattern.match(f.split('/')[-1])]
        # ----------------------------

        if not matching_files:
            print(f"No .pkl files with prefix '{prefix}' found at {base_url}")
            return None

        latest_file_name = max(matching_files)
        print(f"Latest .pkl file found: {latest_file_name}")

        if prefix=="fundamentals_df":
            raw_file_url = f"https://github.com/apod-1/ZoltarFinancial/main/data/{latest_file_name}"
        elif prefix=="ratings_detail_df":
            raw_file_url = f"https://github.com/apod-1/ZoltarFinancial/main/data/{latest_file_name}"
        else:
            raw_file_url = f"https://github.com/apod-1/ZoltarFinancial/main/daily_ranks/{latest_file_name}"

        response = requests.get(raw_file_url)
        response.raise_for_status()
        df = pd.read_pickle(BytesIO(response.content))
        return df

    except requests.exceptions.RequestException as e:
        print(f"Request error: {e}")
        return None
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

# def load_data(file_path):
#     return pd.read_pickle(file_path)


# def get_latest_prod_files(data_dir=None):
#     try:
#         if data_dir is None:
#             if os.path.exists(r'C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks'):
#                 data_dir = r'C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks'
#             else:
#                 data_dir = '/mount/src/zoltarfinancial/daily_ranks'
    
#         latest_files = {}
#         for category in ['high_risk', 'low_risk']:
#             files = [f for f in os.listdir(data_dir) if f.startswith(f"{category}_PROD_") and f.endswith(".pkl")]
#             if files:
#                 latest_file = max(files, key=lambda x: os.path.getmtime(os.path.join(data_dir, x)))
#                 latest_files[category] = latest_file
#             else:
#                 latest_files[category] = None

#     except FileNotFoundError:
#     #     with st.spinner("New version of Zoltar Ranks is loading. The process usually takes ~1 min to complete. Please try again..."):
#     #         sleep(60)  # Wait for 60 seconds
#         st.error("Unable to load the latest files. Please try again later.")
#     #     return None, None

#     return latest_files, data_dir

# latest_files, data_dir = get_latest_prod_files() 
# high_risk_df_long = load_data(os.path.join(data_dir, latest_files['high_risk'])) if latest_files['high_risk'] else None
# low_risk_df_long = load_data(os.path.join(data_dir, latest_files['low_risk'])) if latest_files['low_risk'] else None

# high_risk_df_long['Date'] = high_risk_df_long['Date'].astype(str)
# low_risk_df_long['Date'] = low_risk_df_long['Date'].astype(str)    



# def get_latest_prod_files(data_dir=None):
#     try:
#         if data_dir is None:
#             if os.path.exists(r'C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks'):
#                 data_dir = r'C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks'
#             else:
#                 data_dir = '/mount/src/zoltarfinancial/daily_ranks'
    
#         latest_files = {}
#         for category in ['high_risk', 'low_risk']:
#             files = [f for f in os.listdir(data_dir) if f.startswith(f"{category}_PROD_") and f.endswith(".pkl")]
#             if files:
#                 latest_file = max(files, key=lambda x: os.path.getmtime(os.path.join(data_dir, x)))
#                 latest_files[category] = latest_file
#             else:
#                 latest_files[category] = None

#     except FileNotFoundError:
#     #     with st.spinner("New version of Zoltar Ranks is loading. The process usually takes ~1 min to complete. Please try again..."):
#     #         sleep(60)  # Wait for 60 seconds
#         st.error("Unable to load the latest files. Please try again later.")
#     #     return None, None

#     return latest_files, data_dir

# latest_files, data_dir = get_latest_prod_files()

# Define paths and prefixes for different file types
# paths = [
#     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "all_high_risk_PROD", "table": "all_high_risk"},
#     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "all_low_risk_PROD", "table": "all_low_risk"},
#     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "high_risk_PROD", "table": "high_risk"},
#     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "low_risk_PROD", "table": "low_risk"},
#     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "combined_SHAP_summary_Large", "table": "shap_summary_Large"},
#     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "combined_SHAP_summary_Mid", "table": "shap_summary_Mid"},
#     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "combined_SHAP_summary_Small", "table": "shap_summary_Small"},
#     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\data", "prefix": "fundamentals_df", "table": "fundamentals"},
#     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\data", "prefix": "ratings_detail_df", "table": "ratings_detail"}
# ]

# Inspect each file and print its columns
# def inspect_files():
#     for entry in paths:
#         directory = entry["path"]
#         prefix = entry["prefix"]
#         table = entry["table"]

#         # Get latest file matching prefix
#         latest_file = get_latest_file(directory, prefix)
#         if not latest_file:
#             print(f"No file found for {prefix} in {directory}")
#             continue

#         print(f"Inspecting {latest_file} for table '{table}'...")

#         # Load the .pkl file into a DataFrame
#         try:
#             df = pd.read_pickle(latest_file)
#             print(f"Columns for table '{table}': {df.columns.tolist()}")
#             print(f"Sample data (head) for table '{table}':")
#             print(df.head())
#             print("\n")
#         except Exception as e:
#             print(f"Error loading {latest_file}: {e}")

# # Run the inspection
# inspect_files()
import random
import string
# from datetime import datetime, timedelta, date, time
from time import sleep


def random_db_filename(base_name="zoltar_financial.db"):
    name, ext = os.path.splitext(base_name)
    suffix = ''.join(random.choices(string.ascii_lowercase + string.digits, k=6))
    return f"{name}_{suffix}{ext}"

def get_sqlite_connection_with_random_on_lock(db_file, max_retries=3, retry_delay=0.5):
    for attempt in range(max_retries):
        try:
            conn = sqlite3.connect(db_file, timeout=10)
            # Try a simple operation to check if locked
            conn.execute("PRAGMA quick_check;")
            return conn, db_file
        except sqlite3.OperationalError as e:
            if "database is locked" in str(e):
                print(f"Database is locked, creating new db file with random suffix (attempt {attempt+1})...")
                db_file = random_db_filename(db_file)
                sleep(retry_delay)
            else:
                raise
    raise RuntimeError("Could not acquire database connection after multiple retries (database is locked).")


# 1. Set up databases we'll need (5 total)
# Define database connection
db_file = "zoltar_financial.db"
db_conn, db_file_used = get_sqlite_connection_with_random_on_lock(db_file)
# db_conn = sqlite3.connect(db_file)

print(f"Using database file: {db_file_used}")




# Create tables in SQLite database
def create_tables():
    with db_conn:
        db_conn.executescript("""
            -- Drop tables if they exist
            DROP TABLE IF EXISTS high_risk;
            DROP TABLE IF EXISTS low_risk;
            DROP TABLE IF EXISTS all_high_risk;
            DROP TABLE IF EXISTS all_low_risk;
            DROP TABLE IF EXISTS fundamentals;
            DROP TABLE IF EXISTS ratings_detail;                              
            DROP TABLE IF EXISTS shap_summary_Large;                              
            DROP TABLE IF EXISTS shap_summary_Mid;                              
            DROP TABLE IF EXISTS shap_summary_Small;                              
        """)
        db_conn.executescript("""
            -- Table for high-risk stocks
            CREATE TABLE IF NOT EXISTS high_risk (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Date DATETIME,
                Symbol TEXT,
                Score REAL,
                Score_Sharpe REAL,
                Score_HoldPeriod REAL,
                Close_Price REAL,
                Cap_Size TEXT,
                Sector TEXT,
                Industry TEXT,
                source TEXT
            );

            -- Table for low-risk stocks
            CREATE TABLE IF NOT EXISTS low_risk (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Date DATETIME,
                Symbol TEXT,
                Score REAL,
                Score_Sharpe REAL,
                Score_HoldPeriod REAL,
                Close_Price REAL,
                Cap_Size TEXT,
                Sector TEXT,
                Industry TEXT,
                source TEXT
            );
            -- Table for high-risk stocks
            CREATE TABLE IF NOT EXISTS all_high_risk (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Date DATETIME,
                Symbol TEXT,
                Score REAL,
                Score_Sharpe REAL,
                Score_HoldPeriod REAL,
                Close_Price REAL,
                Cap_Size TEXT,
                Sector TEXT,
                Industry TEXT,
                source TEXT
            );

            -- Table for low-risk stocks
            CREATE TABLE IF NOT EXISTS all_low_risk (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Date DATETIME,
                Symbol TEXT,
                Score REAL,
                Score_Sharpe REAL,
                Score_HoldPeriod REAL,
                Close_Price REAL,
                Cap_Size TEXT,
                Sector TEXT,
                Industry TEXT,
                source TEXT
            );

            -- Table for fundamentals
            CREATE TABLE IF NOT EXISTS fundamentals (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Symbol TEXT,
                Fundamentals_OverallRating REAL,
                total_ratings INTEGER,
                Fundamentals_Sector TEXT,
                Fundamentals_Industry TEXT,
                Fundamentals_Dividends REAL,
                Fundamentals_PE REAL,
                Fundamentals_PB REAL,
                Fundamentals_MarketCap REAL,
                Fundamentals_avgVolume2Weeks REAL,
                Fundamentals_avgVolume30Days REAL,
                Fundamentals_52WeekHigh REAL,
                Fundamentals_52WeekLow REAL,
                Fundamentals_52WeekHighDate DATE,
                Fundamentals_52WeekLowDate DATE,
                Fundamentals_Float REAL,
                Fundamentals_SharesOutstanding INTEGER,
                Fundamentals_CEO TEXT,
                Fundamentals_NumEmployees INTEGER,
                Fundamentals_YearFounded INTEGER,
                Fundamentals_ExDividendDate DATE,
                Fundamentals_PayableDate DATE,
                Fundamentals_Description TEXT
            );

            -- Table for ratings detail
            CREATE TABLE IF NOT EXISTS ratings_detail (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Symbol TEXT,
                RatingType TEXT,
                RatingText TEXT,
                RatingPublishedAt DATETIME
            );
        """)

# Helper function to find the latest file by timestamp in a directory
# def drop_tables():
#     """
#     Explicitly drop specific tables from the database.
#     """
#     tables_to_drop = ["high_risk", "low_risk", "fundamentals", "ratings_detail", "shap_summary"]
    
#     with db_conn:
#         cursor = db_conn.cursor()
#         for table in tables_to_drop:
#             print(f"Dropping table: {table}")
#             cursor.execute(f"DROP TABLE IF EXISTS {table};")
#         db_conn.commit()
#         print("Specified tables dropped successfully.")

# Function to list all tables in the database
def list_tables() -> list[str]:
    """Retrieve the names of all tables in the database."""
    st.write(' - DB CALL: list_tables()')

    cursor = db_conn.cursor()
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
    tables = cursor.fetchall()
    return [t[0] for t in tables]

# Function to describe a table's schema
def describe_table(table_name: str) -> list[tuple[str, str]]:
    """Look up the table schema.

    Returns:
      List of columns, where each entry is a tuple of (column, type).
    """
    st.write(f' - DB CALL: describe_table({table_name})')

    cursor = db_conn.cursor()
    cursor.execute(f"PRAGMA table_info({table_name});")
    schema = cursor.fetchall()
    return [(col[1], col[2]) for col in schema]

# Function to execute an SQL query
# def execute_query(sql: str) -> list[list[str]]:
#     """Execute an SQL statement, returning the results."""
#     st.write(f' - DB CALL: execute_query({sql})')

#     cursor = db_conn.cursor()
#     cursor.execute(sql)
#     return cursor.fetchall()
def execute_query(sql: str) -> list[list[str]]:
    """Execute an SQL statement, returning the results."""
    # Don't st.write here!
    cursor = db_conn.cursor()
    cursor.execute(sql)
    results = cursor.fetchall()
    # Return both the call string and results
    return {"call": f"execute_query({sql})", "results": results}

# Function to dynamically create shap_summary table based on DataFrame columns
def create_shap_table(df):
    """
    Dynamically create the shap_summary table based on DataFrame columns.
    """
    with db_conn:
        # Drop the shap_summary table if it exists
        db_conn.execute("DROP TABLE IF EXISTS shap_summary;")        
        # Dynamically generate column definitions
        columns = ", ".join([f'"{col}" REAL' for col in df.columns])
            # CREATE TABLE IF NOT EXISTS shap_summary (
        create_table_query = f"""
            CREATE TABLE {df} (
                id INTEGER PRIMARY KEY AUTOINCREMENT, 
                {columns}
            );
        """
        db_conn.execute(create_table_query)
        print("Created shap_summary table with dynamic schema.")



# Static tables schema
# STATIC_SCHEMA = {
#     'high_risk': '''
#         (id INTEGER PRIMARY KEY AUTOINCREMENT,
#         Date DATETIME, Symbol TEXT, Score REAL,
#         Score_Sharpe REAL, Score_HoldPeriod REAL,
#         Close_Price REAL, Cap_Size TEXT,
#         Sector TEXT, Industry TEXT, source TEXT)''',
#     # Add other static table definitions here
# }


# Load .pkl files into SQLite tables
# def load_data_into_db():
#     # Define paths and prefixes for different file types
#     paths = [
#         {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "all_high_risk_PROD", "table": "high_risk"},
#         {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "all_low_risk_PROD", "table": "low_risk"},
#         {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\data", "prefix": "fundamentals_df", "table": "fundamentals"},
#         {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\data", "prefix": "ratings_detail_df", "table": "ratings_detail"},
#         {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "combined_SHAP_summary_Large", "table": "shap_summary"},
#         {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "combined_SHAP_summary_Mid", "table": "shap_summary"},
#         {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "combined_SHAP_summary_Small", "table": "shap_summary"}
#     ]

#     # Iterate through paths and load each file into its corresponding table
#     for entry in paths:
#         directory = entry["path"]
#         prefix = entry["prefix"]
#         table = entry["table"]

#         # Get the latest file matching the prefix
#         latest_file = get_latest_file(directory, prefix)
#         if not latest_file:
#             print(f"No file found for {prefix} in {directory}")
#             continue

#         print(f"Loading {latest_file} into {table}...")

#         # Load .pkl file into a DataFrame
#         df = pd.read_pickle(latest_file)

#         # Handle dynamic creation for shap_summary table
#         if table == "shap_summary":
#             create_shap_table(df)

#         # Insert data into database
#         with db_conn:
#             try:
#                 df.to_sql(table, db_conn, if_exists="append", index=False)
#             except Exception as e:
#                 print(f"Error inserting into {table}: {e}")
# def load_data_into_db():
#     # Define paths and prefixes for different file types
#     paths = [
#         {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "all_high_risk_PROD", "table": "all_high_risk"},
#         {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "all_low_risk_PROD", "table": "all_low_risk"},
#         {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "high_risk_PROD", "table": "high_risk"},
#         {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "low_risk_PROD", "table": "low_risk"},
#         {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\data", "prefix": "fundamentals_df", "table": "fundamentals"},
#         {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\data", "prefix": "ratings_detail_df", "table": "ratings_detail"},
#         {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "combined_SHAP_summary_Large", "table": "shap_summary_Large"},
#         {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "combined_SHAP_summary_Mid", "table": "shap_summary_Mid"},
#         {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "combined_SHAP_summary_Small", "table": "shap_summary_Small"}
#     ]
#     # Iterate through paths and load each file into its corresponding table
#     for entry in paths:
#         directory = entry["path"]
#         prefix = entry["prefix"]
#         table = entry["table"]

#         # Get the latest file matching the prefix
#         latest_file = get_latest_file(directory, prefix)
#         if not latest_file:
#             print(f"No file found for {prefix} in {directory}")
#             continue

#         print(f"Loading {latest_file} into {table}...")

#         # Load .pkl file into a DataFrame
#         df = pd.read_pickle(latest_file)

#         # Special handling for ratings_detail table
#         if table == 'ratings_detail':
#             # Decode RatingText from bytes to UTF-8 strings
#             df['RatingText'] = df['RatingText'].apply(lambda x: x.decode('utf-8') if isinstance(x, bytes) else x)

#             # Convert RatingPublishedAt to SQLite-compatible DATETIME format
#             df['RatingPublishedAt'] = pd.to_datetime(df['RatingPublishedAt'], errors='coerce')

#         # Insert data into database
#         with db_conn:
#             try:
#                 df.to_sql(table, db_conn, if_exists="replace", index=False)
#                 print(f"Successfully inserted data into {table}.")
#                 # Verify data insertion
#                 cursor = db_conn.cursor()
#                 cursor.execute(f"SELECT COUNT(*) FROM {table};")
#                 count = cursor.fetchone()[0]
#                 print(f"Table {table} now contains {count} rows.")
#             except Exception as e:
#                 print(f"Error inserting into {table}: {e}")

# # Run setup functions
# create_tables()
# load_data_into_db()

def infer_sqlite_type(col, sample_val=None):
    """Infer SQLite type based on column name or sample value."""
    col_lower = col.lower()
    if "date" in col_lower or "time" in col_lower:
        return "DATETIME"
    if "symbol" in col_lower or "sector" in col_lower or "industry" in col_lower or "source" in col_lower:
        return "TEXT"
    if sample_val is not None:
        if isinstance(sample_val, (int, float)):
            return "REAL"
        if isinstance(sample_val, str):
            return "TEXT"
    return "REAL"

def recreate_table_from_df(conn, table, df):
    """Drop and recreate a table with columns matching the DataFrame."""
    cols = []
    for col in df.columns:
        sample_val = df[col].dropna().iloc[0] if not df[col].dropna().empty else None
        col_type = infer_sqlite_type(col, sample_val)
        cols.append(f'"{col}" {col_type}')
    schema = ", ".join(cols)
    sql = f'CREATE TABLE IF NOT EXISTS "{table}" ({schema});'
    with conn:
        conn.execute(f'DROP TABLE IF EXISTS "{table}";')
        conn.execute(sql)

def load_data_into_db():
    paths = [
        {"path": "/mount/src/zoltarfinancial/daily_ranks/", "prefix": "all_high_risk_PROD", "table": "all_high_risk"},
        {"path": "/mount/src/zoltarfinancial/daily_ranks/", "prefix": "all_low_risk_PROD", "table": "all_low_risk"},
        {"path": "/mount/src/zoltarfinancial/daily_ranks/", "prefix": "high_risk_PROD", "table": "high_risk"},
        {"path": "/mount/src/zoltarfinancial/daily_ranks/", "prefix": "low_risk_PROD", "table": "low_risk"},
        {"path": "/mount/src/zoltarfinancial/data/", "prefix": "fundamentals_df", "table": "fundamentals"},
        {"path": "/mount/src/zoltarfinancial/data/", "prefix": "ratings_detail_df", "table": "ratings_detail"},
        {"path": "/mount/src/zoltarfinancial/daily_ranks/", "prefix": "combined_SHAP_summary_Large", "table": "shap_summary_Large"},
        {"path": "/mount/src/zoltarfinancial/daily_ranks/", "prefix": "combined_SHAP_summary_Mid", "table": "shap_summary_Mid"},
        {"path": "/mount/src/zoltarfinancial/daily_ranks/", "prefix": "combined_SHAP_summary_Small", "table": "shap_summary_Small"}
    ]
    # paths = [
    #     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "all_high_risk_PROD", "table": "all_high_risk"},
    #     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "all_low_risk_PROD", "table": "all_low_risk"},
    #     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "high_risk_PROD", "table": "high_risk"},
    #     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "low_risk_PROD", "table": "low_risk"},
    #     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\data", "prefix": "fundamentals_df", "table": "fundamentals"},
    #     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\data", "prefix": "ratings_detail_df", "table": "ratings_detail"},
    #     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "combined_SHAP_summary_Large", "table": "shap_summary_Large"},
    #     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "combined_SHAP_summary_Mid", "table": "shap_summary_Mid"},
    #     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "combined_SHAP_summary_Small", "table": "shap_summary_Small"}
    # ]
    shap_tables = {"shap_summary_Large", "shap_summary_Mid", "shap_summary_Small"}

    for entry in paths:
        directory = entry["path"]
        prefix = entry["prefix"]
        table = entry["table"]
        print(f"Fetching latest file for table '{table}' with prefix '{prefix}'...")

        # Construct BASE_URL dynamically based on path
        
        latest_file = get_latest_file(directory, prefix)
        if not latest_file:
            print(f"No file found for {prefix} in {directory}")
            continue

        print(f"Loading {latest_file} into {table}...")
        df = pd.read_pickle(latest_file)

        # Special handling for ratings_detail table
        if table == 'ratings_detail':
            df['RatingText'] = df['RatingText'].apply(lambda x: x.decode('utf-8') if isinstance(x, bytes) else x)
            df['RatingPublishedAt'] = pd.to_datetime(df['RatingPublishedAt'], errors='coerce')

        # SHAP tables: drop "Feature Category" and reset index if needed
        if table in shap_tables:
            if "Feature Category" in df.columns:
                df = df.drop(columns=["Feature Category"])
            if df.index.name is not None or not df.index.equals(pd.RangeIndex(len(df))):
                df = df.reset_index()
            df = df.rename(columns={'index': 'Symbol'})

            # Debug: print DataFrame info before inserting
            #st.dataframe(df.style.format(precision=9))
            # st.write("\n--- DEBUG: DataFrame to be written to", table, "---")
            # st.write("Columns:", df.columns.tolist())
            # st.write("Dtypes:\n", df.dtypes)
            # st.write("Sample rows:\n", df.head(10))
            # st.write("Describe:\n", df.describe())
            # st.write("--- END DEBUG ---\n")

        # Insert data into database (let pandas create table and types)
        with db_conn:
            try:
                df.to_sql(table, db_conn, if_exists="replace", index=False)
                print(f"Successfully inserted data into {table}.")
                cursor = db_conn.cursor()
                cursor.execute(f"SELECT COUNT(*) FROM {table};")
                count = cursor.fetchone()[0]
                print(f"Table {table} now contains {count} rows.")
            except Exception as e:
                print(f"Error inserting into {table}: {e}")

        # Optional: read back and check first few rows
        # if table in shap_tables:
        #     df_check = pd.read_sql_query(f"SELECT * FROM {table} LIMIT 10", db_conn)
        #     print(f"\n--- DEBUG: Data read back from {table} ---")
        #     print(df_check)
        #     print("--- END DEBUG ---\n")
# Usage
# create_tables()
# load_data_into_db()

# print("Database setup complete.")

try:
    create_tables()
    load_data_into_db()
    print(f"Tables created and data loaded successfully in {db_file_used}.")
except Exception as e:
    print(f"Error during table creation or data loading: {e}")

    # Define database connection
    db_file = random_db_filename(db_file)
    # db_file = "zoltar_financial2.db"
    db_conn, db_file_used = get_sqlite_connection_with_random_on_lock(db_file)
    # db_conn = sqlite3.connect(db_file)
    
    print(f"Using database file: {db_file_used}")



# Verify high_risk table contents
# st.write("Sample Data from 'high_risk':", execute_query("SELECT * FROM high_risk LIMIT 5;"))

# Verify low_risk table contents
# st.write("Sample Data from 'low_risk':", execute_query("SELECT * FROM low_risk LIMIT 5;"))

# # Database connection
# db_file = "zoltar_financial.db"
# db_conn = sqlite3.connect(db_file)

# Test database functions
# print("Tables in Database:", list_tables())
# print("Schema for 'shap_summary':", describe_table("shap_summary"))
# print("Sample Data from 'shap_summary':", execute_query("SELECT * FROM shap_summary LIMIT 5;"))

# Define tools for chatbot
db_tools = [list_tables, describe_table, execute_query]

# Instruction for chatbot
instruction = """You are a helpful chatbot that can interact with an SQL database
for Stock trading education app. You will take the users' questions and turn them into SQL
queries using the tools available. Once you have the information you need, you will
answer the user's question using the data returned.  
high risk scores should be communicated as high Zoltar Ranks in context, and low risk scores are low Zoltar Ranks for context.  
These scores predict returns - high is for best return in next 14 days, and low is average expected return for the next 14 days. 
User is usually interested in high returns, and if stable returns are preferred, low risk scores (low zoltar rank) should be used, 
with sorting always done with highest values on top.  
If user is interested in ratings, go to ratings_detail and get necessary data (by Symbol). the RatingsPlblishedAt example format(2025-03-21T11:52:24Z) is not a timestamp format (but can extract timestamp info)


Use list_tables to see what tables are present, describe_table to understand the schema, and execute_query to issue an SQL SELECT query. When recommending an action, you have to take that action.
Be mindful of space used and limit as much as possible upfront in SQL queries output.


Here's avaliable data:
            -- Table for high-risk stocks
            CREATE TABLE IF NOT EXISTS high_risk (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Date DATETIME,
                Symbol TEXT,
                Score REAL,
                Score_Sharpe REAL,
                Score_HoldPeriod REAL,
                Close_Price REAL,
                Cap_Size TEXT,
                Sector TEXT,
                Industry TEXT,
                source TEXT
            );

            -- Table for low-risk stocks
            CREATE TABLE IF NOT EXISTS low_risk (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Date DATETIME,
                Symbol TEXT,
                Score REAL,
                Score_Sharpe REAL,
                Score_HoldPeriod REAL,
                Close_Price REAL,
                Cap_Size TEXT,
                Sector TEXT,
                Industry TEXT,
                source TEXT
            );
            -- Table for high-risk stocks
            CREATE TABLE IF NOT EXISTS all_high_risk (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Date DATETIME,
                Symbol TEXT,
                Score REAL,
                Score_Sharpe REAL,
                Score_HoldPeriod REAL,
                Close_Price REAL,
                Cap_Size TEXT,
                Sector TEXT,
                Industry TEXT,
                source TEXT
            );

            -- Table for low-risk stocks
            CREATE TABLE IF NOT EXISTS all_low_risk (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Date DATETIME,
                Symbol TEXT,
                Score REAL,
                Score_Sharpe REAL,
                Score_HoldPeriod REAL,
                Close_Price REAL,
                Cap_Size TEXT,
                Sector TEXT,
                Industry TEXT,
                source TEXT
            );

            -- Table for fundamentals
            CREATE TABLE IF NOT EXISTS fundamentals (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Symbol TEXT,
                Fundamentals_OverallRating REAL,
                total_ratings INTEGER,
                Fundamentals_Sector TEXT,
                Fundamentals_Industry TEXT,
                Fundamentals_Dividends REAL,
                Fundamentals_PE REAL,
                Fundamentals_PB REAL,
                Fundamentals_MarketCap REAL,
                Fundamentals_avgVolume2Weeks REAL,
                Fundamentals_avgVolume30Days REAL,
                Fundamentals_52WeekHigh REAL,
                Fundamentals_52WeekLow REAL,
                Fundamentals_52WeekHighDate DATE,
                Fundamentals_52WeekLowDate DATE,
                Fundamentals_Float REAL,
                Fundamentals_SharesOutstanding INTEGER,
                Fundamentals_CEO TEXT,
                Fundamentals_NumEmployees INTEGER,
                Fundamentals_YearFounded INTEGER,
                Fundamentals_ExDividendDate DATE,
                Fundamentals_PayableDate DATE,
                Fundamentals_Description TEXT
            );

            -- Table for ratings detail
            CREATE TABLE IF NOT EXISTS ratings_detail (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Symbol TEXT,
                RatingType TEXT,
                RatingText TEXT,
                RatingPublishedAt DATETIME
            
            CREATE TABLE IF NOT EXISTS shap_summary_Large (CHECK CONTENTS FOR COLUMN NAMES)
            CREATE TABLE IF NOT EXISTS shap_summary_Mid (CHECK CONTENTS FOR COLUMN NAMES)
            CREATE TABLE IF NOT EXISTS shap_summary_Small (CHECK CONTENTS FOR COLUMN NAMES)

    SHAP REASONS ARE NOT IN FUNDAMENTALS - THEY ARE LOCATED IN 3 SEPARATE DATASETS: Additionally, use database tools to examine shap_summary_Large, shap_summary_Small, shap_summary_Mid in SQLite3 database (using the database tool) that contain SHAPLEY explanations for ML results on top stocks in corresponding tables with Symbol to merge with other tabes - if not in there it is not in top stocks currently.:

    all_high_risk and all_low_risk contain intraday production runs of Zoltar Ranks (Date column), and the high_risk and low_risk without "all" contain only daily production runs (but go further back from now) - also with Date column. Unless there is a reason to look at only the most recent intraday data, there is no reason to use 'all' datasets.
    when user asks for time-related tasks, Date column should be used in congjuction with Symbol, which represent Tickers, or Stocks.
    fundamentals dataset is updated only once a day, all_ datasets contains intraday data (Date) , and low_risk and high_risk contain daily data.  
    When user wants the most recent trends, always take the mox(Date) for the answer for each Symbol, and all_ files usually provide a better answer. For long-term trends the other ones are used.
    

The tables are related to each other by Symbol, and additionally by Date if available.  data for SHAP can be merged by knowing Cap_Size in high_risk and low_risk (and all_ versions also have these).
Important: Since many dates are available for same Symbol in high_risk and low_risk data, only the latest date should be used for most queries (unless explicitly stated otherwise)
Always order by descending date first (pick only records with max date unless stated otherwise), then descending Returns for the final answer, and sometimes in order of descending dividends.
When user requests Top stocks, they mean stocks with highest expected returns (highest Zoltar Ranks - low or high, depending on preference), and prefers
Ensure final answer meets all criteria set by the user request, and the answer contains non-duplicate symbols that look at the most recent data point, and mention the date used in the answer.
When stocks symbols are presented, also mention current price, and a few ratings/explanations, and when some information is missing, work with the information that is available (fundamentals and SHAP data could be missing).
When the user asks for Top stocks without mentioning High or Low, assume High Zoltar Ranks is needed.
When user asks for reasons for stocks being selected, refer to SHAP datasets using Cap_Size and Symbol (can check all 3 by Symbol)
When user asks for alpha, the comparison with SPY returns needs to be made.

User prefers the answer in a table format with relevant statistics, and a summary brief.  Response always ends with the phrase 'May the riches be with you...'

"""



# 5.7.25 - new from jupyter notebook


# # Main section with using genai gemini model
# is_retriable = lambda e: (isinstance(e, genai.errors.APIError) and e.code in {429, 503})

# if not hasattr(genai.models.Models.generate_content, '__wrapped__'):
#   genai.models.Models.generate_content = retry.Retry(
#       predicate=is_retriable)(genai.models.Models.generate_content)
# client = genai.Client(api_key=GOOGLE_API_KEY)



# # Model configuration for Gemini-2.0-Flash
# model_config = types.GenerateContentConfig(
#     temperature=0.1,
#     top_p=0.95,
#     system_instruction=instruction,
#     tools=db_tools,
# )

#5.26.25 -  Agent settings
st.sidebar.header("Agent Configuration")
st.sidebar.write("News sources selection:")

# Create 3 columns in the sidebar
col1side, col2side = st.sidebar.columns(2)

with col1side:
    google_trends = st.checkbox("Google Trends", value=True)
    stocktwits = st.checkbox("StockTwits", value=True)
    nasdaq = st.checkbox("Zacks", value=True)

with col2side:
    reddit = st.checkbox("Reddit", value=True)
    sentimentrader = st.checkbox("Yahoo Finance", value=True)
    tipranks = st.checkbox("TipRanks", value=True)

selected_sources = []
if google_trends: selected_sources.append("Google Trends (https://trends.google.com/)")
if stocktwits: selected_sources.append("StockTwits (https://stocktwits.com/")
if sentimentrader: selected_sources.append("Yahoo (https://finance.yahoo.com/)")
if tipranks: selected_sources.append("TipRanks (https://www.tipranks.com/")
if nasdaq: selected_sources.append("Zacks (https://www.zacks.com/)")
if reddit: selected_sources.append("Reddit (https://www.reddit.com/)")
    
source_str = ", ".join(selected_sources) if selected_sources else "no sources selected"




st.sidebar.write("Visualization selection:")

# Create 3 columns in the sidebar
col1side, col2side = st.sidebar.columns(2)

with col1side:
    Pie_chart = st.checkbox("Pie Chart", value=True)
    Return_hold = st.checkbox("Returns", value=True)

with col2side:
    low_ranks_trend = st.checkbox("Ranks Trend", value=True)
    recommendations_table = st.checkbox("Summary", value=False)


# Map checkbox variables to their prompt instructions
viz_instructions = []

if Pie_chart:
    viz_instructions.append("- Industry: Pie Chart of Industries of selected stocks (add data for this one last)")
if Return_hold:
    viz_instructions.append("- Expected Returns: line chart for each of the selected stocks with two points for each - first point starting at (0,0) and second point X is number of days to hold (Score_HoldPeriod in high_risk and all_high_risk tables) vs High Zoltar Rank (y-axis), making starting point for x-axis max(Date) and iterating days forward from that point.")
if low_ranks_trend:
    viz_instructions.append("- Low Zoltar Rank Over Time: a pretty line chart of Low Zoltar Rank of each stock over time")
if recommendations_table:
    viz_instructions.append("- Recommendations: Table of model recommendations for each stock")

# Join instructions for prompt
if viz_instructions:
    viz_section = "\n".join(viz_instructions)
else:
    viz_section = "- No visualizations selected."


# # Sidebar sliders for tuning model parameters
# st.sidebar.header("Model Configuration")
# temperature = st.sidebar.slider(
#     "Temperature", 
#     min_value=0.0, 
#     max_value=1.0, 
#     value=0.1,  # Default value
#     step=0.05
# )
# top_p = st.sidebar.slider(
#     "Top-p", 
#     min_value=0.0, 
#     max_value=1.0, 
#     value=0.95,  # Default value
#     step=0.05
# )

st.sidebar.header("Model Configuration")

# --- Define levels and their mappings ---
temp_levels = [("0.0 - Exact", 0.0), ("0.1 - Middle", 0.1), ("1.0 - Wild", 1.0)]
top_p_levels = [("0.5 - Wild", 0.7), ("0.9 - Middle", 0.9), ("1.0 - Exact", 1.0)]

# --- Helper function for segmented buttons ---
def segmented_buttons(label, levels, key_prefix):
    cols = st.sidebar.columns(len(levels))
    selected = st.session_state.get(f"{key_prefix}_selected", levels[0][0])
    for i, (level, _) in enumerate(levels):
        button_kwargs = {}
        if selected == level:
            button_kwargs["type"] = "primary"
        if cols[i].button(level, key=f"{key_prefix}_{level}", **button_kwargs):
            st.session_state[f"{key_prefix}_selected"] = level
            selected = level
    return dict(levels)[selected]

# --- Render segmented buttons ---
st.sidebar.write("Temperature setting:")
temperature = segmented_buttons("Temperature Level", temp_levels, "temp")
# st.sidebar.markdown(
#     f"**Temperature:** {st.session_state['temp_selected']} ({temperature})\n\n"
# )

st.sidebar.write("Top-p setting:")

top_p = segmented_buttons("Top-p Level", top_p_levels, "top_p")

# --- Show current values ---
# st.sidebar.markdown(
#     f"**Top-p:** {st.session_state['top_p_selected']} ({top_p})"
# )

# Show the selected values (for debugging/demo)
# st.sidebar.write(f"Temperature value: {temperature}")
# st.sidebar.write(f"Top-p value: {top_p}")



# Model configuration for Gemini-2.0-Flash
model_config = types.GenerateContentConfig(
    temperature=temperature,
    top_p=top_p,
    system_instruction=instruction,
    tools=db_tools,
)
# Initialize Google GenAI client (ensure GOOGLE_API_KEY is set)
# from google import genai

client = genai.Client(api_key=GOOGLE_API_KEY)

# Start a chat with automatic function calling enabled
chat = client.chats.create(
    model="gemini-2.0-flash",
    config=model_config,
)
result=None

# 5.24 - helper functions
def to_json_serializable(obj):
    if isinstance(obj, str):
        return obj
    try:
        return json.dumps(obj, default=str)
    except Exception as e:
        return str(obj)

def is_blank_png(img_bytes):
    # Quick check for empty or very small files
    if not img_bytes or len(img_bytes) < 8500:
        return True

    try:
        with Image.open(BytesIO(img_bytes)) as img:
            img = img.convert("RGBA")  # Ensure 4 channels
            extrema = img.getextrema()  # Returns (min, max) for each channel

            # Check if all pixels are fully transparent
            if extrema[3] == (0, 0):
                return True

            # Check if all pixels are white (255,255,255,255)
            if all(channel == (255, 255) for channel in extrema):
                return True

            # For grayscale/other modes, check if all pixels are same value
            if all(e[0] == e[1] for e in extrema):
                return True

            # Optionally, check if all pixels are the same color
            if len(set(img.getdata())) == 1:
                return True

    except Exception as e:
        # If PIL fails to open, treat as blank/invalid
        print(f"Image validation error: {e}")
        return True

    return False

import sys
def debug_payload(message):
    try:
        msg_str = to_json_serializable(message)
        if isinstance(msg_str, str):
            size = len(msg_str.encode('utf-8'))
        else:
            size = sys.getsizeof(msg_str)
        print(f"Payload size: {size} bytes")
        if size > 1_000_000:
            st.warning("Payload is very large and may be rejected by Gemini API.")
    except Exception as e:
        st.write(f"Could not serialize payload: {e}")


def prepare_image_for_gemini(image_path):
    if not os.path.exists(image_path):
        return None
    with open(image_path, "rb") as f:
        img_bytes = f.read()
    if is_blank_png(img_bytes):
        return None
    # Downscale if too large
    if len(img_bytes) > 2_000_000:
        img = Image.open(BytesIO(img_bytes))
        img.thumbnail((800, 800))
        buffer = BytesIO()
        img.save(buffer, format="PNG", optimize=True)
        img_bytes = buffer.getvalue()
    img_b64 = base64.b64encode(img_bytes).decode("utf-8")
    return img_b64

# # Example query to chatbot
# response = chat.send_message("what stocks have highest low Zoltar Rank, averaged over the last 5 data points? put in a table with Low and High Zoltar Ranks shown.")
# # Streamlit UI for user input
# st.title("Chatbot Query Interface")


with col2:
   
    # Input text box for the user's query
    user_query = st.text_input(
        "",
        # value="Provide stocks with best expected returns (latest date) and a dividend yield above 5, with dividend date coming up within a month from now.  create a table and a summary.",
        value="Best stocks to get now?",
        help="Ask about best stocks, dividends, sectors, explanations (anything stocks related)"
        ,placeholder = "Ask your stock-related question..."
    )
    # Submit button
    if st.button("Submit Query"):
        # Send the query to the chatbot
        #response = chat.send_message(user_query)
    
        # Display the chatbot's response
        #st.write("Chatbot Response:")
        #st.write(response.text)
    
        # Optional: Print chat history for debugging
        def print_chat_turns(chat):
            """Prints out each turn in the chat history, including function calls and responses."""
            for event in chat.get_history():
                st.write(f"{event.role.capitalize()}:")
    
                # for part in event.parts:
                #     if txt := part.text:
                #         st.write(f'  "{txt}"')
                #     elif fn := part.function_call:
                #         args = ", ".join(f"{key}={val}" for key, val in fn.args.items())
                #         st.write(f"  Function call: {fn.name}({args})")
                #     elif resp := part.function_response:
                #         st.write("  Function response:")
                #         st.write(resp.response['result'])
                for part in event.parts:
                    if code := part.executable_code:
                        st.markdown(f"### Code\n``````")
                
                    elif result := part.code_execution_result:
                        st.markdown(f"### Result: {result.outcome}\n``````")
                
                    elif img := part.inline_data:
                        try:
                            # Validate and display the image
                            image = Image.open(io.BytesIO(img.data))
                            st.image(image, caption="Generated Image")
                        except Exception as e:
                            st.error(f"Error displaying image: {e}")
                st.write()
    
        # # Display chatbot response
        # # st.write("Chatbot Response:", response)
        # st.write("Chatbot Response:")
        # st.write(f"\n{response.text}")
        # # Print chat turns (optional)
        # print_chat_turns(chat)
    
    
        class DefaultAPI:
            def execute_query(self, sql):
                # Mock result for demonstration purposes
                return {
                    "result": [
                        ["2025-04-07 00:00:00", "MNSO", 0.0514, 16.61, "Large"],
                        ["2025-03-24 22:23:03", "SMCI", 0.0383, 41.42, "Mid"],
                        ["2025-03-21 19:20:35", "HPQ", 0.0259, 28.68, "Large"],
                        ["2025-03-21 19:20:35", "SMCI", 0.0257, 42.42, "Mid"],
                        ["2025-03-26 19:10:39", "MIRM", 0.0257, 46.28, "Small"]
                    ]
                }
        
        # Initialize default_api
        default_api = GOOGLE_API_KEY #DefaultAPI() 
        
     #     return all_responses
        async def handle_response(stream, tool_impl=None):
            """Stream output and handle any tool calls during the session."""
            all_responses = []
        
            async for msg in stream.receive():
                all_responses.append(msg)
        
                if text := msg.text:
                    # Output any text chunks that are streamed back.
                    if len(all_responses) < 2 or not all_responses[-2].text:
                        # Display a header if this is the first text chunk.
                        st.markdown('### Text')
                    st.write(text)
        
                elif tool_call := msg.tool_call:
                    # Handle tool-call requests.
                    for fc in tool_call.function_calls:
                        st.markdown('### Tool call')
        
                        # Execute the tool and collect the result to return to the model.
                        if callable(tool_impl):
                            try:
                                result = tool_impl(**fc.args)
                            except Exception as e:
                                result = str(e)
                        else:
                            result = 'ok'
        
                        tool_response = types.LiveClientToolResponse(
                            function_responses=[types.FunctionResponse(
                                name=fc.name,
                                id=fc.id,
                                response={'result': result},
                            )]
                        )
                        await stream.send(input=tool_response)
        
                elif msg.server_content and msg.server_content.model_turn:
                    for part in msg.server_content.model_turn.parts:
                        #st.write("Available attributes in Part:", dir(part))  # Debugging
        
                        if code := part.executable_code:
                            st.markdown("### Code")
                            st.code(code.code)
        
                            # Dynamically execute provided code
                            try:
                                exec_globals = {
                                    "pd": pd,
                                    "sns": sns,
                                    "plt": plt,
                                    "base64": base64,
                                    "BytesIO": BytesIO,
                                    "default_api": default_api,  # Pass default_api into context
                                }
                                exec(code.code, exec_globals)
        
                                # Decode base64 string and display image
                                if "image_base64" in exec_globals:
                                    image_base64 = exec_globals["image_base64"]
                                    img_bytes = base64.b64decode(image_base64)
                                    st.image(img_bytes, caption="Generated Plot", use_column_width=True)
                                else:
                                    st.warning("No plot was generated.")
                            except Exception as e:
                                st.error(f"Error executing code: {e}")
                        elif text := part.text:  # Fallback for text-based instructions
                            st.markdown("### Instructions for Plotting")
                            st.write(text)
        
            return all_responses
    
        
    #5.4.25 additions
    
        import base64
        #!pip install -U textblob
        #from textblob import TextBlob
        
        # Global state variable
        global_state = {
            "collected_text": "",
            "tool_calls": [],
            "code_results": [],
            "images": []
        }
        
        # Counter to track updates
        update_counter = 0
        
        def update_state(key, value):
            """Update the global state."""
            global global_state
            global_state[key] = value
            
        # def update_state(key, value):
        #     global global_state
        #     # Replace entire value instead of appending
        #     global_state[key] = value if not isinstance(value, list) else [value[-1]]      
        
        # def display_state():
        #     """Refresh the display with the latest state."""
        #     #st.rerun()
        #     #clear_output(wait=True)  # Clear previous output
        
        #     # Display images (latest only)
        #     if global_state["images"]:
        #         for img in global_state["images"]:
        #             st.image(img)
        
        #     # Display collected text
        #     if global_state["collected_text"].strip():
        #         st.markdown(f"### Text\n\n{global_state['collected_text'].strip()}")
        
        #     # Display tool calls (latest only)
        #     #if global_state["tool_calls"]:
        #      #   for tool_call in global_state["tool_calls"]:
        #       #      display(Markdown(f"### Tool Call\n\n{tool_call}"))
        
        #     # Display code results (latest only)
        #     if global_state["code_results"]:
        #         for code_result in global_state["code_results"]:
        #             st.markdown(f"### Code Result\n\n{code_result}")
        placeholder_container = st.empty()  # Master container for refreshable content
        
        # def display_state():
        #     """Dynamic refresh using placeholder replacement"""
        #     with placeholder_container.container():
        #         # Clear previous content
        #         #st.empty()  # Creates temporary empty space
                
        #         # Images with auto-clear
        #         if global_state["images"]:
        #             img_placeholder = st.empty()
        #             with img_placeholder:
        #                 for img in global_state["images"]:
        #                     # Your existing image handling logic
        #                     if isinstance(img, str):
        #                         img_bytes = base64.b64decode(img)
        #                         image = Image.open(BytesIO(img_bytes))
        #                         st.image(image)
        #                     elif isinstance(img, (bytes, bytearray)):
        #                         image = Image.open(BytesIO(img))
        #                         st.image(image)
        #                     elif isinstance(img, Image.Image):
        #                         st.image(img)
        #             img_placeholder.empty()  # Clear after render
        
        #         # Text with incremental updates  
        #         if global_state["collected_text"]:
        #             text_placeholder = st.empty()
        #             cleaned_text = "\n".join([
        #                 line for line in global_state["collected_text"].split("\n")
        #                 if "End of User Query" not in line
        #             ])
        #             text_placeholder.markdown(f"**Analysis**\n\n{cleaned_text}")    
        def is_base64_bytes(data):
            # PNG header is: b'\x89PNG\r\n\x1a\n'
            if data.startswith(b'\x89PNG\r\n\x1a\n'):
                return False
            # If it's all ASCII and decodes cleanly, likely base64
            try:
                base64.b64decode(data, validate=True)
                return True
            except Exception:
                return False
            
        def display_state():
            """Dynamic refresh using placeholder replacement"""
            with placeholder_container.container():
                # Images with auto-clear
                if global_state["images"]:
                    img_placeholder = st.empty()
                    with img_placeholder:
                        for img in global_state["images"]:
                            try:
                                # If it's a base64 string
                                if isinstance(img, str):
                                    img_bytes = base64.b64decode(img.strip())
                                    if img_bytes and len(img_bytes) > 0:
                                        st.image(img_bytes)
                                    else:
                                        st.warning("Decoded image is empty.")
                                # If it's raw bytes or bytearray
                                elif isinstance(img, (bytes, bytearray)):
                                    if img and len(img) > 0:
                                        if is_base64_bytes(img):
                                            print("Detected base64-encoded bytes, decoding...")
                                            img_bytes = base64.b64decode(img.strip())
                                            st.image(img_bytes)
                                        else:
                                            st.image(img)
                                    else:
                                        st.warning("Image bytes are empty.")
                                # If it's a PIL Image
                                elif isinstance(img, Image.Image):
                                    st.image(img)
                                else:
                                    st.warning(f"Unsupported image type: {type(img)}")
                            except Exception as e:
                                st.error(f"Could not display image: {e}")
                    img_placeholder.empty()  # Clear after render
        
                # Text with incremental updates  
                if global_state["collected_text"]:
                    text_placeholder = st.empty()
                    cleaned_text = "\n".join([
                        line for line in global_state["collected_text"].split("\n")
                        if "End of User Query" not in line
                    ])
                    text_placeholder.markdown(f"**Analysis**\n\n{cleaned_text}")
    
    
        def display_state():
            """Dynamic refresh using placeholder replacement"""
            with placeholder_container.container():
                # --- Display Images FIRST ---
                if global_state["images"]:
                    for img in global_state["images"]:
                        try:
                            # If it's a base64 string
                            if isinstance(img, str):
                                img_bytes = base64.b64decode(img.strip())
                                if img_bytes and len(img_bytes) > 0:
                                    st.image(img_bytes)
                                else:
                                    st.warning("Decoded image is empty.")
                            # If it's raw bytes or bytearray
                            elif isinstance(img, (bytes, bytearray)):
                                if img and len(img) > 0:
                                    if is_base64_bytes(img):
                                        print("Detected base64-encoded bytes, decoding...")
                                        img_bytes = base64.b64decode(img.strip())
                                        st.image(img_bytes)
                                    else:
                                        st.image(img)
                                else:
                                    st.warning("Image bytes are empty.")
                            # If it's a PIL Image
                            elif isinstance(img, Image.Image):
                                st.image(img)
                            else:
                                st.warning(f"Unsupported image type: {type(img)}")
                        except Exception as e:
                            st.error(f"Could not display image: {e}")
        
                # --- Then Display Text and Other Content ---
                if global_state["collected_text"]:
                    text_placeholder = st.empty()
                    cleaned_text = "\n".join([
                        line for line in global_state["collected_text"].split("\n")
                        if "End of User Query" not in line
                    ])
                    text_placeholder.markdown(f"**Analysis**\n\n{cleaned_text}")
    
                    
        def format_global_state(global_state):
            """Convert global_state into a structured string for agent prompts."""
            sections = []
    
            if global_state["images"]:
                for img in global_state["images"]:
                    print("img type:", type(img))
                    if isinstance(img, str):
                        print("First 100 chars:", img[:100])
                    elif isinstance(img, (bytes, bytearray)):
                        print("First 20 bytes:", img[:20])
                        print("Bytes length:", len(img))
                    else:
                        print("img is not str or bytes!")
            
                    # Handle base64 strings
                    if isinstance(img, str):
                        try:
                            img_bytes = base64.b64decode(img.strip())
                            st.image(img_bytes)  # Streamlit can handle PNG/JPEG bytes directly
                        except Exception as e:
                            st.error(f"Base64 decode error: {str(e)}")
            
                    # Handle raw bytes
                    elif isinstance(img, (bytes, bytearray)):
                        try:
                            if is_base64_bytes(img):
                                print("Detected base64-encoded bytes, decoding for display...")
                                img_bytes = base64.b64decode(img.strip())
                                st.image(img_bytes)
                            else:
                                st.image(img)
                        except Exception as e:
                            st.error(f"Bytes conversion error: {str(e)}")
            
                    # Handle PIL Images directly
                    elif isinstance(img, Image.Image):
                        st.image(img)
                        
            # Text Content
            if global_state["collected_text"]:
                sections.append(f"## Collected Text\n{global_state['collected_text']}")
            
            # Tool Calls
            if global_state["tool_calls"]:
                tool_list = "\n".join(f"- {call}" for call in global_state["tool_calls"])
                sections.append(f"## Tool Calls\n{tool_list}")
            
            # Code Results
            if global_state["code_results"]:
                code_list = "\n".join(f"- {result}" for result in global_state["code_results"])
                sections.append(f"## Code Execution Results\n{code_list}")
            
            # Images (reference filenames)
            # if global_state["images"]:
            #     img_list = "\n".join(f"- Image saved as: plot_{i+1}.png" for i in range(len(global_state["images"])))
            #     sections.append(f"## Generated Visualizations\n{img_list}")
    
            return "\n\n".join(sections)
        def format_global_state_notool(global_state):
            """Convert global_state into a structured string for agent prompts, omitting tool calls."""
            sections = []
            
            # Text Content (remove tool call lines)
            if global_state["collected_text"]:
                filtered_text = "\n".join(
                    line for line in global_state["collected_text"].splitlines()
                    if not line.strip().lower().startswith("tool call") and "db call" not in line.lower()
                )
                sections.append(f"## Collected Text\n{filtered_text}")
            
            # Code Results
            if global_state["code_results"]:
                code_list = "\n".join(f"- {result}" for result in global_state["code_results"])
                sections.append(f"## Code Execution Results\n{code_list}")
            
            # Images (reference filenames or other info if needed)
            # If you want to show info about images, you can uncomment below:
            # if global_state["images"]:
            #     img_list = "\n".join(f"- Image saved as: plot_{i+1}.png" for i in range(len(global_state["images"])))
            #     sections.append(f"## Generated Visualizations\n{img_list}")
            
            return "\n\n".join(sections)   
    
    
    
        async def handle_response_refresh(stream, tool_impl=None):
            """Stream output and handle any tool calls during the session."""
            global update_counter  # Use the global counter for tracking updates
            all_responses = []
            collected_text = ""  # Collect all text responses
            tool_call_results = []  # Temporary list for tool call results
            code_results = []  # Temporary list for code results
            images = []  # Temporary list for inline images
            MAX_BYTES = 1000000  # Leave 20% buffer
            current_size = 0
            retries = 3
            backoff = 1  # seconds                

            while retries > 0:
                try:
        
                    async for msg in stream.receive():
                        all_responses.append(msg)
                        msg_size = len(str(msg).encode('utf-8'))
                        if current_size + msg_size > MAX_BYTES:
                            print("Approaching size limit - truncating response")
                            #break
                        current_size += msg_size
                        
                        if text := msg.text:
                            # Collect text chunks into a single string
                            collected_text += text + " "
                            update_state("collected_text", collected_text)  # Update state with new text
                
                        # elif tool_call := msg.tool_call:
                        #     # Handle tool-call requests
                        #     tool_call_results = []  # Reset temporary list for tool calls
                        #     for fc in tool_call.function_calls:
                        #         st.markdown('### Tool call')
                
                        #         # Execute the tool and collect the result to return to the model
                        #         if callable(tool_impl):
                        #             try:
                        #                 result = tool_impl(**fc.args)
                        #             except Exception as e:
                        #                 result = str(e)
                        #         else:
                        #             result = 'ok'
                        elif tool_call := msg.tool_call:
                            tool_call_results = []
                            # for fc in tool_call.function_calls:
                            #     if callable(tool_impl):
                            #         try:
                            #             result = tool_impl(**fc.args)
                            #             # If result is a dict with 'call' and 'results'
                            #             if isinstance(result, dict) and 'call' in result:
                            #                 tool_call_results.append(result['call'])
                            #                 # Optionally, store the actual results somewhere else
                            #                 code_results.append(result['results'])
                            #             else:
                            #                 tool_call_results.append(str(result))
                            #         except Exception as e:
                            #             tool_call_results.append(str(e))
                            #     else:
                            #         tool_call_results.append('ok')
                            # # update_state("tool_calls", tool_call_results)    
                            #     tool_response = types.LiveClientToolResponse(
                            #         function_responses=[types.FunctionResponse(
                            #             name=fc.name,
                            #             id=fc.id,
                            #             response={'result': result},
                            #         )]
                            #     )
                            #     await stream.send(input=tool_response)
                
                            #     # Add result to temporary list
                            #     tool_call_results.append(result)
                            for fc in tool_call.function_calls:
                                if callable(tool_impl):
                                    try:
                                        result = tool_impl(**fc.args)
                                        if isinstance(result, dict) and 'call' in result:
                                            tool_call_results.append(result['call'])
                                            code_results.append(result['results'])
                                        else:
                                            tool_call_results.append(str(result))
                                    except Exception as e:
                                        result = str(e)
                                        tool_call_results.append(result)
                                else:
                                    result = 'ok'
                                    tool_call_results.append(result)
                            
                                # tool_response = types.LiveClientToolResponse(
                                #     function_responses=[types.FunctionResponse(
                                #         name=fc.name,
                                #         id=fc.id,
                                #         response={'result': result},
                                #     )]
                                # )
                                tool_response = types.LiveClientToolResponse(
                                    function_responses=[types.FunctionResponse(
                                        name=fc.name,
                                        id=fc.id,
                                        response={
                                            'result': json.dumps(result)  # Serialize to string
                                        }
                                    )]
                                )                                
                                await stream.send(input=tool_response)        
                            # Replace previous tool calls with the latest ones
                            update_state("tool_calls", tool_call_results)
                            # update_state("tool_calls", [tool_call_results[-1]])
                
                        elif msg.server_content and msg.server_content.model_turn:
                            # Handle code execution results and inline images
                            code_results = []  # Reset temporary list for code results
                            images = []  # Reset temporary list for inline images
                
                            for part in msg.server_content.model_turn.parts:
                                if code := part.executable_code:
                                    code_results.append(code)
                
                                elif result := part.code_execution_result:
                                    code_results.append(result.outcome)
                
                                elif img := part.inline_data:
                                    images.append(img.data)
                                # Save the first image (or all images)
                                # if images:
                                #     first_img = images[0]
                                #     # Check if base64 string or bytes
                                #     if isinstance(first_img, str):
                                #         img_bytes = base64.b64decode(first_img)
                                #     else:
                                #         img_bytes = first_img  # assume bytes
                            
                                #     with open("stock_price_plot.png", "wb") as f:
                                #         f.write(img_bytes)
                                #     print("Image saved as stock_price_plot.png")
                                #     update_state("images", images)
                                if images:
                                    first_img = images[0]
                                    img_bytes = None
                                    # String: decode base64
                                    if isinstance(first_img, str):
                                        img_bytes = base64.b64decode(first_img.strip())
                                    # Bytes: check if base64 or PNG
                                    elif isinstance(first_img, (bytes, bytearray)):
                                        if is_base64_bytes(first_img):
                                            print("Detected base64-encoded bytes, decoding...")
                                            img_bytes = base64.b64decode(first_img.strip())
                                            print(img_bytes)
                                        else:
                                            img_bytes = first_img
                                    else:
                                        print(f"Unsupported image type: {type(first_img)}")
                                        img_bytes = b""
                                
                                    # print(f"Type: {type(img_bytes)}, Length: {len(img_bytes)}")
                                    # print("First 20 bytes:", img_bytes[:20])
                                    # print("Last 20 bytes:", img_bytes[-20:])
                                
                                    # Save and validate as before
                                    if img_bytes and len(img_bytes) > 8500:
                                        try:
                                            with open("stock_price_plot.png", "wb") as f:
                                                f.write(img_bytes)
                                            print("Image saved as stock_price_plot.png")
                                            # Try to open with PIL to check validity
                                            image = Image.open(BytesIO(img_bytes))
                                            image.verify()
                                            print("Image is valid!")
                                        except Exception as e:
                                            print(f"Invalid image data: {e}")
                                        update_state("images", images)
                                        st.session_state.image = img_bytes

                                    else:
                                        print("No valid image bytes to save.")
                                
                            # Replace previous code results and images with the latest ones
                            update_state("code_results", code_results)
                
                        # Increment counter and refresh display every other message
                        update_counter += 1
                        if update_counter % 2 == 0:  # Refresh after every second message
                            display_state()
                
                    # Display concatenated text at the end (final refresh)
                    if collected_text.strip():
                        update_state("collected_text", collected_text.strip())  # Update final text state
                        display_state()  # Refresh final state dynamically
                
                    print()
                    return all_responses    

                except (ConnectionResetError, ConnectionClosedError) as e:
                    print(f"Connection error: {e}, retries left: {retries}")
                    await asyncio.sleep(backoff)
                    retries -= 1
                    backoff *= 2
                    continue
            raise RuntimeError("Max retries exceeded")    
        # with open("stock_price_plot.png", "rb") as f:
        #     data = f.read()
        # print("First 8 bytes:", data[:8])
        # print("Last 8 bytes:", data[-8:])
        # print("File size:", len(data))
        
        model = 'gemini-2.0-flash-exp'
        live_client = genai.Client(api_key=GOOGLE_API_KEY,
                                   http_options=types.HttpOptions(api_version='v1alpha'))
        
        # Wrap the existing execute_query tool you used in the earlier example.
        execute_query_tool_def = types.FunctionDeclaration.from_callable(
            client=live_client, callable=execute_query)
        
        # Provide the model with enough information to use the tool, such as describing
        # the database so it understands which SQL syntax to use.
        # sys_int = """Your role is this: You are a database interface. Use the `execute_query` function
        # to answer the users questions by looking up information in the database, running any necessary queries and responding to the user.
        # Be mindful of systme limitation with large output: Example why: ConnectionClosedError: sent 1009 (message too big) frame with 2275770 bytes exceeds limit of 1048576 bytes; no close frame received    
        
        # You need to look up table schema using sqlite3 syntax SQL, then once an answer is found be sure to tell the user. 
        # Important: If the user is requesting an action, you must also execute the actions.
        
        # """
        # sys_int = """Your role is this: You are a database interface. Use the `execute_query` function
        # to answer the user's questions by looking up information in the database, running any necessary queries, and responding to the user.
        # Important: If the user requests a visualization (e.g., a Seaborn chart), you must generate and provide Python code that can be executed directly to create the plot. Ensure the code uses Pandas for data manipulation and Seaborn for plotting, and that it includes all necessary imports. The generated code should be executable without modification.
    
        # """
    
        sys_int = instruction + """
        Your role is this: You are a database interface. Use the [execute_query_tool_def.to_json_dict()] to understand the database/table contents and be able to pull relevant information from the tables.
        and then to answer the user's questions by looking up information in the database, running any necessary queries, and responding to the user. 
        Provide a comprehensive report on each of the selected stocks with data available on the database to be used by subsequent agents to summarize further.
        After the stock symbols of interest are known, include all company information on them form Zoltar Ranks Database, including descriptions, sector, P/E, Dividends, 52Week highs and Lows, Ratings, and other fundamentals info.
        If you recommend an action, you must take that action.
        """
        #Important: If the user requests a visualization (e.g., a Seaborn chart), you must generate Python code that uses Pandas for data manipulation and Seaborn for plotting. Ensure that the generated code includes all necessary imports and replaces plt.show() with logic to return a base64-encoded string of the plot image.
        # Replace plt.show():
            
        #     plt.show() is used for GUI-based backends and does not work in Streamlit.
            
        #     Instead, use st.pyplot(fig) where fig is a Matplotlib Figure object.
    
        config = {
            "response_modalities": ["TEXT"],
            "system_instruction": {"parts": [{"text": sys_int}]},
            "tools": [
                {"code_execution": {}},
                {"function_declarations": [execute_query_tool_def.to_json_dict()]},
            ],
        }
    
        config = {
            "response_modalities": ["TEXT"],
            "system_instruction": sys_int,
            "tools": [
                {"code_execution": {}},
                {"function_declarations": [execute_query_tool_def.to_json_dict()]},
                types.Tool(google_search=types.GoogleSearch())  # Add the Google Search tool
            ],
            "temperature": temperature,
            "top_p": top_p,
        }
    
        async def main(user_query):
            result=None
            async with live_client.aio.live.connect(model=model, config=config) as session:
                placeholder_container = st.empty()  # Master container for refreshable content
                # st.toast("AGENT 1...ZOLTAR DATABASE", icon="⏳")  # Shows a floating toast message
                agent1_toast = st.toast("AGENT 1...ZOLTAR DATABASE", icon="⏳")
                # sleep(30)
                message = user_query #"Can you figure out the number of orders that were made by each of the staff?"
                print(f"> {message}\n")
                await session.send(input=to_json_serializable(message), end_of_turn=True)
                all_responses = await handle_response_refresh(session, tool_impl=execute_query)
                agent_result = "\n".join(msg.text for msg in all_responses if msg.text)            
                formatted_state = format_global_state(global_state)
                agent1_toast.toast("AGENT 1...ZOLTAR DATABASE", icon="✅")
                agent2_toast = st.toast("AGENT 2...NEWS ARTICLES", icon="⏳")
                # st.toast("AGENT 2...NEWS ARTICLES", icon="⏳")  # Shows a floating toast message
                # sleep(30)
                message = f"Search for latest News and analyze Sentiment using types.Tool(google_search=types.GoogleSearch() tool that you have. When searching, only look at the sources specifically selected by the user: {source_str}. Create a table with best 3 links for detailed search, related to the stocks the user asked about found from Zoltar Ranks Database for stocks found by prior agent. Here is the result of the first agent findings: {agent_result}"
                print(f"> {message}\n")
                await session.send(input=to_json_serializable(message), end_of_turn=True)
                all_responses2 = await handle_response_refresh(session, tool_impl=execute_query)
                agent_result2 = "\n".join(msg.text for msg in all_responses2 if msg.text)  
                formatted_state = format_global_state(global_state)
                agent1_toast.toast("AGENT 1...ZOLTAR DATABASE", icon="✅")
                agent2_toast.toast("AGENT 2...NEWS ARTICLES", icon="✅")
                agent3_toast = st.toast("AGENT 3...OVERVIEW PLOTS", icon="⏳")
                # st.toast("AGENT 3...OVERVIEW PLOTS", icon="⏳")  # Shows a floating toast message
                # sleep(30)

                # message = f"""Use the result of the first agent findings: {agent_result}. ** end of first agent result ** 
                #      Your task is to create a seaborn plot.
                #      You can interact with Zoltar SQL database for Stock trading education app using [execute_query_tool_def.to_json_dict()] tool and should become an expert on the contents of the database and the formats of all variables; and you have access to results found by prior Agent (initial Agent findings: section below) 
                #     Use daily data unless specified otherwise (not 'all_' - since that one which contains intraday data).
                #     Once you have the information you need, you will generate and run some code to get data for the  plot from Zoltar Database tables on the stocks found by Agent #1 as a python seaborn chart, preferrably over time, 
                #     Then generate the plot:
                #         all plot components need to fit in one frame/image - an informative chart with 3 equal horizontal sections:
                #             - left -  Industry: Pie Chart of Industries of selected stocks
                #             - Middle - Expected Returns: line chart for each of the selected stocks with two points for each - first point starting at (0,0) and second point X is number of days to hold (Score_HoldPeriod in high_risk and all_high_risk tables) vs High Zoltar Rank (y-axis), making starting point for x-axis max(Date) and iterating days forward from that point.
                #             - right - Low Zoltar Rank Over Time: a pretty line chart of Low Zoltar Rank of each stock over time; 
                #           Turn x-axis labels -45 degrees.
             
                #     You should analyze data used for plotting and and create a section "References to visualization", the discussion of the new visualization.
            
                #     AND THIS IS ABSOLUTELY CRUCIAL: limit Date ranges to less than 3 months, use complex and nested query logic to FILTER UPFRONT and use aggregation logic in queries when possible.
                #     to get data from db in every SQL query and communication instead of transmitting actual data, or everything will crash.  Estimate size of output using Zoltar database tables detail and expected query output. (be cautious not to hit the total limit of 808576 bytes) 
                #     and don't use textblob.  If plotting fails more than 2 times, simplify significantly and send only 1 month of data to reduce transmitted payload.
                #     Generate Python code and execute to create a matplotlib/seaborn plot.
                #     """
                message = f"""Use the result of the first agent findings: {agent_result}. ** end of first agent result ** 
                     Your task is to create a seaborn plot.
                     You can interact with Zoltar SQL database for Stock trading education app using [execute_query_tool_def.to_json_dict()] tool and should become an expert on the contents of the database and the formats of all variables; and you have access to results found by prior Agent (initial Agent findings: section below) 
                    Use daily data unless specified otherwise (not 'all_' - since that one which contains intraday data).
                    Once you have the information you need, you will generate and run some code to get data for the  plot from Zoltar Database tables on the stocks found by Agent #1 as a python seaborn chart, preferrably over time, 
                    Then generate the plot:
                    all plot components need to fit in one landscape positioned frame/image - an informative chart with the following sections:
                    {viz_section}
                    Turn x-axis labels -45 degrees.
             
                    You should analyze data used for plotting and and create a section "References to visualization", the discussion of the new visualization.
            
                    AND THIS IS ABSOLUTELY CRUCIAL: limit Date ranges to less than 3 months, use complex and nested query logic to FILTER UPFRONT and use aggregation logic in queries when possible.
                    to get data from db in every SQL query and communication instead of transmitting actual data, or everything will crash.  Estimate size of output using Zoltar database tables detail and expected query output. (be cautious not to hit the total limit of 808576 bytes) 
                    and don't use textblob. use integers instead of string for indicies. in the past, this has been the issue and helped fix: the structure of the output now. It's a dictionary with a "result" key, whose value is a string containing a JSON-like structure. Inside that string, there's a "results" key containing a list of lists , where each inner list represents a row of data.
                    high_risk_data['result']  and low_risk_data['result'] are strings, not dictionaries. use the json.loads() function to parse the strings.
                    If plotting fails more than 2 times, simplify significantly and send only 1 month of data to reduce transmitted payload.
                    Generate Python code and execute to create a matplotlib/seaborn plot.
                    """
     
                print(f"> {message}\n")
                await session.send(input=to_json_serializable(message), end_of_turn=True)
                all_responses2b = await handle_response_refresh(session, tool_impl=execute_query)
                agent_result2b = "\n".join(msg.text for msg in all_responses2b if msg.text)  
                agent1_toast.toast("AGENT 1...ZOLTAR DATABASE", icon="✅")
                agent2_toast.toast("AGENT 2...NEWS ARTICLES", icon="✅")
                agent3_toast.toast("AGENT 3...OVERVIEW PLOTS", icon="✅")
                #formatted_state = format_global_state(global_state)

                # while (
                #     not global_state["images"] or
                #     all(
                #         (img is None) or
                #         (isinstance(img, str) and not img.strip()) or
                #         (isinstance(img, (bytes, bytearray)) and len(img) == 0)
                #         for img in global_state["images"]
                #     )
                # ):
                max_tries = 3
                tries = 0
                agent4_toasts = []
                try:

                    # while (
                    #     (tries < max_tries) and (
                    #         not global_state["images"] or
                    #         all(
                    #             (img is None)
                    #             or (isinstance(img, str) and not img.strip())
                    #             or (isinstance(img, (bytes, bytearray)) and (len(img) == 0 or is_blank_png(img)))
                    #             for img in global_state["images"]
                    #         )                    
                    #     )
                    # ):
    

                        
                    while (
                        (tries < max_tries) and (
                            (not st.session_state.image) or
                            is_blank_png(st.session_state.image)                   
                        ) and (Pie_chart or Return_hold  or low_ranks_trend or recommendations_table)
                    ):
                        # Your loop code here
                        tries += 1                 
                        toast_msg = f"AGENT 4...FALLBACK PLOTS (TRY #{tries})"
                        agent1_toast.toast("AGENT 1...ZOLTAR DATABASE", icon="✅")
                        agent2_toast.toast("AGENT 2...NEWS ARTICLES", icon="✅")
                        agent3_toast.toast("AGENT 3...OVERVIEW PLOTS", icon="✅")
                        agent4_toast = st.toast(toast_msg, icon="⏳")
                        agent4_toasts.append(agent4_toast)                        
                        # st.toast(f"AGENT 4...FALLBACK PLOTS (TRY #{tries})", icon="⏳")  # Shows a floating toast message
                        # sleep(30)

                        message = f"""Use the result of the first agent findings: {agent_result}. ** end of first agent result ** 
                             Your task is to create a plot. This is attempt number {tries}
                             You can interact with Zoltar SQL database for Stock trading education app using [execute_query_tool_def.to_json_dict()] tool and should become an expert on the contents of the database and the formats of all variables; and you have access to results found by prior Agent (initial Agent findings: section below) 
                            Use daily data unless specified otherwise (not 'all_' - since that one which contains intraday data).
                            can interact with an SQL database for Stock trading education app. You will take the users' questions and turn them into SQL
                            queries using the tools available. Once you have the information you need, you will generate and run some code to plot data from Zoltar Database tables on the stocks found by Agent #1 as a python seaborn chart, preferrably over time, 
                            Then generate the plot with only two sections from the requested vizualizations below, which need to fit in one landscape positioned frame/image - an informative chart with the following sections:
                            {viz_section}
                            Turn x-axis labels -45 degrees.                     
                            You should analyze data used for plotting and and create a section "References to visualization", the discussion of the new visualization.
                    
                            AND THIS IS ABSOLUTELY CRUCIAL: The prior attempt to generate the plot failed due to exceeding payload limit and being careless, even after taking this into account.. limit Date ranges to less than 3 months, use complex and nested query logic to FILTER UPFRONT and use aggregating functions in queries when possible
                            to get data from db in every SQL query and communication instead of transmitting actual data, or everything will crash.  Estimate size of output using Zoltar database tables detail and expected query output. (be cautious not to hit the total limit of 808576 bytes) 
                            and don't use textblob.  in the past, this has been the issue and helped fix: the structure of the output now. It's a dictionary with a "result" key, whose value is a string containing a JSON-like structure. Inside that string, there's a "results" key containing a list of lists , where each inner list represents a row of data.
                            high_risk_data['result']  and low_risk_data['result'] are strings, not dictionaries. use the json.loads() function to parse the strings.
                            If plotting fails more than 2 times, simplify significantly and send only 1 month of data. 
                            Generate Python code and execute to create matplotlib/seaborn plot.
                            """
            
                        print(f"> {message}\n")
                        # Before sending:
                        # debug_payload(message)
                        await session.send(input=to_json_serializable(message), end_of_turn=True)
                        all_responses2c = await handle_response_refresh(session, tool_impl=execute_query)
                        agent_result2c = "\n".join(msg.text for msg in all_responses2c if msg.text)  
                        agent1_toast.toast("AGENT 1...ZOLTAR DATABASE", icon="✅")
                        agent2_toast.toast("AGENT 2...NEWS ARTICLES", icon="✅")
                        agent3_toast.toast("AGENT 3...OVERVIEW PLOTS", icon="✅")
                        agent4_toast.toast(toast_msg, icon="✅")
                        #formatted_state = format_global_state(global_state)    
                except RuntimeError as e:
                    st.error(f"Stage 2c failed: {e}")
                    agent_result2c = "Stage 2c failed. No plot generated due to exceeding payload limit."
                    st.toast("AGENT 4 failed: Could not generate plot.", icon="❌")                    
                    # Optionally: continue to next stage    
                # st.toast("AGENT 5...COMPILE REPORT", icon="⏳")  # Shows a floating toast message
                agent1_toast.toast("AGENT 1...ZOLTAR DATABASE", icon="✅")
                agent2_toast.toast("AGENT 2...NEWS ARTICLES", icon="✅")
                agent3_toast.toast("AGENT 3...OVERVIEW PLOTS", icon="✅")
                agent5_toast = st.toast("AGENT 5...COMPILE REPORT", icon="⏳")
                # sleep(30)
    
                #message = f"Generate and run some code to pull necessary data from Zoltar Ranks Database for stocks found by prior agent. Plot the Price and Zoltar Ranks over time as a python seaborn chart. Return base64-encoded images.  Here is the result of the first agent findings: {agent_result2}. ***IMPORTANT*** there is a limit of 4000 characters on output so use efficient sub-queries to filter and limit timeframe to 30 days."
                message = f"""Combine the results of prior agants into a comprehensive report, and make sure to use all information synthesized by prior agents to answer this original query: {user_query}. ** End of User Query ** Here is the result of the first agent findings: {agent_result}. Here is the result of the second agent findings: {agent_result2}. ***End of AGENT 2 results**** And this is commentary of the supporting plots: {agent_result2b} *** End of Agent Results *** The final report needs to have an executive structure, containing 1. Summary section with a sentence and table of Fundamentals/About Information and overall recommendation column (Buy, Mixed, Sell), 2. News and Ratings section with Summary table for News and for Analyst Ratings with columns: Analyst Consensus,	Blogger Sentiment,	Crowd Wisdom,	News Sentiment; 3. Quant Section with Zoltar Ranks and SHAP discussion; 4. Conclusion based on contents of prior section. Return just the Final Executive Report and nothing else."""
                print(f"> {message}\n")
                # debug_payload(message)
                await session.send(input=to_json_serializable(message), end_of_turn=True)
                all_responses3 = await handle_response_refresh(session, tool_impl=execute_query)
                agent_result3 = "\n".join(msg.text for msg in all_responses3 if msg.text)  
                st.session_state.final_agent_result = agent_result3
                agent1_toast.toast("AGENT 1...ZOLTAR DATABASE", icon="✅")
                agent2_toast.toast("AGENT 2...NEWS ARTICLES", icon="✅")
                agent3_toast.toast("AGENT 3...OVERVIEW PLOTS", icon="✅")
                agent5_toast.toast("AGENT 5...COMPILE REPORT", icon="✅")
                st.toast("Final report completed!", icon="✅")
                st.balloons()
                #formatted_state = format_global_state(global_state)
    
        # with col2:
        # Run the async code
        asyncio.run(main(user_query))

    if st.session_state.final_agent_result:
        with st.popover("✅ Ready to share the results?"):   
     # still continuing with col2
           
            ## 5.24.25: new section to email results
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            import smtplib
            from email.mime.multipart import MIMEMultipart
            from email.mime.text import MIMEText
            from email.mime.base import MIMEBase
            from email import encoders
            
            def add_bold_runs(paragraph, text):
                import re
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = paragraph.add_run(part[2:-2])
                        run.bold = True
                    else:
                        paragraph.add_run(part)
            
            def save_to_docx(content, filename="Bot_Output.docx", image_path="stock_price_plot.png"):
                doc = Document()
                doc.add_heading('Your Zoltar Financial Research', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if os.path.exists(image_path):
                    doc.add_picture(image_path, width=Inches(6))
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    if line.startswith('## '):
                        header_text = line[3:].strip()
                        p = doc.add_heading(header_text, level=2)
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    elif line.startswith('* '):
                        bullet_text = line[2:].strip()
                        p = doc.add_paragraph(style='List Bullet')
                        add_bold_runs(p, bullet_text)
                    else:
                        p = doc.add_paragraph()
                        add_bold_runs(p, line)
                doc.save(filename)
                return filename
            
            def send_email(sender, password, recipient, doc_path):
                msg = MIMEMultipart()
                msg['From'] = f"Zoltar Financial <{sender}>"
                msg['To'] = recipient
                msg['Subject'] = "Your Zoltar Research Report"
                body = "Thank you for using Zoltar Financial Research Assistant. Please find attached the generated report."
                msg.attach(MIMEText(body, 'plain'))
                with open(doc_path, "rb") as attachment:
                    part = MIMEBase("application", "octet-stream")
                    part.set_payload(attachment.read())
                encoders.encode_base64(part)
                part.add_header("Content-Disposition", f"attachment; filename={os.path.basename(doc_path)}")
                msg.attach(part)
                try:
                    server = smtplib.SMTP_SSL('smtp.gmail.com', 465)
                    server.login(sender, password)
                    server.send_message(msg)
                    server.close()
                    return True
                except Exception as e:
                    st.error(f"Failed to send email: {e}")
                    return False
            
            # --- Streamlit App ---
            
            st.header("Share your research results", help="Save this report as a .docx file and email it to yourself.")
        
        
            
            # Example: get your research content and image path
            # Replace this with your actual result variable
            content = st.session_state.get("final_agent_result", "Your research content goes here.")
            image_path = "stock_price_plot.png"
            
            try:
                sender = st.secrets["GMAIL"]["GMAIL_ACCT"]
                password = st.secrets["GMAIL"]["GMAIL_PASS"]
            except:
                # If Streamlit secrets are not available, use environment variables
                sender = os.getenv('GMAIL_ACCT')
                password = os.getenv('GMAIL_PASS') 
                # st.error("Gmail credentials not found in secrets. Please check your configuration.")
            
            with st.form("email_form"):
                recipient = st.text_input("Recipient email address")
                # sender = st.text_input("Sender Gmail address")
                # password = st.text_input("Sender Gmail password (use App Password)", type="password")
                submitted = st.form_submit_button("Send Report")
                if submitted:
                    if not recipient or not sender or not password:
                        st.error("Please fill in all fields.")
                    else:
                        doc_path = save_to_docx(content, filename="zoltar_financial_report.docx", image_path=image_path)
                        st.success(f"Document saved as {doc_path}")
                        if send_email(sender, password, recipient, doc_path):
                            st.success(f"Report sent successfully to {recipient}!")