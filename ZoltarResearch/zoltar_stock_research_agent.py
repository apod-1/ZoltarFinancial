# -*- coding: utf-8 -*-
"""
Created on Thu Apr  3 14:46:53 2025

General description:
    This is a capstone project for genAI 5-day Intensive Google-led course
    
    Main goal: Show knowledge of concepts to establish a grounding database, vectorize it, and utilize available functions to be envoked by AI Agent at the appropriate time (as-needed)
    Secondary goal: Show knowledge of advanced concepts, including: 
            multi-agent system, 
            langgraph representation
            transparent chain-of-thought steps, 
            API calls/search in real-time data, 


Application: Helpful bot that can analyze stocks and present a comprehensive report of the analysis to the user, utilizing proprietary Zoltar Ranks as grounding database.

launch (at home)
    activate myenv
    streamlit_env\Scripts\activate
    cd C:\ Users\apod7\StockPicker\app\ZoltarFinancial\ZoltarResearch    
    streamlit run zoltar_stock_research_agent.py
@author: apod
"""
import os
import openai
import pandas as pd
import streamlit as st
import markdown2    
import json
import sqlite3
import textwrap
import asyncio
# import IPython
import io
import altair as alt
import seaborn as sns
import base64
import re
import matplotlib.pyplot as plt  # Import Matplotlib globally
import requests
import random
# from IPython.display import display, Image, Markdown
from pprint import pprint
from pprint import pformat
from dotenv import load_dotenv
from google import genai 
from google.genai import types
from google.api_core import retry
from datetime import datetime
from io import BytesIO
from PIL import Image  # Now safe from namespace collision
from websockets.exceptions import ConnectionClosedError
# load_dotenv()
# # GOOGLE_API = os.getenv('GOOGLE_API_KEY')
# GMAIL_ACCT = os.getenv('GMAIL_ACCT')
# GMAIL_PASS = os.getenv('GMAIL_PASS')

try:
    favicon = "https://github.com/apod-1/ZoltarFinancial/raw/main/docs/ZoltarSurf_48x48.png"
except (KeyError, FileNotFoundError):
    favicon = st.secrets["browser"]["favicon"]

st.set_page_config(page_title="Zoltar Stock Research Agent", page_icon=favicon, layout="wide", initial_sidebar_state="collapsed")


# Load environment variables
try:
    GOOGLE_API = None #os.getenv('GOOGLE_API_KEY')
    if GOOGLE_API:
        GOOGLE_API_KEY = GOOGLE_API        
    else: 
        GOOGLE_API_KEY = st.secrets["google_api"]["api_key"]
except:
    print("Error")
# from langchain_openai import ChatOpenAI
# from langchain.globals import set_verbose, set_debug, set_llm_cache
# from langchain.cache import InMemoryCache  # For in-memory caching

# Set verbosity and debugging options
# set_verbose(True)  # Enables detailed logs globally
# set_debug(False)   # Disables deep debugging for now

# Set up in-memory LLM caching to avoid redundant API calls
# set_llm_cache(InMemoryCache())


# st.markdown("""
# <style>
#     [data-testid="collapsedControl"] {
#         display: none !important;
#     }
# </style>
# """, unsafe_allow_html=True)

hide_streamlit_style = """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    </style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)    

# # 6.5.25 - collapse bottom
# st.markdown("""
# <style>
#     [data-testid="collapsedControl"] {
#         display: none !important;
#     }
# </style>
# """, unsafe_allow_html=True)
    
# These are the Python functions defined above.
# db_tools = [list_tables, describe_table, execute_query]

# instruction = """You are a helpful chatbot that can interact with an SQL database
# for a computer store. You will take the users questions and turn them into SQL
# queries using the tools available. Once you have the information you need, you will
# answer the user's question using the data returned.

# Use list_tables to see what tables are present, describe_table to understand the
# schema, and execute_query to issue an SQL SELECT query."""
col1, col2, col3 = st.columns([1, 5, 1])    


with col2:
# Streamlit UI for user input
    st.title("US Equities Zoltar Research Agent 🤖", help="I am here to help you make better decisions! Don't be shy - ask away...")


# 5.26.25 - initialize session state variables 
if 'final_agent_result' not in st.session_state:
    st.session_state.final_agent_result = ""
if 'image' not in st.session_state:
    st.session_state.image = None
if "temp_selected" not in st.session_state:
    st.session_state["temp_selected"] = "0.1 - Middle"  # or your desired default

if "top_p_selected" not in st.session_state:
    st.session_state["top_p_selected"] = "0.9 - Middle"   # or your desired default
if "agent_repo" not in st.session_state:
    st.session_state.agent_repo = {
        "agents": {},
        "execution_order": []
    }
# A little pre-work to set up what we need:

    # 1. Set up databases we'll need (5 total)

    # 2. context and metadata
    
    # 3. define functions and other tools available, including live API
    
    # 4. create sliders/selectors for tuning model parameters (to be expanded later)
    
    # 5. define and create interactive structure and guidelines / tool use instructions for agents



# background creatives

def set_bg_video(video_file):
    st.markdown(
        f"""
        <style>
        .stApp {{
            background: transparent !important;
        }}
        .block-container {{
            background: transparent !important;
        }}
        .main {{
            background: transparent !important;
        }}
        video.bgvid {{
            position: fixed;
            top: 50%;
            left: 50%;
            min-width: 100vw;
            min-height: 100vh;
            width: auto;
            height: auto;
            z-index: -1;
            object-fit: cover;
            opacity: 0.7;
            pointer-events: none;
            /* Zoom in by scaling the video */
            transform: translate(-50%, -50%) scale(1.27);  /* change 1.2 to any zoom factor you want */
        }}
        </style>
        <video autoplay loop muted class="bgvid">
            <source src="data:video/mp4;base64,{video_file}" type="video/mp4">
        </video>
        """,
        unsafe_allow_html=True
    )

    
# Load video and encode as base64
url = "https://github.com/apod-1/ZoltarFinancial/raw/main/docs/wave_vid.mp4"
response = requests.get(url)
video_bytes = response.content
encoded = base64.b64encode(video_bytes).decode()

set_bg_video(encoded)



# Helper function to verify files and print their columns

def get_latest_file(data_dir=None, prefix=None):
    """
    Finds the latest file in a directory based on a given prefix.

    Args:
        data_dir (str): Directory to search for files. Defaults to Zoltar Financial's daily ranks directory.
        prefix (str): Prefix to filter files (e.g., "high_risk_PROD", "low_risk_PROD").

    Returns:
        str: Path to the latest file matching the prefix, or None if no valid files are found.
    """
    try:
        # Default directory setup
        if data_dir is None:
            # if os.path.exists("https://github.com/apod-1/ZoltarFinancial/main/daily_ranks"):
            #     data_dir = 'https://github.com/apod-1/ZoltarFinancial/main/daily_ranks/'
            # else:
            data_dir = '/mount/src/zoltarfinancial/daily_ranks'

        # Find files matching the prefix
        files = [f for f in os.listdir(data_dir) if f.startswith(prefix) and f.endswith(".pkl")]
        if not files:
            print(f"No valid files found for prefix '{prefix}' in directory '{data_dir}'.")
            return None

        # Find the latest file based on modification time
        latest_file = max(files, key=lambda x: os.path.getmtime(os.path.join(data_dir, x)))

        # Return the full path to the latest file
        return os.path.join(data_dir, latest_file)

    except FileNotFoundError:
        st.error("Unable to load the latest files. Please try again later.")
        return None

def get_latest_file_from_github(base_url, prefix):
    """
    Finds and downloads the latest .pkl file from a GitHub directory listing page,
    based on file name prefix, using raw file URLs.
    """
    try:
        response = requests.get(base_url)
        response.raise_for_status()
        html_content = response.text
        pattern_href = re.compile(r'href="([^"]*\.pkl)"')
        matches = pattern_href.findall(html_content)

        # --- FIXED MATCHING LOGIC ---
        pattern = re.compile(rf"^{re.escape(prefix)}(_|\..*$)")
        matching_files = [f.split('/')[-1] for f in matches if pattern.match(f.split('/')[-1])]
        # ----------------------------

        if not matching_files:
            print(f"No .pkl files with prefix '{prefix}' found at {base_url}")
            return None

        latest_file_name = max(matching_files)
        print(f"Latest .pkl file found: {latest_file_name}")

        if prefix=="fundamentals_df":
            raw_file_url = f"https://github.com/apod-1/ZoltarFinancial/main/data/{latest_file_name}"
        elif prefix=="ratings_detail_df":
            raw_file_url = f"https://github.com/apod-1/ZoltarFinancial/main/data/{latest_file_name}"
        else:
            raw_file_url = f"https://github.com/apod-1/ZoltarFinancial/main/daily_ranks/{latest_file_name}"

        response = requests.get(raw_file_url)
        response.raise_for_status()
        df = pd.read_pickle(BytesIO(response.content))
        return df

    except requests.exceptions.RequestException as e:
        print(f"Request error: {e}")
        return None
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

import random
import string
# from datetime import datetime, timedelta, date, time
from time import sleep


def random_db_filename(base_name="zoltar_financial.db"):
    name, ext = os.path.splitext(base_name)
    suffix = ''.join(random.choices(string.ascii_lowercase + string.digits, k=6))
    return f"{name}_{suffix}{ext}"

def get_sqlite_connection_with_random_on_lock(db_file, max_retries=3, retry_delay=0.5):
    for attempt in range(max_retries):
        try:
            conn = sqlite3.connect(db_file, timeout=10)
            # Try a simple operation to check if locked
            conn.execute("PRAGMA quick_check;")
            return conn, db_file
        except sqlite3.OperationalError as e:
            if "database is locked" in str(e):
                print(f"Database is locked, creating new db file with random suffix (attempt {attempt+1})...")
                db_file = random_db_filename(db_file)
                sleep(retry_delay)
            else:
                raise
    raise RuntimeError("Could not acquire database connection after multiple retries (database is locked).")


# 1. Set up databases we'll need (5 total)
# Define database connection
db_file = "zoltar_financial.db"
db_conn, db_file_used = get_sqlite_connection_with_random_on_lock(db_file)
# db_conn = sqlite3.connect(db_file)

print(f"Using database file: {db_file_used}")




# Create tables in SQLite database
def create_tables():
    with db_conn:
        db_conn.executescript("""
            -- Drop tables if they exist
            DROP TABLE IF EXISTS high_risk;
            DROP TABLE IF EXISTS low_risk;
            DROP TABLE IF EXISTS all_high_risk;
            DROP TABLE IF EXISTS all_low_risk;
            DROP TABLE IF EXISTS fundamentals;
            DROP TABLE IF EXISTS ratings_detail;                              
            DROP TABLE IF EXISTS shap_summary_Large;                              
            DROP TABLE IF EXISTS shap_summary_Mid;                              
            DROP TABLE IF EXISTS shap_summary_Small;                              
        """)
        db_conn.executescript("""
            -- Table for high-risk stocks
            CREATE TABLE IF NOT EXISTS high_risk (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Date DATETIME,
                Symbol TEXT,
                Score REAL,
                Score_Sharpe REAL,
                Score_HoldPeriod REAL,
                Close_Price REAL,
                Cap_Size TEXT,
                Sector TEXT,
                Industry TEXT,
                source TEXT
            );

            -- Table for low-risk stocks
            CREATE TABLE IF NOT EXISTS low_risk (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Date DATETIME,
                Symbol TEXT,
                Score REAL,
                Score_Sharpe REAL,
                Score_HoldPeriod REAL,
                Close_Price REAL,
                Cap_Size TEXT,
                Sector TEXT,
                Industry TEXT,
                source TEXT
            );
            -- Table for high-risk stocks
            CREATE TABLE IF NOT EXISTS all_high_risk (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Date DATETIME,
                Symbol TEXT,
                Score REAL,
                Score_Sharpe REAL,
                Score_HoldPeriod REAL,
                Close_Price REAL,
                Cap_Size TEXT,
                Sector TEXT,
                Industry TEXT,
                source TEXT
            );

            -- Table for low-risk stocks
            CREATE TABLE IF NOT EXISTS all_low_risk (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Date DATETIME,
                Symbol TEXT,
                Score REAL,
                Score_Sharpe REAL,
                Score_HoldPeriod REAL,
                Close_Price REAL,
                Cap_Size TEXT,
                Sector TEXT,
                Industry TEXT,
                source TEXT
            );

            -- Table for fundamentals
            CREATE TABLE IF NOT EXISTS fundamentals (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Symbol TEXT,
                Fundamentals_OverallRating REAL,
                total_ratings INTEGER,
                Fundamentals_Sector TEXT,
                Fundamentals_Industry TEXT,
                Fundamentals_Dividends REAL,
                Fundamentals_PE REAL,
                Fundamentals_PB REAL,
                Fundamentals_MarketCap REAL,
                Fundamentals_avgVolume2Weeks REAL,
                Fundamentals_avgVolume30Days REAL,
                Fundamentals_52WeekHigh REAL,
                Fundamentals_52WeekLow REAL,
                Fundamentals_52WeekHighDate DATE,
                Fundamentals_52WeekLowDate DATE,
                Fundamentals_Float REAL,
                Fundamentals_SharesOutstanding INTEGER,
                Fundamentals_CEO TEXT,
                Fundamentals_NumEmployees INTEGER,
                Fundamentals_YearFounded INTEGER,
                Fundamentals_ExDividendDate DATE,
                Fundamentals_PayableDate DATE,
                Fundamentals_Description TEXT
            );

            -- Table for ratings detail
            CREATE TABLE IF NOT EXISTS ratings_detail (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Symbol TEXT,
                RatingType TEXT,
                RatingText TEXT,
                RatingPublishedAt DATETIME
            );
        """)

# Helper function to find the latest file by timestamp in a directory
# def drop_tables():
#     """
#     Explicitly drop specific tables from the database.
#     """
#     tables_to_drop = ["high_risk", "low_risk", "fundamentals", "ratings_detail", "shap_summary"]
    
#     with db_conn:
#         cursor = db_conn.cursor()
#         for table in tables_to_drop:
#             print(f"Dropping table: {table}")
#             cursor.execute(f"DROP TABLE IF EXISTS {table};")
#         db_conn.commit()
#         print("Specified tables dropped successfully.")

# Function to list all tables in the database
def list_tables() -> list[str]:
    """Retrieve the names of all tables in the database."""
    st.write(' - DB CALL: list_tables()')

    cursor = db_conn.cursor()
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
    tables = cursor.fetchall()
    return [t[0] for t in tables]

# Function to describe a table's schema
def describe_table(table_name: str) -> list[tuple[str, str]]:
    """Look up the table schema.

    Returns:
      List of columns, where each entry is a tuple of (column, type).
    """
    st.write(f' - DB CALL: describe_table({table_name})')

    cursor = db_conn.cursor()
    cursor.execute(f"PRAGMA table_info({table_name});")
    schema = cursor.fetchall()
    return [(col[1], col[2]) for col in schema]

# Function to execute an SQL query
# def execute_query(sql: str) -> list[list[str]]:
#     """Execute an SQL statement, returning the results."""
#     st.write(f' - DB CALL: execute_query({sql})')

#     cursor = db_conn.cursor()
#     cursor.execute(sql)
#     return cursor.fetchall()
def execute_query(sql: str) -> list[list[str]]:
    """Execute an SQL statement, returning the results."""
    # Don't st.write here!
    cursor = db_conn.cursor()
    cursor.execute(sql)
    results = cursor.fetchall()
    # Return both the call string and results
    return {"call": f"execute_query({sql})", "results": results}

# Function to dynamically create shap_summary table based on DataFrame columns
def create_shap_table(df):
    """
    Dynamically create the shap_summary table based on DataFrame columns.
    """
    with db_conn:
        # Drop the shap_summary table if it exists
        db_conn.execute("DROP TABLE IF EXISTS shap_summary;")        
        # Dynamically generate column definitions
        columns = ", ".join([f'"{col}" REAL' for col in df.columns])
            # CREATE TABLE IF NOT EXISTS shap_summary (
        create_table_query = f"""
            CREATE TABLE {df} (
                id INTEGER PRIMARY KEY AUTOINCREMENT, 
                {columns}
            );
        """
        db_conn.execute(create_table_query)
        print("Created shap_summary table with dynamic schema.")



def infer_sqlite_type(col, sample_val=None):
    """Infer SQLite type based on column name or sample value."""
    col_lower = col.lower()
    if "date" in col_lower or "time" in col_lower:
        return "DATETIME"
    if "symbol" in col_lower or "sector" in col_lower or "industry" in col_lower or "source" in col_lower:
        return "TEXT"
    if sample_val is not None:
        if isinstance(sample_val, (int, float)):
            return "REAL"
        if isinstance(sample_val, str):
            return "TEXT"
    return "REAL"

def recreate_table_from_df(conn, table, df):
    """Drop and recreate a table with columns matching the DataFrame."""
    cols = []
    for col in df.columns:
        sample_val = df[col].dropna().iloc[0] if not df[col].dropna().empty else None
        col_type = infer_sqlite_type(col, sample_val)
        cols.append(f'"{col}" {col_type}')
    schema = ", ".join(cols)
    sql = f'CREATE TABLE IF NOT EXISTS "{table}" ({schema});'
    with conn:
        conn.execute(f'DROP TABLE IF EXISTS "{table}";')
        conn.execute(sql)

def load_data_into_db():
    paths = [
        {"path": "/mount/src/zoltarfinancial/daily_ranks/", "prefix": "all_high_risk_PROD", "table": "all_high_risk"},
        {"path": "/mount/src/zoltarfinancial/daily_ranks/", "prefix": "all_low_risk_PROD", "table": "all_low_risk"},
        {"path": "/mount/src/zoltarfinancial/daily_ranks/", "prefix": "high_risk_PROD", "table": "high_risk"},
        {"path": "/mount/src/zoltarfinancial/daily_ranks/", "prefix": "low_risk_PROD", "table": "low_risk"},
        {"path": "/mount/src/zoltarfinancial/data/", "prefix": "fundamentals_df", "table": "fundamentals"},
        {"path": "/mount/src/zoltarfinancial/data/", "prefix": "ratings_detail_df", "table": "ratings_detail"},
        {"path": "/mount/src/zoltarfinancial/daily_ranks/", "prefix": "combined_SHAP_summary_Large", "table": "shap_summary_Large"},
        {"path": "/mount/src/zoltarfinancial/daily_ranks/", "prefix": "combined_SHAP_summary_Mid", "table": "shap_summary_Mid"},
        {"path": "/mount/src/zoltarfinancial/daily_ranks/", "prefix": "combined_SHAP_summary_Small", "table": "shap_summary_Small"}
    ]
    # paths = [
    #     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "all_high_risk_PROD", "table": "all_high_risk"},
    #     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "all_low_risk_PROD", "table": "all_low_risk"},
    #     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "high_risk_PROD", "table": "high_risk"},
    #     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "low_risk_PROD", "table": "low_risk"},
    #     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\data", "prefix": "fundamentals_df", "table": "fundamentals"},
    #     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\data", "prefix": "ratings_detail_df", "table": "ratings_detail"},
    #     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "combined_SHAP_summary_Large", "table": "shap_summary_Large"},
    #     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "combined_SHAP_summary_Mid", "table": "shap_summary_Mid"},
    #     {"path": r"C:\Users\apod7\StockPicker\app\ZoltarFinancial\daily_ranks", "prefix": "combined_SHAP_summary_Small", "table": "shap_summary_Small"}
    # ]
    shap_tables = {"shap_summary_Large", "shap_summary_Mid", "shap_summary_Small"}

    for entry in paths:
        directory = entry["path"]
        prefix = entry["prefix"]
        table = entry["table"]
        print(f"Fetching latest file for table '{table}' with prefix '{prefix}'...")

        # Construct BASE_URL dynamically based on path
        
        latest_file = get_latest_file(directory, prefix)
        if not latest_file:
            print(f"No file found for {prefix} in {directory}")
            continue

        print(f"Loading {latest_file} into {table}...")
        df = pd.read_pickle(latest_file)

        # Special handling for ratings_detail table
        if table == 'ratings_detail':
            df['RatingText'] = df['RatingText'].apply(lambda x: x.decode('utf-8') if isinstance(x, bytes) else x)
            df['RatingPublishedAt'] = pd.to_datetime(df['RatingPublishedAt'], errors='coerce')

        # SHAP tables: drop "Feature Category" and reset index if needed
        if table in shap_tables:
            if "Feature Category" in df.columns:
                df = df.drop(columns=["Feature Category"])
            if df.index.name is not None or not df.index.equals(pd.RangeIndex(len(df))):
                df = df.reset_index()
            df = df.rename(columns={'index': 'Symbol'})

            # Debug: print DataFrame info before inserting
            #st.dataframe(df.style.format(precision=9))
            # st.write("\n--- DEBUG: DataFrame to be written to", table, "---")
            # st.write("Columns:", df.columns.tolist())
            # st.write("Dtypes:\n", df.dtypes)
            # st.write("Sample rows:\n", df.head(10))
            # st.write("Describe:\n", df.describe())
            # st.write("--- END DEBUG ---\n")

        # Insert data into database (let pandas create table and types)
        with db_conn:
            try:
                df.to_sql(table, db_conn, if_exists="replace", index=False)
                print(f"Successfully inserted data into {table}.")
                cursor = db_conn.cursor()
                cursor.execute(f"SELECT COUNT(*) FROM {table};")
                count = cursor.fetchone()[0]
                print(f"Table {table} now contains {count} rows.")
            except Exception as e:
                print(f"Error inserting into {table}: {e}")

        # Optional: read back and check first few rows
        # if table in shap_tables:
        #     df_check = pd.read_sql_query(f"SELECT * FROM {table} LIMIT 10", db_conn)
        #     print(f"\n--- DEBUG: Data read back from {table} ---")
        #     print(df_check)
        #     print("--- END DEBUG ---\n")
# Usage
# create_tables()
# load_data_into_db()

# print("Database setup complete.")

try:
    create_tables()
    load_data_into_db()
    print(f"Tables created and data loaded successfully in {db_file_used}.")
except Exception as e:
    print(f"Error during table creation or data loading: {e}")

    # Define database connection
    db_file = random_db_filename(db_file)
    # db_file = "zoltar_financial2.db"
    db_conn, db_file_used = get_sqlite_connection_with_random_on_lock(db_file)
    # db_conn = sqlite3.connect(db_file)
    
    print(f"Using database file: {db_file_used}")

with st.sidebar:
    st.sidebar.markdown(
        """
        <style>
        .zoltar-btn {
            background: linear-gradient(135deg, #301934 0%, #9370DB 100%);
            border: none;
            color: #fff;
            padding: 14px 28px;
            text-align: center;
            text-decoration: none;
            display: inline-block;
            font-size: 18px;
            font-weight: 600;
            margin: 8px 2px;
            border-radius: 12px;
            cursor: pointer;
            box-shadow:
                0 4px 14px 0 rgba(80, 40, 120, 0.45),
                0 1.5px 8px 2px rgba(255,255,255,0.06) inset;
            transition: all 0.18s cubic-bezier(.4,0,.2,1);
            position: relative;
            outline: none;
        }
        .zoltar-btn:hover, .zoltar-btn:focus {
            background: linear-gradient(135deg, #9370DB 0%, #301934 100%);
            box-shadow:
                0 8px 24px 0 rgba(80, 40, 120, 0.55),
                0 2px 12px 3px rgba(255,255,255,0.10) inset;
            transform: translateY(-2px) scale(1.04);
        }
        .zoltar-btn:active {
            box-shadow:
                0 2px 6px 0 rgba(80, 40, 120, 0.30),
                0 1px 4px 1px rgba(255,255,255,0.08) inset;
            transform: translateY(1px) scale(0.98);
        }
        </style>
        <a href="https://zoltar.streamlit.app" target="_blank" title="Open the main Zoltar Financial Research Platform in a new tab.">
            <button class="zoltar-btn" title="Open the main Zoltar Financial Research Platform in a new tab.">
                Open Zoltar Research Platform
            </button>
        </a>
        """,
        unsafe_allow_html=True
    )    
    show_top_symbols = st.sidebar.toggle("Show Top Symbols Section", value=True)
    with st.expander("Bubble Display Settings",expanded=False):
        if show_top_symbols:
            c1, c2 = st.columns(2)
            with c1:
                top_n1 = st.number_input("Symbols for Low Rank", min_value=1, max_value=20, value=5, step=1)
            with c2:
                top_n2 = st.number_input("Symbols for High Rank", min_value=1, max_value=20, value=5, step=1)

# 5.28.25 - new section for bubbles with top stocks in bubbles on the sides
def generate_top_10_stream(db_path='zoltar_financial.db'):
    conn = sqlite3.connect(db_path)
    
    # try:
    #     # Get latest date
    #     latest_date = conn.execute(
    #         "SELECT MAX(Date) FROM low_risk"
    #     ).fetchone()[0]
        
    #     # Get top 10 symbols from low_risk
    #     top_symbols = conn.execute(f"""
    #         SELECT Symbol, Score as Low_Risk_Score 
    #         FROM low_risk 
    #         WHERE Date = '{latest_date}'
    #         ORDER BY Low_Risk_Score DESC 
    #         LIMIT 10
    #     """).fetchall()
    try:
        # Get latest date
        latest_date = conn.execute(
            "SELECT MAX(Date) FROM low_risk"
        ).fetchone()[0]
        
        # Get top 10 symbols from low_risk
        top_low = conn.execute(f"""
            SELECT Symbol, Score as Low_Risk_Score 
            FROM low_risk 
            WHERE Date = '{latest_date}'
            GROUP BY 1,2
            ORDER BY Low_Risk_Score DESC 
            LIMIT {top_n1}
        """).fetchall()
        
        # Get top 10 symbols from low_risk
        top_high = conn.execute(f"""
             SELECT Symbol, Score as High_Risk_Score 
             FROM high_risk 
             WHERE Date = '{latest_date}'
             GROUP BY 1,2
             ORDER BY High_Risk_Score DESC 
             LIMIT {top_n2}
        """).fetchall()
        top_symbols = top_low + top_high
    # try:
    #     # Get latest date
    #     latest_date = conn.execute(
    #         "SELECT MAX(Date) FROM low_risk"
    #     ).fetchone()[0]

    #     # Get top N symbols from low_risk
    #     top_low = conn.execute(f"""
    #         SELECT Symbol, Score as Low_Risk_Score 
    #         FROM low_risk 
    #         WHERE Date = ?
    #         ORDER BY Low_Risk_Score DESC 
    #         LIMIT ?
    #     """, (latest_date, top_n1)).fetchall()

    #     # Get top N symbols from high_risk
    #     top_high = conn.execute(f"""
    #         SELECT Symbol, Score as High_Risk_Score 
    #         FROM high_risk 
    #         WHERE Date = ?
    #         ORDER BY High_Risk_Score DESC 
    #         LIMIT ?
    #     """, (latest_date, top_n2)).fetchall()

        # Combine, avoiding duplicates (keep order: low_risk first, then high_risk additions)
        symbols_seen = set()
        combined = []
        for symbol, score in top_low + top_high:
            if symbol not in symbols_seen:
                symbols_seen.add(symbol)
                combined.append(symbol)        
        stream_content = []
        
        for symbol, low_score in top_symbols:
            try:
                # Get high risk data
                high_data = conn.execute(f"""
                    SELECT Score as High_Risk_Score, Score_HoldPeriod as High_Risk_Score_HoldPeriod 
                    FROM high_risk 
                    WHERE Symbol = '{symbol}' AND Date = '{latest_date}'
                """).fetchone()
                
                # Get fundamentals
                fundamentals = conn.execute(f"""
                    SELECT Fundamentals_Industry, Fundamentals_Sector,
                           Fundamentals_PE, Fundamentals_PB,
                           Fundamentals_Dividends, Fundamentals_ExDividendDate,
                           Fundamentals_MarketCap, Fundamentals_Description
                    FROM fundamentals 
                    WHERE Symbol = '{symbol}'
                """).fetchone()
                
                if not high_data or not fundamentals:
                    continue
                
                # Unpack data
                high_score, hold_period = high_data
                (industry, sector, pe, pb, 
                 dividend, ex_div, mcap, desc) = fundamentals
                
                # Format values
                dividend_pct = f"{dividend:.2f}%" if dividend else "none"
                ex_div_date = pd.to_datetime(ex_div).strftime('%m-%d-%Y') if ex_div else 'N/A'
                mcap_formatted = f"${mcap/1e9:.2f}B" if mcap else 'N/A'
                truncated_desc = f"{desc[:300]}..." if desc else ""
                
                stream_content.append({
                    "symbol": symbol,
                    "low_score": f"{low_score:.2%}",
                    "high_score": f"{high_score:.2%}",
                    "hold_period": f"{hold_period:.0f}d",
                    "industry": industry,
                    "sector": sector,
                    "pe": f"{pe:.2f}",
                    "pb": f"{pb:.2f}",
                    "dividend": dividend_pct,
                    "ex_div": ex_div_date,
                    "mcap": mcap_formatted,
                    "desc": truncated_desc
                })
                
            except Exception as e:
                print(f"Error processing {symbol}: {str(e)}")
                
        return stream_content
        
    finally:
        conn.close()



def generate_top_10_stream(db_path='zoltar_financial.db'):
    conn = sqlite3.connect(db_path)
    
    # try:
    #     # Get latest date
    #     latest_date = conn.execute(
    #         "SELECT MAX(Date) FROM low_risk"
    #     ).fetchone()[0]
        
    #     # Get top 10 symbols from low_risk
    #     top_symbols = conn.execute(f"""
    #         SELECT Symbol, Score as Low_Risk_Score 
    #         FROM low_risk 
    #         WHERE Date = '{latest_date}'
    #         ORDER BY Low_Risk_Score DESC 
    #         LIMIT 10
    #     """).fetchall()
    try:
        # Get latest date
        latest_date = conn.execute(
            "SELECT MAX(Date) FROM low_risk"
        ).fetchone()[0]
        
        # Get top 10 symbols from low_risk
        top_low = conn.execute(f"""
            SELECT Symbol, Score as Low_Risk_Score 
            FROM low_risk 
            WHERE Date = '{latest_date}'
            GROUP BY 1,2
            ORDER BY Low_Risk_Score DESC 
            LIMIT {top_n1}
        """).fetchall()
        
        # Get top 10 symbols from low_risk
        top_high = conn.execute(f"""
             SELECT Symbol, Score as High_Risk_Score 
             FROM high_risk 
             WHERE Date = '{latest_date}'
             GROUP BY 1,2
             ORDER BY High_Risk_Score DESC 
             LIMIT {top_n2}
        """).fetchall()
        top_symbols = top_low + top_high

        # Combine, avoiding duplicates (keep order: low_risk first, then high_risk additions)
        symbols_seen = set()
        combined = []
        for symbol, score in top_low + top_high:
            if symbol not in symbols_seen:
                symbols_seen.add(symbol)
                combined.append(symbol)        
        stream_content = []
        
        for symbol, _ in top_symbols:
            try:
                # Always get the low risk score for this symbol
                low_data = conn.execute(f"""
                    SELECT Score FROM low_risk 
                    WHERE Symbol = '{symbol}' AND Date = '{latest_date}'
                """).fetchone()
                low_score = low_data[0] if low_data else None
        
                # Get high risk data
                high_data = conn.execute(f"""
                    SELECT Score as High_Risk_Score, Score_HoldPeriod as High_Risk_Score_HoldPeriod 
                    FROM high_risk 
                    WHERE Symbol = '{symbol}' AND Date = '{latest_date}'
                """).fetchone()
                
                # Get fundamentals
                fundamentals = conn.execute(f"""
                    SELECT Fundamentals_Industry, Fundamentals_Sector,
                           Fundamentals_PE, Fundamentals_PB,
                           Fundamentals_Dividends, Fundamentals_ExDividendDate,
                           Fundamentals_MarketCap, Fundamentals_Description
                    FROM fundamentals 
                    WHERE Symbol = '{symbol}'
                """).fetchone()
                
                if not high_data or not fundamentals:
                    continue
                
                # Unpack data
                high_score, hold_period = high_data
                (industry, sector, pe, pb, 
                 dividend, ex_div, mcap, desc) = fundamentals
                
                # Format values
                dividend_pct = f"{dividend:.2f}%" if dividend else "none"
                ex_div_date = pd.to_datetime(ex_div).strftime('%m-%d-%Y') if ex_div else 'N/A'
                mcap_formatted = f"${mcap/1e9:.2f}B" if mcap else 'N/A'
                truncated_desc = f"{desc[:300]}..." if desc else ""
                
                stream_content.append({
                    "symbol": symbol,
                    "low_score": f"{low_score:.2%}" if low_score is not None else "N/A",
                    "high_score": f"{high_score:.2%}",
                    "hold_period": f"{hold_period:.0f}d",
                    "industry": industry,
                    "sector": sector,
                    "pe": f"{pe:.2f}",
                    "pb": f"{pb:.2f}",
                    "dividend": dividend_pct,
                    "ex_div": ex_div_date,
                    "mcap": mcap_formatted,
                    "desc": truncated_desc
                })
                
            except Exception as e:
                print(f"Error processing {symbol}: {str(e)}")
                
        return stream_content
        
    finally:
        conn.close()


# def generate_top_10_stream(db_path='zoltar_financial.db', top_n1=10, top_n2=10):
#     conn = sqlite3.connect(db_path)
#     try:
#         # Get latest date
#         latest_date = conn.execute(
#             "SELECT MAX(Date) FROM low_risk"
#         ).fetchone()[0]

#         # Get top N symbols from low_risk
#         top_low = conn.execute(f"""
#             SELECT Symbol, Score as Low_Risk_Score 
#             FROM low_risk 
#             WHERE Date = ?
#             ORDER BY Low_Risk_Score DESC 
#             LIMIT ?
#         """, (latest_date, top_n1)).fetchall()

#         # Get top N symbols from high_risk
#         top_high = conn.execute(f"""
#              SELECT Symbol, Score as High_Risk_Score 
#              FROM high_risk 
#              WHERE Date = ?
#              ORDER BY High_Risk_Score DESC 
#              LIMIT ?
#         """, (latest_date, top_n2)).fetchall()

#         # Combine, avoiding duplicates
#         symbols_seen = set()
#         combined = []
#         for symbol, _ in top_low + top_high:
#             if symbol not in symbols_seen:
#                 symbols_seen.add(symbol)
#                 combined.append(symbol)

#         stream_content = []
#         for symbol in combined:
#             try:
#                 # Get low risk score
#                 low_data = conn.execute("""
#                     SELECT Score FROM low_risk 
#                     WHERE Symbol = ? AND Date = ?
#                 """, (symbol, latest_date)).fetchone()
#                 low_score = low_data[0] if low_data else None

#                 # Get high risk data
#                 high_data = conn.execute("""
#                     SELECT Score, Score_HoldPeriod 
#                     FROM high_risk 
#                     WHERE Symbol = ? AND Date = ?
#                 """, (symbol, latest_date)).fetchone()
#                 high_score, hold_period = high_data if high_data else (None, None)

#                 # Get fundamentals
#                 fundamentals = conn.execute("""
#                     SELECT Fundamentals_Industry, Fundamentals_Sector,
#                            Fundamentals_PE, Fundamentals_PB,
#                            Fundamentals_Dividends, Fundamentals_ExDividendDate,
#                            Fundamentals_MarketCap, Fundamentals_Description
#                     FROM fundamentals 
#                     WHERE Symbol = ?
#                 """, (symbol,)).fetchone()

#                 if not fundamentals:
#                     continue

#                 (industry, sector, pe, pb, 
#                  dividend, ex_div, mcap, desc) = fundamentals

#                 # Format values
#                 dividend_pct = f"{dividend:.2f}%" if dividend else "none"
#                 ex_div_date = pd.to_datetime(ex_div).strftime('%m-%d-%Y') if ex_div else 'N/A'
#                 mcap_formatted = f"${mcap/1e9:.2f}B" if mcap else 'N/A'
#                 truncated_desc = f"{desc[:120]}..." if desc else ""

#                 stream_content.append({
#                     "symbol": symbol,
#                     "low_score": f"{low_score:.2%}" if low_score is not None else "N/A",
#                     "high_score": f"{high_score:.2%}" if high_score is not None else "N/A",
#                     "hold_period": f"{hold_period:.0f}d" if hold_period is not None else "N/A",
#                     "industry": industry,
#                     "sector": sector,
#                     "pe": f"{pe:.2f}" if pe is not None else "N/A",
#                     "pb": f"{pb:.2f}" if pb is not None else "N/A",
#                     "dividend": dividend_pct,
#                     "ex_div": ex_div_date,
#                     "mcap": mcap_formatted,
#                     "desc": truncated_desc
#                 })

#             except Exception as e:
#                 print(f"Error processing {symbol}: {str(e)}")

#         return stream_content

#     finally:
#         conn.close()


# def bubble_style():
#     return """
#     <style>
#         @keyframes float {
#             0%   { transform: translateY(0px);}
#             50%  { transform: translateY(-20px);}
#             100% { transform: translateY(0px);}
#         }
#         @keyframes focusBlur {
#             0%   { filter: blur(2.5px);}
#             25%  { filter: blur(2.5px);}
#             35%  { filter: blur(0.5px);}
#             65%  { filter: blur(0.5px);}
#             75%  { filter: blur(2.5px);}
#             100% { filter: blur(2.5px);}        }
#         .bubble-container {
#             position: relative;
#             width: 100%;
#             max-width: 100%;
#             height: 1100px;
#             box-sizing: border-box;
#         }
#         .bubble {
#             border-radius: 50%;
#             margin: 10px;
#             position: absolute;
#             /* Animate both floating and focus/blur */
#             animation: float 6s ease-in-out infinite, focusBlur 6s ease-in-out infinite;
#             backdrop-filter: blur(5px);
#             border: 1px solid rgba(255,255,255,0.18);
#             box-shadow: 0 8px 24px rgba(0,0,0,0.14), 0 1.5px 8px 2px rgba(255,255,255,0.08) inset;
#             transition: transform 0.3s ease;
#             display: flex;
#             flex-direction: column;
#             align-items: center;
#             justify-content: center;
#             overflow: hidden;
#         }
#         .bubble:hover {
#             transform: scale(1.06);
#         }
#         .bubble h3, .bubble p {
#             text-align: center;
#             margin: 0;
#             padding: 0 10px;
#             word-break: break-word;
#         }
#     </style>
#     """

# def display_bubbles(col, items):
#     html = "<div class='bubble-container'>"
#     n = len(items)
#     base_height = 1100  # px
#     # base_height = 750  # px
#     container_height = base_height + int((n-5)*220)
#     bubble_diameter = 220    # px
#     # bubble_diameter = 150    # px

#     if n > 1:
#         space_between = (container_height - bubble_diameter) // (n - 1)
#     else:
#         space_between = 0

#     for i, item in enumerate(items):
#         container_width = 150  # px, set to your actual container width
#         max_left_percent = 100 - (bubble_diameter / container_width * 100)
#         left = random.uniform(0, max_left_percent)        
#         # left = random.randint(5, 75)
#         hue = random.randint(0, 360)
#         # Much darker, richer gradient for a professional look
#         gradient = (
#             f"radial-gradient(circle at 35% 30%, "
#             f"hsla({hue}, 80%, 22%, 0.95) 0%, "     # highlight (dark, saturated)
#             f"hsla({hue}, 80%, 14%, 0.92) 55%, "    # mid-tone (very dark)
#             f"hsla({hue}, 90%, 7%, 0.95) 100%)"     # shadow (almost black)
#         )
#         top_px = i * space_between + random.randint(-8, 8)
#         duration = random.uniform(2.5, 5.5)
#         delay = random.uniform(0, 2)
#         html += f"""
#             <div class="bubble" style="
#                 left: {left}%;
#                 top: {top_px}px;
#                 width: {bubble_diameter}px;
#                 height: {bubble_diameter}px;
#                 background: {gradient};
#                 animation-duration: {duration}s;
#                 animation-delay: {delay}s;
#             ">
#                 <h3 style='color: hsl({hue}, 60%, 70%); font-size:1.2em; margin:0; padding:0;'>{item['symbol']}</h3>
#                 <div class="bubble-desc-scroll-x" style="margin:0; padding:0; margin-top:2px;">
#                     <div class="bubble-desc-scroll-x-inner">
#                         {item['desc']} &nbsp;&nbsp;&nbsp; {item['desc']}
#                     </div>
#                 </div>
#                 <p style='font-size: 0.92em; margin:0; padding:0 10px; text-align:center; word-break:break-word; color: #e0e0e0;'>
#                     🏭 {item['industry']}<br>
#                     📈 Low Risk: {item['low_score']}<br>
#                     🚀 High Risk: {item['high_score']}<br>
#                     💰 P/E: {item['pe']} | P/B: {item['pb']}<br>
#                     📅 Ex-Div: {item['ex_div']}<br>
#                     💵 Div: {item['dividend']}
#                 </p>
#             </div>
#         """
#         html += "</div>"
#     html += "</div>"
#     col.markdown(bubble_style() + html, unsafe_allow_html=True)


def display_bubbles(col, items):
    html = "<div class='bubble-container'>"
    n = len(items)
    base_height = 750  # px
    container_height = base_height + int((n-5)*150)
    bubble_diameter = 150    # px  <-- updated

    if n > 1:
        space_between = (container_height - bubble_diameter) // (n - 1)
    else:
        space_between = 0

    for i, item in enumerate(items):
        container_width = 200  # px, set to your actual container width
        max_left_percent = 100 - (bubble_diameter / container_width * 100)
        left = random.uniform(0, max_left_percent)
        hue = random.randint(0, 360)
        gradient = (
            f"radial-gradient(circle at 35% 30%, "
            f"hsla({hue}, 80%, 22%, 0.95) 0%, "
            f"hsla({hue}, 80%, 14%, 0.92) 55%, "
            f"hsla({hue}, 90%, 7%, 0.95) 100%)"
        )
        top_px = i * space_between + random.randint(-8, 8)
        duration = random.uniform(2.5, 5.5)
        delay = random.uniform(0, 2)
        html += f"""
            <div class="bubble" style="
                left: {left}%;
                top: {top_px}px;
                width: {bubble_diameter}px;
                height: {bubble_diameter}px;
                background: {gradient};
                animation-duration: {duration}s;
                animation-delay: {delay}s;
            ">
                <h3 style='color: hsl({hue}, 60%, 70%); font-size:0.9em; margin:0; padding:0;'>{item['symbol']}</h3>
                <div class="bubble-desc-scroll-x" style="margin:0; padding:0; margin-top:2px;">
                    <div class="bubble-desc-scroll-x-inner" style="font-size:0.8em;">
                        {item['desc']} &nbsp;&nbsp;&nbsp; {item['desc']}
                    </div>
                </div>
                <p style='font-size: 0.7em; margin:0; padding:0 6px; text-align:center; word-break:break-word; color: #e0e0e0;'>
                    🏭 {item['industry']}<br>
                    🚀 High Rank: {item['high_score']}<br>
                    📈 Low Rank: {item['low_score']}<br>
                    💰 P/E: {item['pe']} | P/B: {item['pb']}<br>
                    📅 Ex-Div: {item['ex_div']}<br>
                    💵 Div: {item['dividend']}
                </p>
            </div>
        """
        html += "</div>"
    html += "</div>"
    col.markdown(bubble_style() + html, unsafe_allow_html=True)


# def bubble_style():
#     return """
#     <style>
#         @keyframes float {
#             0%   { transform: translateY(0px);}
#             50%  { transform: translateY(-20px);}
#             100% { transform: translateY(0px);}
#         }
#         @keyframes focusBlur {
#             0%   { filter: blur(2.5px);}
#             25%  { filter: blur(2.5px);}
#             35%  { filter: blur(0.5px);}
#             65%  { filter: blur(0.5px);}
#             75%  { filter: blur(2.5px);}
#             100% { filter: blur(2.5px);}
#         }
#         @keyframes scroll-horizontal {
#             0%   { transform: translateX(0%);}
#             100% { transform: translateX(-50%);}
#         }
#         .bubble-container {
#             position: relative;
#             width: 100%;
#             max-width: 100%;
#             height: 1100px;
#             box-sizing: border-box;
#         }
#         .bubble {
#             border-radius: 50%;
#             margin: 10px;
#             position: absolute;
#             animation: float 6s ease-in-out infinite, focusBlur 6s ease-in-out infinite;
#             backdrop-filter: blur(5px);
#             border: 1px solid rgba(255,255,255,0.18);
#             box-shadow: 0 8px 24px rgba(0,0,0,0.14), 0 1.5px 8px 2px rgba(255,255,255,0.08) inset;
#             transition: transform 0.3s ease;
#             display: flex;
#             flex-direction: column;
#             align-items: center;
#             justify-content: center;
#             overflow: hidden;
#         }
#         .bubble-desc-scroll-x {
#             width: 90%;
#             height: 2em;
#             overflow: hidden;
#             background: transparent;
#             margin: 0 auto;
#             position: relative;
#             white-space: nowrap;
#         }
#         .bubble-desc-scroll-x-inner {
#             display: inline-block;
#             white-space: nowrap;
#             animation: scroll-horizontal 18s linear infinite;
#         }
#     </style>
#     """
            # border: 1px solid rgba(255,255,255,0.18);

def bubble_style():
    return """
    <style>
        @keyframes float {
            0%   { transform: translateY(0px);}
            50%  { transform: translateY(-20px);}
            100% { transform: translateY(0px);}
        }
        @keyframes focusBlur {
            0%   { filter: blur(2.5px);}
            25%  { filter: blur(2.5px);}
            35%  { filter: blur(0.5px);}
            65%  { filter: blur(0.5px);}
            75%  { filter: blur(2.5px);}
            100% { filter: blur(2.5px);}
        }
        @keyframes scroll-horizontal {
            0%   { transform: translateX(0%);}
            100% { transform: translateX(-50%);}
        }
        .bubble-container {
            position: relative;
            width: 100%;
            max-width: 100%;
            height: 1100px;
            box-sizing: border-box;
        }
        .bubble {
            border-radius: 50%;
            margin: 10px;
            position: absolute;
            animation: float 6s ease-in-out infinite;
            backdrop-filter: blur(5px);
            box-shadow: 0 8px 24px rgba(0,0,0,0.14), 0 1.5px 8px 2px rgba(255,255,255,0.08) inset;
            transition: transform 0.3s ease;
            display: flex;
            flex-direction: column;
            align-items: center;
            justify-content: center;
            overflow: hidden;
        }
        .bubble-desc-scroll-x {
            width: 90%;
            height: 2em;
            overflow: hidden;
            background: transparent;
            margin: 0 auto;
            position: relative;
            white-space: nowrap;
        }
        .bubble-desc-scroll-x-inner {
            display: inline-block;
            white-space: nowrap;
            animation: scroll-horizontal 37s linear infinite;
        }
    </style>
    """


    
if show_top_symbols:    
    top_symbols = generate_top_10_stream(db_file_used)

# Verify high_risk table contents
# st.write("Sample Data from 'high_risk':", execute_query("SELECT * FROM high_risk LIMIT 5;"))

# Verify low_risk table contents
# st.write("Sample Data from 'low_risk':", execute_query("SELECT * FROM low_risk LIMIT 5;"))

# # Database connection
# db_file = "zoltar_financial.db"
# db_conn = sqlite3.connect(db_file)

# Test database functions
# print("Tables in Database:", list_tables())
# print("Schema for 'shap_summary':", describe_table("shap_summary"))
# print("Sample Data from 'shap_summary':", execute_query("SELECT * FROM shap_summary LIMIT 5;"))

# Define tools for chatbot
db_tools = [list_tables, describe_table, execute_query]

# Instruction for chatbot
instruction = """You are a helpful chatbot that can interact with an SQL database
for Stock trading education app. You will take the users' questions and turn them into SQL
queries using the tools available. Once you have the information you need, you will
answer the user's question using the data returned.  
high risk scores should be communicated as high Zoltar Ranks in context, and low risk scores are low Zoltar Ranks for context.  
These scores predict returns - high is for best return in next 14 days, and low is average expected return for the next 14 days. 
User is usually interested in high returns, and if stable returns are preferred, low risk scores (low zoltar rank) should be used, 
with sorting always done with highest values on top.  
If user is interested in ratings, go to ratings_detail and get necessary data (by Symbol). the RatingsPlblishedAt example format(2025-03-21T11:52:24Z) is not a timestamp format (but can extract timestamp info)


Use list_tables to see what tables are present, describe_table to understand the schema, and execute_query to issue an SQL SELECT query. When recommending an action, you have to take that action.
Be mindful of space used and limit as much as possible upfront in SQL queries output.


Here's avaliable data:
            -- Table for high-risk stocks
            CREATE TABLE IF NOT EXISTS high_risk (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Date DATETIME,
                Symbol TEXT,
                Score REAL,
                Score_Sharpe REAL,
                Score_HoldPeriod REAL,
                Close_Price REAL,
                Cap_Size TEXT,
                Sector TEXT,
                Industry TEXT,
                source TEXT
            );

            -- Table for low-risk stocks
            CREATE TABLE IF NOT EXISTS low_risk (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Date DATETIME,
                Symbol TEXT,
                Score REAL,
                Score_Sharpe REAL,
                Score_HoldPeriod REAL,
                Close_Price REAL,
                Cap_Size TEXT,
                Sector TEXT,
                Industry TEXT,
                source TEXT
            );
            -- Table for high-risk stocks
            CREATE TABLE IF NOT EXISTS all_high_risk (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Date DATETIME,
                Symbol TEXT,
                Score REAL,
                Score_Sharpe REAL,
                Score_HoldPeriod REAL,
                Close_Price REAL,
                Cap_Size TEXT,
                Sector TEXT,
                Industry TEXT,
                source TEXT
            );

            -- Table for low-risk stocks
            CREATE TABLE IF NOT EXISTS all_low_risk (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Date DATETIME,
                Symbol TEXT,
                Score REAL,
                Score_Sharpe REAL,
                Score_HoldPeriod REAL,
                Close_Price REAL,
                Cap_Size TEXT,
                Sector TEXT,
                Industry TEXT,
                source TEXT
            );

            -- Table for fundamentals
            CREATE TABLE IF NOT EXISTS fundamentals (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Symbol TEXT,
                Fundamentals_OverallRating REAL,
                total_ratings INTEGER,
                Fundamentals_Sector TEXT,
                Fundamentals_Industry TEXT,
                Fundamentals_Dividends REAL,
                Fundamentals_PE REAL,
                Fundamentals_PB REAL,
                Fundamentals_MarketCap REAL,
                Fundamentals_avgVolume2Weeks REAL,
                Fundamentals_avgVolume30Days REAL,
                Fundamentals_52WeekHigh REAL,
                Fundamentals_52WeekLow REAL,
                Fundamentals_52WeekHighDate DATE,
                Fundamentals_52WeekLowDate DATE,
                Fundamentals_Float REAL,
                Fundamentals_SharesOutstanding INTEGER,
                Fundamentals_CEO TEXT,
                Fundamentals_NumEmployees INTEGER,
                Fundamentals_YearFounded INTEGER,
                Fundamentals_ExDividendDate DATE,
                Fundamentals_PayableDate DATE,
                Fundamentals_Description TEXT
            );

            -- Table for ratings detail
            CREATE TABLE IF NOT EXISTS ratings_detail (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                Symbol TEXT,
                RatingType TEXT,
                RatingText TEXT,
                RatingPublishedAt DATETIME
            
            CREATE TABLE IF NOT EXISTS shap_summary_Large (CHECK CONTENTS FOR COLUMN NAMES)
            CREATE TABLE IF NOT EXISTS shap_summary_Mid (CHECK CONTENTS FOR COLUMN NAMES)
            CREATE TABLE IF NOT EXISTS shap_summary_Small (CHECK CONTENTS FOR COLUMN NAMES)

    SHAP REASONS ARE NOT IN FUNDAMENTALS - THEY ARE LOCATED IN 3 SEPARATE DATASETS: Additionally, use database tools to examine shap_summary_Large, shap_summary_Small, shap_summary_Mid in SQLite3 database (using the database tool) that contain SHAPLEY explanations for ML results on top stocks in corresponding tables with Symbol to merge with other tabes - if not in there it is not in top stocks currently.:

    all_high_risk and all_low_risk contain intraday production runs of Zoltar Ranks (Date column), and the high_risk and low_risk without "all" contain only daily production runs (but go further back from now) - also with Date column. Unless there is a reason to look at only the most recent intraday data, there is no reason to use 'all' datasets.
    when user asks for time-related tasks, Date column should be used in congjuction with Symbol, which represent Tickers, or Stocks.
    fundamentals dataset is updated only once a day, all_ datasets contains intraday data (Date) , and low_risk and high_risk contain daily data.  
    When user wants the most recent trends, always take the mox(Date) for the answer for each Symbol, and all_ files usually provide a better answer. For long-term trends the other ones are used.
    

The tables are related to each other by Symbol, and additionally by Date if available.  data for SHAP can be merged by knowing Cap_Size in high_risk and low_risk (and all_ versions also have these).
Important: Since many dates are available for same Symbol in high_risk and low_risk data, only the latest date should be used for most queries (unless explicitly stated otherwise)
Always order by descending date first (pick only records with max date unless stated otherwise), then descending Returns for the final answer, and sometimes in order of descending dividends.
When user requests Top stocks, they mean stocks with highest expected returns (highest Zoltar Ranks - low or high, depending on preference), and prefers
Ensure final answer meets all criteria set by the user request, and the answer contains non-duplicate symbols that look at the most recent data point, and mention the date used in the answer.
When stocks symbols are presented, also mention current price, and a few ratings/explanations, and when some information is missing, work with the information that is available (fundamentals and SHAP data could be missing).
When the user asks for Top stocks without mentioning High or Low, assume stocks that are in the top 10 for both Low and High Zoltar Ranks are needed.
When user asks for reasons for stocks being selected, refer to SHAP datasets using Cap_Size and Symbol (can check all 3 by Symbol)
When user asks for alpha, the comparison with SPY returns needs to be made.

User prefers the answer in a table format with relevant statistics, and a summary brief.  Response always ends with the phrase 'May the riches be with you...'

"""



# 5.7.25 - new from jupyter notebook


# # Main section with using genai gemini model
# is_retriable = lambda e: (isinstance(e, genai.errors.APIError) and e.code in {429, 503})

# if not hasattr(genai.models.Models.generate_content, '__wrapped__'):
#   genai.models.Models.generate_content = retry.Retry(
#       predicate=is_retriable)(genai.models.Models.generate_content)
# client = genai.Client(api_key=GOOGLE_API_KEY)



# # Model configuration for Gemini-2.0-Flash
# model_config = types.GenerateContentConfig(
#     temperature=0.1,
#     top_p=0.95,
#     system_instruction=instruction,
#     tools=db_tools,
# )

#5.26.25 -  Agent settings
st.sidebar.header("Agent Configuration")
st.sidebar.write("**News sources selection:**")

# Create 3 columns in the sidebar
col1side, col2side = st.sidebar.columns(2)

with col1side:
    google_trends = st.checkbox("Google Trends", value=False)
    stocktwits = st.checkbox("StockTwits", value=True)
    zacks = st.checkbox("Zacks", value=False)
    seeking = st.checkbox("SeekingAlpha", value=True)

with col2side:
    reddit = st.checkbox("Reddit", value=True)
    sentimentrader = st.checkbox("Yahoo Finance", value=False)
    tipranks = st.checkbox("TipRanks", value=True)
    nasdaq = st.checkbox("NASDAQ", value=True)

selected_sources = []
if google_trends: selected_sources.append("Google Trends (https://trends.google.com/)")
if stocktwits: selected_sources.append("StockTwits (https://stocktwits.com/")
if sentimentrader: selected_sources.append("Yahoo (https://finance.yahoo.com/)")
if tipranks: selected_sources.append("TipRanks (https://www.tipranks.com/")
if zacks: selected_sources.append("Zacks (https://www.zacks.com/)")
if reddit: selected_sources.append("Reddit (https://www.reddit.com/)")
if seeking: selected_sources.append("SeekingAlpha (https://seekingalpha.com/)")
if nasdaq: selected_sources.append("NASDAQ.com (https://www.nasdaq.com/market-activity/stocks)")
    
source_str = ", ".join(selected_sources) if selected_sources else "no sources selected"




st.sidebar.write("**Visualization selection:**")

# Create 3 columns in the sidebar
col1side, col2side = st.sidebar.columns(2)

with col1side:
    Pie_chart = st.checkbox("Pie Chart", value=False)
    Return_hold = st.checkbox("Returns", value=True)
    returns_trend = st.checkbox("Returns Trend", value=False)

with col2side:
    low_ranks_trend = st.checkbox("Ranks Trend", value=True)
    Price_trend = st.checkbox("Price", value=True)
    recommendations_table = st.checkbox("Summary", value=False)


# Map checkbox variables to their prompt instructions
viz_instructions = []

if Pie_chart:
    viz_instructions.append("- Industry: Pie Chart of Industries of selected stocks")
if Return_hold:
    viz_instructions.append("- Expected Returns: line chart for each of the selected stocks with two points for each - first point starting at (0,0) and second point X is number of days to hold (Score_HoldPeriod in high_risk and all_high_risk tables) vs High Zoltar Rank (y-axis), making starting point for x-axis max(Date) and iterating days forward from that point.")
if low_ranks_trend:
    viz_instructions.append("- Low Zoltar Rank Over Time: a pretty line chart of Low Zoltar Rank of each stock over time")
if recommendations_table:
    viz_instructions.append("- Recommendations: Table of model recommendations for each stock")
if returns_trend:
    viz_instructions.append("- High Zoltar Rank Over Time: a pretty line chart of High Zoltar Rank of each stock over time")
if Price_trend:
    viz_instructions.append("- Price Over Time: a pretty line chart of Price of each stock over time (from high_risk table)")


# Join instructions for prompt
if viz_instructions:
    viz_section = "\n".join(viz_instructions)
else:
    viz_section = "- No visualizations selected."


# # Sidebar sliders for tuning model parameters
# st.sidebar.header("Model Configuration")
# temperature = st.sidebar.slider(
#     "Temperature", 
#     min_value=0.0, 
#     max_value=1.0, 
#     value=0.1,  # Default value
#     step=0.05
# )
# top_p = st.sidebar.slider(
#     "Top-p", 
#     min_value=0.0, 
#     max_value=1.0, 
#     value=0.95,  # Default value
#     step=0.05
# )

st.sidebar.header("Model Configuration")

# --- Define levels and their mappings ---
temp_levels = [("0.0 - Exact", 0.0), ("0.1 - Middle", 0.1), ("1.0 - Wild", 1.0)]
top_p_levels = [("0.5 - Wild", 0.7), ("0.9 - Middle", 0.9), ("1.0 - Exact", 1.0)]

# --- Helper function for segmented buttons ---
def segmented_buttons(label, levels, key_prefix):
    cols = st.sidebar.columns(len(levels))
    selected = st.session_state.get(f"{key_prefix}_selected", levels[0][0])
    for i, (level, _) in enumerate(levels):
        button_kwargs = {}
        if selected == level:
            button_kwargs["type"] = "primary"
        if cols[i].button(level, key=f"{key_prefix}_{level}", **button_kwargs):
            st.session_state[f"{key_prefix}_selected"] = level
            selected = level
    return dict(levels)[selected]

# --- Render segmented buttons ---
st.sidebar.write("Temperature setting:")
temperature = segmented_buttons("Temperature Level", temp_levels, "temp")
# st.sidebar.markdown(
#     f"**Temperature:** {st.session_state['temp_selected']} ({temperature})\n\n"
# )

st.sidebar.write("Top-p setting:")

top_p = segmented_buttons("Top-p Level", top_p_levels, "top_p")



st.sidebar.markdown(
    """
    <style>
    .disclaimer-btn {
        background: linear-gradient(135deg, #301934 0%, #9370DB 100%);
        border: none;
        color: #fff;
        padding: 10px 22px;
        text-align: center;
        text-decoration: none;
        display: inline-block;
        font-size: 15px;
        font-weight: 600;
        margin: 10px 2px;
        border-radius: 10px;
        cursor: pointer;
        box-shadow: 0 4px 14px 0 rgba(80, 40, 120, 0.25);
        transition: all 0.18s cubic-bezier(.4,0,.2,1);
        outline: none;
    }
    .disclaimer-btn:hover, .disclaimer-btn:focus {
        background: linear-gradient(135deg, #9370DB 0%, #301934 100%);
        box-shadow: 0 8px 24px 0 rgba(80, 40, 120, 0.35);
        transform: translateY(-1px) scale(1.03);
    }
    </style>
    <a href="https://github.com/apod-1/ZoltarFinancial/raw/main/docs/User_Agreement.txt"
       target="_blank"
       title="By using this app, you agree to the terms and conditions. Not investment advice.">
        <button class="disclaimer-btn"
                title="By using this app, you agree to the terms and conditions. This is not investment advice.">
            View Disclaimer
        </button>
    </a>
    """,
    unsafe_allow_html=True
)
                                    
# --- Show current values ---
# st.sidebar.markdown(
#     f"**Top-p:** {st.session_state['top_p_selected']} ({top_p})"
# )

# Show the selected values (for debugging/demo)
# st.sidebar.write(f"Temperature value: {temperature}")
# st.sidebar.write(f"Top-p value: {top_p}")



# Model configuration for Gemini-2.0-Flash
model_config = types.GenerateContentConfig(
    temperature=temperature,
    top_p=top_p,
    system_instruction=instruction,
    tools=db_tools,
)
# Initialize Google GenAI client (ensure GOOGLE_API_KEY is set)
# from google import genai

client = genai.Client(api_key=GOOGLE_API_KEY)

# Start a chat with automatic function calling enabled
chat = client.chats.create(
    model="gemini-2.0-flash",
    config=model_config,
)
result=None

# 5.24 - helper functions
def to_json_serializable(obj):
    if isinstance(obj, str):
        return obj
    try:
        return json.dumps(obj, default=str)
    except Exception as e:
        return str(obj)

def is_blank_png(img_bytes):
    # Quick check for empty or very small files
    if not img_bytes or len(img_bytes) < 8500:
        return True

    try:
        with Image.open(BytesIO(img_bytes)) as img:
            img = img.convert("RGBA")  # Ensure 4 channels
            extrema = img.getextrema()  # Returns (min, max) for each channel

            # Check if all pixels are fully transparent
            if extrema[3] == (0, 0):
                return True

            # Check if all pixels are white (255,255,255,255)
            if all(channel == (255, 255) for channel in extrema):
                return True

            # For grayscale/other modes, check if all pixels are same value
            if all(e[0] == e[1] for e in extrema):
                return True

            # Optionally, check if all pixels are the same color
            if len(set(img.getdata())) == 1:
                return True

    except Exception as e:
        # If PIL fails to open, treat as blank/invalid
        print(f"Image validation error: {e}")
        return True

    return False

import sys
def debug_payload(message):
    try:
        msg_str = to_json_serializable(message)
        if isinstance(msg_str, str):
            size = len(msg_str.encode('utf-8'))
        else:
            size = sys.getsizeof(msg_str)
        print(f"Payload size: {size} bytes")
        if size > 1_000_000:
            st.warning("Payload is very large and may be rejected by Gemini API.")
    except Exception as e:
        st.write(f"Could not serialize payload: {e}")


def prepare_image_for_gemini(image_path):
    if not os.path.exists(image_path):
        return None
    with open(image_path, "rb") as f:
        img_bytes = f.read()
    if is_blank_png(img_bytes):
        return None
    # Downscale if too large
    if len(img_bytes) > 2_000_000:
        img = Image.open(BytesIO(img_bytes))
        img.thumbnail((800, 800))
        buffer = BytesIO()
        img.save(buffer, format="PNG", optimize=True)
        img_bytes = buffer.getvalue()
    img_b64 = base64.b64encode(img_bytes).decode("utf-8")
    return img_b64

# # Example query to chatbot
# response = chat.send_message("what stocks have highest low Zoltar Rank, averaged over the last 5 data points? put in a table with Low and High Zoltar Ranks shown.")
# # Streamlit UI for user input
# st.title("Chatbot Query Interface")

# 5.28.25 - bubbles
    # Split symbols between columns
# if top_symbols:
#     mid = len(top_symbols) // 2
#     # mid = top_n1
#     # mid=3
#     with col1:
#         st.write("Top Low Zoltar Rank Stocks")
#         display_bubbles(col1, top_symbols[:mid])
#     with col3:
#         st.write("Top High Zoltar Rank Stocks")
#         display_bubbles(col3, top_symbols[mid:])
if top_symbols:
    mid = len(top_symbols) // 2
    # mid = top_n1  #7.12.25 - made it half again (need to adjust later to actual size and figure out why it doesnt work all the time for n1)
    # mid = top_low
    with col1:
        st.markdown(
            "<div style='text-align:center; font-size:1em; font-weight:600; color:#b22222; margin-bottom:0.2em;'>"
            "Top <span style='color:#DAA520;'>Low Zoltar Rank</span> Stocks"
            "</div>",
            unsafe_allow_html=True
        )
        display_bubbles(col1, top_symbols[:mid])
    
    with col3:
        st.markdown(
            "<div style='text-align:center; font-size:1em; font-weight:600; color:#b22222; margin-bottom:0.2em;'>"
            "Top <span style='color:#DAA520;'>High Zoltar Rank</span> Stocks"
            "</div>",
            unsafe_allow_html=True
        )
        display_bubbles(col3, top_symbols[mid:])     
print("Top symbols:", top_symbols)  # Add this line
# print("Latest date in low_risk:", latest_date)


# #5.31.25 -  testing of shap tables
# def check_shap_tables(db_path='zoltar_database.sqlite3'):
#     # conn = sqlite3.connect(db_path)
    
#     # Check if SHAP tables exist
#     tables = pd.read_sql("""
#         SELECT name 
#         FROM sqlite_master 
#         WHERE type='table' 
#         AND name LIKE 'shap_summary_%'
#     """, db_conn)
    
#     if tables.empty:
#         print("ERROR: No SHAP tables found in the database")
#         return
    
#     # Check columns in first SHAP table
#     table_name = tables['name'].iloc[0]
#     columns = pd.read_sql(f"PRAGMA table_info({table_name})", db_conn)
    
#     # conn.close()
#     return tables, columns

# tables, columns = check_shap_tables()
# print("SHAP Tables:\n", tables)
# print("\nColumns in First SHAP Table:\n", columns)

# def create_shap_table(symbols, db_path='zoltar_database.sqlite3'):
#     # conn = sqlite3.connect(db_path)
#     all_shap = []
    
#     # Get list of SHAP tables
#     shap_tables = pd.read_sql("""
#         SELECT name 
#         FROM sqlite_master 
#         WHERE type='table' 
#         AND name LIKE 'shap_summary_%'
#     """, db_conn)['name'].tolist()
    
#     for symbol in symbols:
#         symbol_data = []
        
#         for table in shap_tables:
#             # Get most recent data for symbol
#             query = f"""
#                 SELECT * 
#                 FROM {table} 
#                 WHERE Symbol = '{symbol}'
#                 LIMIT 1
#             """
#             df = pd.read_sql(query, db_conn)
            
#             if not df.empty:
#                 # Process SHAP values
#                 numeric_cols = df.select_dtypes(include='number').columns
#                 for col in numeric_cols:
#                     value = df[col].values[0]
#                     if pd.notnull(value) and value != 0:
#                         symbol_data.append({
#                             'Symbol': symbol,
#                             'SHAP Table': table,
#                             'Feature': col,
#                             'SHAP Value': f"{value:.9f}",
#                             'Impact': "Increasing" if value > 0 else "Decreasing"
#                         })
        
#         if not symbol_data:
#             all_shap.append(pd.DataFrame({
#                 'Symbol': [symbol],
#                 'Status': ['No SHAP data found']
#             }))
#         else:
#             all_shap.append(pd.DataFrame(symbol_data))
    
#     # conn.close()
#     return pd.concat(all_shap).reset_index(drop=True)

# # Test with your symbols
# symbols = ['F', 'RIVN', 'NIO', 'XPEV', 'LI']
# shap_results = create_shap_table(symbols)
# st.write(shap_results)


with col2:
   
    # Input text box for the user's query
    user_query = st.text_input(
        "",
        # value="Provide stocks with best expected returns (latest date) and a dividend yield above 5, with dividend date coming up within a month from now.  create a table and a summary.",
        value="Best stocks to get now?",
        help="Ask about best stocks, dividends, sectors, explanations (anything stocks related)"
        ,placeholder = "Ask your stock-related question..."
    )
    # Submit button
    if st.button("Submit Query"):
        prep_db = st.toast("UPDATING ZOLTAR DATABASE...", icon="⏳")

        #reset the repo
        # st.session_state.agent_repo2 =  st.session_state.agent_repo

        # Save to JSON file
        with open("agent_repo_t.json", "w") as f:
            json.dump(st.session_state.agent_repo, f)
        

        #reset the repo
        st.session_state.agent_repo = {
            "agents": {},
            "execution_order": []
        }

        # 5.26.25 - initialize session state variables 
        st.session_state.final_agent_result = ""
        st.session_state.agent_progress = {}
        def add_agent_result(agent_key, agent_data):
            # Add agent result to the repo if not already present in execution_order
            if agent_key not in st.session_state.agent_repo["agents"]:
                st.session_state.agent_repo["agents"][agent_key] = agent_data
            if agent_key not in st.session_state.agent_repo["execution_order"]:
                st.session_state.agent_repo["execution_order"].append(agent_key)         
        
        # Send the query to the chatbot
        #response = chat.send_message(user_query)
    
        # Display the chatbot's response
        #st.write("Chatbot Response:")
        #st.write(response.text)
    
        # Optional: Print chat history for debugging
        def print_chat_turns(chat):
            """Prints out each turn in the chat history, including function calls and responses."""
            for event in chat.get_history():
                st.write(f"{event.role.capitalize()}:")
    
                # for part in event.parts:
                #     if txt := part.text:
                #         st.write(f'  "{txt}"')
                #     elif fn := part.function_call:
                #         args = ", ".join(f"{key}={val}" for key, val in fn.args.items())
                #         st.write(f"  Function call: {fn.name}({args})")
                #     elif resp := part.function_response:
                #         st.write("  Function response:")
                #         st.write(resp.response['result'])
                for part in event.parts:
                    if code := part.executable_code:
                        st.markdown(f"### Code\n``````")
                
                    elif result := part.code_execution_result:
                        st.markdown(f"### Result: {result.outcome}\n``````")
                
                    elif img := part.inline_data:
                        try:
                            # Validate and display the image
                            image = Image.open(io.BytesIO(img.data))
                            st.image(image, caption="Generated Image")
                        except Exception as e:
                            st.error(f"Error displaying image: {e}")
                st.write()
    
        # # Display chatbot response
        # # st.write("Chatbot Response:", response)
        # st.write("Chatbot Response:")
        # st.write(f"\n{response.text}")
        # # Print chat turns (optional)
        # print_chat_turns(chat)
    
    
        class DefaultAPI:
            def execute_query(self, sql):
                # Mock result for demonstration purposes
                return {
                    "result": [
                        ["2025-04-07 00:00:00", "MNSO", 0.0514, 16.61, "Large"],
                        ["2025-03-24 22:23:03", "SMCI", 0.0383, 41.42, "Mid"],
                        ["2025-03-21 19:20:35", "HPQ", 0.0259, 28.68, "Large"],
                        ["2025-03-21 19:20:35", "SMCI", 0.0257, 42.42, "Mid"],
                        ["2025-03-26 19:10:39", "MIRM", 0.0257, 46.28, "Small"]
                    ]
                }
        
        # Initialize default_api
        default_api = GOOGLE_API_KEY #DefaultAPI() 
        
     #     return all_responses
        async def handle_response_refresh(stream, tool_impl=None):
            """Stream output and handle any tool calls during the session."""
            all_responses = []
        
            async for msg in stream.receive():
                all_responses.append(msg)
        
                if text := msg.text:
                    # Output any text chunks that are streamed back.
                    if len(all_responses) < 2 or not all_responses[-2].text:
                        # Display a header if this is the first text chunk.
                        st.markdown('### Text')
                    st.write(text)
        
                elif tool_call := msg.tool_call:
                    # Handle tool-call requests.
                    for fc in tool_call.function_calls:
                        st.markdown('### Tool call')
        
                        # Execute the tool and collect the result to return to the model.
                        if callable(tool_impl):
                            try:
                                result = tool_impl(**fc.args)
                            except Exception as e:
                                result = str(e)
                        else:
                            result = 'ok'
        
                        tool_response = types.LiveClientToolResponse(
                            function_responses=[types.FunctionResponse(
                                name=fc.name,
                                id=fc.id,
                                response={'result': result},
                            )]
                        )
                        await stream.send(input=tool_response)
        
                elif msg.server_content and msg.server_content.model_turn:
                    for part in msg.server_content.model_turn.parts:
                        #st.write("Available attributes in Part:", dir(part))  # Debugging
        
                        if code := part.executable_code:
                            st.markdown("### Code")
                            st.code(code.code)
        
                            # Dynamically execute provided code
                            try:
                                exec_globals = {
                                    "pd": pd,
                                    "sns": sns,
                                    "plt": plt,
                                    "base64": base64,
                                    "BytesIO": BytesIO,
                                    "default_api": default_api,  # Pass default_api into context
                                }
                                exec(code.code, exec_globals)
        
                                # Decode base64 string and display image
                                if "image_base64" in exec_globals:
                                    image_base64 = exec_globals["image_base64"]
                                    img_bytes = base64.b64decode(image_base64)
                                    st.image(img_bytes, caption="Generated Plot", use_column_width=True)
                                else:
                                    st.warning("No plot was generated.")
                            except Exception as e:
                                st.error(f"Error executing code: {e}")
                        elif text := part.text:  # Fallback for text-based instructions
                            st.markdown("### Instructions for Plotting")
                            st.write(text)
        
            return all_responses
    
        
    #5.4.25 additions
    
        import base64
        #!pip install -U textblob
        #from textblob import TextBlob
        
        # Global state variable
        global_state = {
            "collected_text": "",
            "tool_calls": [],
            "code_results": [],
            "images": []
        }
        
        # Counter to track updates
        update_counter = 0
        
        def update_state(key, value):
            """Update the global state."""
            global global_state
            global_state[key] = value
            
        # def update_state(key, value):
        #     global global_state
        #     # Replace entire value instead of appending
        #     global_state[key] = value if not isinstance(value, list) else [value[-1]]      
        
        # def display_state():
        #     """Refresh the display with the latest state."""
        #     #st.rerun()
        #     #clear_output(wait=True)  # Clear previous output
        
        #     # Display images (latest only)
        #     if global_state["images"]:
        #         for img in global_state["images"]:
        #             st.image(img)
        
        #     # Display collected text
        #     if global_state["collected_text"].strip():
        #         st.markdown(f"### Text\n\n{global_state['collected_text'].strip()}")
        
        #     # Display tool calls (latest only)
        #     #if global_state["tool_calls"]:
        #      #   for tool_call in global_state["tool_calls"]:
        #       #      display(Markdown(f"### Tool Call\n\n{tool_call}"))
        
        #     # Display code results (latest only)
        #     if global_state["code_results"]:
        #         for code_result in global_state["code_results"]:
        #             st.markdown(f"### Code Result\n\n{code_result}")
        placeholder_container = st.empty()  # Master container for refreshable content
        
        # def display_state():
        #     """Dynamic refresh using placeholder replacement"""
        #     with placeholder_container.container():
        #         # Clear previous content
        #         #st.empty()  # Creates temporary empty space
                
        #         # Images with auto-clear
        #         if global_state["images"]:
        #             img_placeholder = st.empty()
        #             with img_placeholder:
        #                 for img in global_state["images"]:
        #                     # Your existing image handling logic
        #                     if isinstance(img, str):
        #                         img_bytes = base64.b64decode(img)
        #                         image = Image.open(BytesIO(img_bytes))
        #                         st.image(image)
        #                     elif isinstance(img, (bytes, bytearray)):
        #                         image = Image.open(BytesIO(img))
        #                         st.image(image)
        #                     elif isinstance(img, Image.Image):
        #                         st.image(img)
        #             img_placeholder.empty()  # Clear after render
        
        #         # Text with incremental updates  
        #         if global_state["collected_text"]:
        #             text_placeholder = st.empty()
        #             cleaned_text = "\n".join([
        #                 line for line in global_state["collected_text"].split("\n")
        #                 if "End of User Query" not in line
        #             ])
        #             text_placeholder.markdown(f"**Analysis**\n\n{cleaned_text}")    
        def is_base64_bytes(data):
            # PNG header is: b'\x89PNG\r\n\x1a\n'
            if data.startswith(b'\x89PNG\r\n\x1a\n'):
                return False
            # If it's all ASCII and decodes cleanly, likely base64
            try:
                base64.b64decode(data, validate=True)
                return True
            except Exception:
                return False
            
        def display_state():
            """Dynamic refresh using placeholder replacement"""
            with placeholder_container.container():
                # Images with auto-clear
                if global_state["images"]:
                    img_placeholder = st.empty()
                    with img_placeholder:
                        for img in global_state["images"]:
                            try:
                                # If it's a base64 string
                                if isinstance(img, str):
                                    img_bytes = base64.b64decode(img.strip())
                                    if img_bytes and len(img_bytes) > 0:
                                        st.image(img_bytes)
                                    else:
                                        st.warning("Decoded image is empty.")
                                # If it's raw bytes or bytearray
                                elif isinstance(img, (bytes, bytearray)):
                                    if img and len(img) > 0:
                                        if is_base64_bytes(img):
                                            print("Detected base64-encoded bytes, decoding...")
                                            img_bytes = base64.b64decode(img.strip())
                                            st.image(img_bytes)
                                        else:
                                            st.image(img)
                                    else:
                                        st.warning("Image bytes are empty.")
                                # If it's a PIL Image
                                elif isinstance(img, Image.Image):
                                    st.image(img)
                                else:
                                    st.warning(f"Unsupported image type: {type(img)}")
                            except Exception as e:
                                st.error(f"Could not display image: {e}")
                    img_placeholder.empty()  # Clear after render
        
                # Text with incremental updates  
                if global_state["collected_text"]:
                    text_placeholder = st.empty()
                    cleaned_text = "\n".join([
                        line for line in global_state["collected_text"].split("\n")
                        if "End of User Query" not in line
                    ])
                    text_placeholder.markdown(f"---\n\n{cleaned_text}")
    
    
        def display_state():
            """Dynamic refresh using placeholder replacement"""
            with placeholder_container.container():
                # --- Display Images FIRST ---
                if global_state["images"]:
                    for img in global_state["images"]:
                        try:
                            # If it's a base64 string
                            if isinstance(img, str):
                                img_bytes = base64.b64decode(img.strip())
                                if img_bytes and len(img_bytes) > 0:
                                    st.image(img_bytes)
                                else:
                                    st.warning("Decoded image is empty.")
                            # If it's raw bytes or bytearray
                            elif isinstance(img, (bytes, bytearray)):
                                if img and len(img) > 0:
                                    if is_base64_bytes(img):
                                        print("Detected base64-encoded bytes, decoding...")
                                        img_bytes = base64.b64decode(img.strip())
                                        st.image(img_bytes)
                                    else:
                                        st.image(img)
                                else:
                                    st.warning("Image bytes are empty.")
                            # If it's a PIL Image
                            elif isinstance(img, Image.Image):
                                st.image(img)
                            else:
                                st.warning(f"Unsupported image type: {type(img)}")
                        except Exception as e:
                            st.error(f"Could not display image: {e}")
        
                # --- Then Display Text and Other Content ---
                if global_state["collected_text"]:
                    text_placeholder = st.empty()
                    cleaned_text = "\n".join([
                        line for line in global_state["collected_text"].split("\n")
                        if "End of User Query" not in line
                    ])
                    text_placeholder.markdown(f"---\n\n{cleaned_text}")
    
                    
        def format_global_state(global_state):
            """Convert global_state into a structured string for agent prompts."""
            sections = []
    
            if global_state["images"]:
                for img in global_state["images"]:
                    print("img type:", type(img))
                    if isinstance(img, str):
                        print("First 100 chars:", img[:100])
                    elif isinstance(img, (bytes, bytearray)):
                        print("First 20 bytes:", img[:20])
                        print("Bytes length:", len(img))
                    else:
                        print("img is not str or bytes!")
            
                    # Handle base64 strings
                    if isinstance(img, str):
                        try:
                            img_bytes = base64.b64decode(img.strip())
                            st.image(img_bytes)  # Streamlit can handle PNG/JPEG bytes directly
                        except Exception as e:
                            st.error(f"Base64 decode error: {str(e)}")
            
                    # Handle raw bytes
                    elif isinstance(img, (bytes, bytearray)):
                        try:
                            if is_base64_bytes(img):
                                print("Detected base64-encoded bytes, decoding for display...")
                                img_bytes = base64.b64decode(img.strip())
                                st.image(img_bytes)
                            else:
                                st.image(img)
                        except Exception as e:
                            st.error(f"Bytes conversion error: {str(e)}")
            
                    # Handle PIL Images directly
                    elif isinstance(img, Image.Image):
                        st.image(img)
                        
            # Text Content
            if global_state["collected_text"]:
                sections.append(f"## Collected Text\n{global_state['collected_text']}")
            
            # Tool Calls
            if global_state["tool_calls"]:
                tool_list = "\n".join(f"- {call}" for call in global_state["tool_calls"])
                sections.append(f"## Tool Calls\n{tool_list}")
            
            # Code Results
            if global_state["code_results"]:
                code_list = "\n".join(f"- {result}" for result in global_state["code_results"])
                sections.append(f"## Code Execution Results\n{code_list}")
            
            # Images (reference filenames)
            # if global_state["images"]:
            #     img_list = "\n".join(f"- Image saved as: plot_{i+1}.png" for i in range(len(global_state["images"])))
            #     sections.append(f"## Generated Visualizations\n{img_list}")
    
            return "\n\n".join(sections)
        def format_global_state_notool(global_state):
            """Convert global_state into a structured string for agent prompts, omitting tool calls."""
            sections = []
            
            # Text Content (remove tool call lines)
            if global_state["collected_text"]:
                filtered_text = "\n".join(
                    line for line in global_state["collected_text"].splitlines()
                    if not line.strip().lower().startswith("tool call") and "db call" not in line.lower()
                )
                sections.append(f"## Collected Text\n{filtered_text}")
            
            # Code Results
            if global_state["code_results"]:
                code_list = "\n".join(f"- {result}" for result in global_state["code_results"])
                sections.append(f"## Code Execution Results\n{code_list}")
            
            # Images (reference filenames or other info if needed)
            # If you want to show info about images, you can uncomment below:
            # if global_state["images"]:
            #     img_list = "\n".join(f"- Image saved as: plot_{i+1}.png" for i in range(len(global_state["images"])))
            #     sections.append(f"## Generated Visualizations\n{img_list}")
            
            return "\n\n".join(sections)   
    
    
    
        async def handle_response_refresh(stream, tool_impl=None):
            """Stream output and handle any tool calls during the session."""
            global update_counter  # Use the global counter for tracking updates
            all_responses = []
            collected_text = ""  # Collect all text responses
            tool_call_results = []  # Temporary list for tool call results
            code_results = []  # Temporary list for code results
            images = []  # Temporary list for inline images
            MAX_BYTES = 1000000  # Leave 20% buffer
            current_size = 0
            retries = 2
            backoff = 1  # seconds                

            while retries > 0:
                try:
        
                    async for msg in stream.receive():
                        all_responses.append(msg)
                        msg_size = len(str(msg).encode('utf-8'))
                        if current_size + msg_size > MAX_BYTES:
                            print("Approaching size limit - truncating response")
                            #break
                        current_size += msg_size
                        
                        if text := msg.text:
                            # Collect text chunks into a single string
                            collected_text += text + " "
                            update_state("collected_text", collected_text)  # Update state with new text
                
                        # elif tool_call := msg.tool_call:
                        #     # Handle tool-call requests
                        #     tool_call_results = []  # Reset temporary list for tool calls
                        #     for fc in tool_call.function_calls:
                        #         st.markdown('### Tool call')
                
                        #         # Execute the tool and collect the result to return to the model
                        #         if callable(tool_impl):
                        #             try:
                        #                 result = tool_impl(**fc.args)
                        #             except Exception as e:
                        #                 result = str(e)
                        #         else:
                        #             result = 'ok'
                        elif tool_call := msg.tool_call:
                            tool_call_results = []
                            # for fc in tool_call.function_calls:
                            #     if callable(tool_impl):
                            #         try:
                            #             result = tool_impl(**fc.args)
                            #             # If result is a dict with 'call' and 'results'
                            #             if isinstance(result, dict) and 'call' in result:
                            #                 tool_call_results.append(result['call'])
                            #                 # Optionally, store the actual results somewhere else
                            #                 code_results.append(result['results'])
                            #             else:
                            #                 tool_call_results.append(str(result))
                            #         except Exception as e:
                            #             tool_call_results.append(str(e))
                            #     else:
                            #         tool_call_results.append('ok')
                            # # update_state("tool_calls", tool_call_results)    
                            #     tool_response = types.LiveClientToolResponse(
                            #         function_responses=[types.FunctionResponse(
                            #             name=fc.name,
                            #             id=fc.id,
                            #             response={'result': result},
                            #         )]
                            #     )
                            #     await stream.send(input=tool_response)
                
                            #     # Add result to temporary list
                            #     tool_call_results.append(result)
                            for fc in tool_call.function_calls:
                                if callable(tool_impl):
                                    try:
                                        result = tool_impl(**fc.args)
                                        if isinstance(result, dict) and 'call' in result:
                                            tool_call_results.append(result['call'])
                                            code_results.append(result['results'])
                                        else:
                                            tool_call_results.append(str(result))
                                    except Exception as e:
                                        result = str(e)
                                        tool_call_results.append(result)
                                else:
                                    result = 'ok'
                                    tool_call_results.append(result)
                            
                                # tool_response = types.LiveClientToolResponse(
                                #     function_responses=[types.FunctionResponse(
                                #         name=fc.name,
                                #         id=fc.id,
                                #         response={'result': result},
                                #     )]
                                # )
                                tool_response = types.LiveClientToolResponse(
                                    function_responses=[types.FunctionResponse(
                                        name=fc.name,
                                        id=fc.id,
                                        response={
                                            'result': json.dumps(result)  # Serialize to string
                                        }
                                    )]
                                )                                
                                await stream.send(input=tool_response)        
                            # Replace previous tool calls with the latest ones
                            update_state("tool_calls", tool_call_results)
                            # update_state("tool_calls", [tool_call_results[-1]])
                
                        elif msg.server_content and msg.server_content.model_turn:
                            # Handle code execution results and inline images
                            code_results = []  # Reset temporary list for code results
                            images = []  # Reset temporary list for inline images
                
                            for part in msg.server_content.model_turn.parts:
                                if code := part.executable_code:
                                    code_results.append(code)
                
                                elif result := part.code_execution_result:
                                    code_results.append(result.outcome)
                
                                elif img := part.inline_data:
                                    images.append(img.data)
                                # Save the first image (or all images)
                                # if images:
                                #     first_img = images[0]
                                #     # Check if base64 string or bytes
                                #     if isinstance(first_img, str):
                                #         img_bytes = base64.b64decode(first_img)
                                #     else:
                                #         img_bytes = first_img  # assume bytes
                            
                                #     with open("stock_price_plot.png", "wb") as f:
                                #         f.write(img_bytes)
                                #     print("Image saved as stock_price_plot.png")
                                #     update_state("images", images)
                                if images:
                                    first_img = images[0]
                                    img_bytes = None
                                    # String: decode base64
                                    if isinstance(first_img, str):
                                        img_bytes = base64.b64decode(first_img.strip())
                                    # Bytes: check if base64 or PNG
                                    elif isinstance(first_img, (bytes, bytearray)):
                                        if is_base64_bytes(first_img):
                                            print("Detected base64-encoded bytes, decoding...")
                                            img_bytes = base64.b64decode(first_img.strip())
                                            print(img_bytes)
                                        else:
                                            img_bytes = first_img
                                    else:
                                        print(f"Unsupported image type: {type(first_img)}")
                                        img_bytes = b""
                                
                                    # print(f"Type: {type(img_bytes)}, Length: {len(img_bytes)}")
                                    # print("First 20 bytes:", img_bytes[:20])
                                    # print("Last 20 bytes:", img_bytes[-20:])
                                
                                    # Save and validate as before
                                    if img_bytes and len(img_bytes) > 8500:
                                        try:
                                            with open("stock_price_plot.png", "wb") as f:
                                                f.write(img_bytes)
                                            print("Image saved as stock_price_plot.png")
                                            # Try to open with PIL to check validity
                                            image = Image.open(BytesIO(img_bytes))
                                            image.verify()
                                            print("Image is valid!")
                                        except Exception as e:
                                            print(f"Invalid image data: {e}")
                                        update_state("images", images)
                                        st.session_state.image = img_bytes

                                    else:
                                        print("No valid image bytes to save.")
                                
                            # Replace previous code results and images with the latest ones
                            update_state("code_results", code_results)
                
                        # Increment counter and refresh display every other message
                        update_counter += 1
                        if update_counter % 2 == 0:  # Refresh after every second message
                            display_state()
                
                    # Display concatenated text at the end (final refresh)
                    if collected_text.strip():
                        update_state("collected_text", collected_text.strip())  # Update final text state
                        display_state()  # Refresh final state dynamically
                
                    print()
                    return all_responses    

                except (ConnectionResetError, ConnectionClosedError) as e:
                    print(f"Connection error: {e}, retries left: {retries}")
                    await asyncio.sleep(backoff)
                    retries -= 1
                    backoff *= 2
                    return None
                    continue
            # st.write("Max retries exceeded... need to retry!")    
        # with open("stock_price_plot.png", "rb") as f:
        #     data = f.read()
        # print("First 8 bytes:", data[:8])
        # print("Last 8 bytes:", data[-8:])
        # print("File size:", len(data))
        
        model = 'gemini-2.0-flash-exp'
        live_client = genai.Client(api_key=GOOGLE_API_KEY,
                                   http_options=types.HttpOptions(api_version='v1alpha'))
        
        # Wrap the existing execute_query tool you used in the earlier example.
        execute_query_tool_def = types.FunctionDeclaration.from_callable(
            client=live_client, callable=execute_query)
        
        # Provide the model with enough information to use the tool, such as describing
        # the database so it understands which SQL syntax to use.
        # sys_int = """Your role is this: You are a database interface. Use the `execute_query` function
        # to answer the users questions by looking up information in the database, running any necessary queries and responding to the user.
        # Be mindful of systme limitation with large output: Example why: ConnectionClosedError: sent 1009 (message too big) frame with 2275770 bytes exceeds limit of 1048576 bytes; no close frame received    
        
        # You need to look up table schema using sqlite3 syntax SQL, then once an answer is found be sure to tell the user. 
        # Important: If the user is requesting an action, you must also execute the actions.
        
        # """
        # sys_int = """Your role is this: You are a database interface. Use the `execute_query` function
        # to answer the user's questions by looking up information in the database, running any necessary queries, and responding to the user.
        # Important: If the user requests a visualization (e.g., a Seaborn chart), you must generate and provide Python code that can be executed directly to create the plot. Ensure the code uses Pandas for data manipulation and Seaborn for plotting, and that it includes all necessary imports. The generated code should be executable without modification.
    
        # """
    
        sys_int = instruction + """
        Your role is this: You are a database interface. Use the [execute_query_tool_def.to_json_dict()] to understand the database/table contents and be able to pull relevant information from the tables.
        and then to answer the user's questions by looking up information in the database, running any necessary queries, and responding to the user. 
        Provide a comprehensive report on each of the selected stocks with data available on the database and provide all final results in text to be used by subsequent agents to summarize further.
        If you recommend an action, you must take that action.
        """
        #Important: If the user requests a visualization (e.g., a Seaborn chart), you must generate Python code that uses Pandas for data manipulation and Seaborn for plotting. Ensure that the generated code includes all necessary imports and replaces plt.show() with logic to return a base64-encoded string of the plot image.
        # Replace plt.show():
            
        #     plt.show() is used for GUI-based backends and does not work in Streamlit.
            
        #     Instead, use st.pyplot(fig) where fig is a Matplotlib Figure object.
    
        config = {
            "response_modalities": ["TEXT"],
            "system_instruction": {"parts": [{"text": sys_int}]},
            "tools": [
                {"code_execution": {}},
                {"function_declarations": [execute_query_tool_def.to_json_dict()]},
            ],
        }
    
        config = {
            "response_modalities": ["TEXT"],
            "system_instruction": sys_int,
            "tools": [
                {"code_execution": {}},
                {"function_declarations": [execute_query_tool_def.to_json_dict()]},
                types.Tool(google_search=types.GoogleSearch())  # Add the Google Search tool
            ],
            "temperature": temperature,
            "top_p": top_p,
        }

        success_T = False            
        async def main(user_query):
            max_attempts_T = 5
            attempt_T = 0
            result=None
            MAX_PAYLOAD_BYTES = 1000000
            prep_db.toast("UPDATED ZOLTAR DATABASE!!!  ", icon="✅")
            # st.session_state.setdefault("agent_progress", {})  # one across all iterations to keep track of what was successful

            for attempt_T in range(1, max_attempts_T + 1):
            
                try:
                    async with live_client.aio.live.connect(model=model, config=config) as session:
                                
                        placeholder_container = st.empty()  # Master container for refreshable content
                        # st.toast("AGENT 1...ZOLTAR DATABASE", icon="⏳")  # Shows a floating toast message
                        if not st.session_state.agent_progress.get("agent1_zoltar") or attempt_T>2:
                                try:
                                    agent1_toast = st.toast("AGENT 1...ZOLTAR DATABASE", icon="⏳")
                                    # sleep(30)
                                    message = user_query+ " ** end of user question** To fully answer this question, after the stock symbols of interest are known, limit to top 5 and in your response include information on them from Zoltar Ranks Database fundamentals table using [execute_query_tool_def.to_json_dict()] for subsequent agents to use, and include sector, P/E, Dividends, 52Week highs and Lows, Overall Rating"
            
                                        
                                    print(f"> {message}\n")
                                    await session.send(input=to_json_serializable(message), end_of_turn=True)
                                    all_responses = await handle_response_refresh(session, tool_impl=execute_query)
                                    if all_responses is None:  # or whatever "bad" value you chose
                                        print(f"Agent failed on attempt {attempt_T}, retrying...")
                                        st.toast("I ran into trouble...RESTARTING", icon="❌")
                                        await asyncio.sleep(1)
                                        continue
                                    else:                                    
                                        agent_result = "\n".join(msg.text for msg in all_responses if msg.text)            
                                        formatted_state = format_global_state(global_state)
                    
                                        # After getting agent_result (Agent 1)
                                        # st.session_state.agent_repo["agents"]["agent1_zoltar"] = {
                                        #     "result": agent_result,
                                        #     "timestamp": datetime.now().isoformat(),
                                        #     "source": "Zoltar Database Query"
                                        # }
                                        # st.session_state.agent_repo["execution_order"].append("agent1_zoltar")
                    
                                        add_agent_result("agent1_zoltar", {
                                            "result": agent_result,
                                            "timestamp": datetime.now().isoformat(),
                                            "source": "Zoltar Database Query"
                                        })
                                        st.session_state.agent_progress["agent1_zoltar"] = True
                                        agent1_toast.toast("AGENT 1...ZOLTAR DATABASE", icon="✅")
                                except Exception as e:
                                    st.toast("I ran into trouble...RESTARTING", icon="❌")
                                    # st.session_state.agent_progress["agent1_zoltar"] = False
                                    return  # Exit, so on next run you'll resume here                        
                        else:
                            agent_result = st.session_state.agent_repo["agents"].get("agent1_zoltar", {}).get("result", None)
                        # Step 2: Ask LLM to check Agent 1's result
                        check_message = user_query + f"""
                            You are checking work performed by Agent #1, whose task it is to: Understand user query, and construct SQL queries and use available tools to gather information from Zoltar Database for requested Summary of Selected Stocks section.
                            Here's Agent 1 task and response: {agent_result}
                            Respond with a single word: ACCURATE or INACCURATE
                        """
                        print(f"> {check_message}\n")
                        await session.send(input=check_message, end_of_turn=True)
                        all_responses_check = await handle_response_refresh(session, tool_impl=execute_query)
                        agent_check_result = "\n".join(msg.text for msg in all_responses_check if msg.text)
                        add_agent_result("agent1_check", {
                            "result": agent_check_result,
                            "timestamp": datetime.now().isoformat(),
                            "source": "Zoltar Database Query Check"
                        })        
                        # Step 3: If INACCURATE, redo Agent 1 with improved instructions
                        if "INACCURATE" in agent_check_result.upper():
                            st.toast("INACCURACY IDENTIFIED, RE-PULLING...", icon="❌")
                            try:
                                agent1_toast = st.toast("AGENT 1...ZOLTAR DATABASE", icon="⏳")
                                # sleep(30)
                                message = user_query+ " ** end of user question** To fully answer this question, after the stock symbols of interest are known, limit to top 5 and in your response include information on them from Zoltar Ranks Database fundamentals table using [execute_query_tool_def.to_json_dict()] for subsequent agents to use, and include sector, P/E, Dividends, 52Week highs and Lows, Overall Rating"
        
                                    
                                print(f"> {message}\n")
                                await session.send(input=to_json_serializable(message), end_of_turn=True)
                                all_responses = await handle_response_refresh(session, tool_impl=execute_query)
                                if all_responses is None:  # or whatever "bad" value you chose
                                    print(f"Agent failed on attempt {attempt_T}, retrying...")
                                    st.toast("I ran into trouble...RESTARTING", icon="❌")
                                    await asyncio.sleep(1)
                                    continue
                                else:                                    
                                    agent_result = "\n".join(msg.text for msg in all_responses if msg.text)            
                                    formatted_state = format_global_state(global_state)
                
                                    # After getting agent_result (Agent 1)
                                    # st.session_state.agent_repo["agents"]["agent1_zoltar"] = {
                                    #     "result": agent_result,
                                    #     "timestamp": datetime.now().isoformat(),
                                    #     "source": "Zoltar Database Query"
                                    # }
                                    # st.session_state.agent_repo["execution_order"].append("agent1_zoltar")
                
                                    add_agent_result("agent1_zoltar", {
                                        "result": agent_result,
                                        "timestamp": datetime.now().isoformat(),
                                        "source": "Zoltar Database Query"
                                    })
                                    st.session_state.agent_progress["agent1_zoltar"] = True
                                    agent1_toast.toast("AGENT 1...ZOLTAR DATABASE", icon="✅")
                            except Exception as e:
                                st.toast("I ran into trouble...RESTARTING", icon="❌")
                                # st.session_state.agent_progress["agent1_zoltar"] = False
                                return  # Exit, so on next run you'll resume here               


                        if not st.session_state.agent_progress.get("agent2_news") or attempt_T>3:
                            try:
                                agent2_toast = st.toast("AGENT 2...NEWS ARTICLES", icon="⏳")
                                # st.toast("AGENT 2...NEWS ARTICLES", icon="⏳")  # Shows a floating toast message
                                # sleep(30)
                                message = f"Search for latest News and analyze Sentiment using types.Tool(google_search=types.GoogleSearch() tool and concise_search that you should use. When searching, only look at the sources specifically selected by the user: {source_str}. Create a table with best 3 links for detailed search, related to the stocks the user asked about found from Zoltar Ranks Database for stocks found by prior agent. Here is the result of the first agent findings: {agent_result}. ** end of prior agent results** And also, provide all final results in text to be used by subsequent agents to summarize further."
                                print(f"> {message}\n")



                                def truncate_to_bytes(s, max_bytes):
                                    encoded = s.encode('utf-8')
                                    if len(encoded) <= max_bytes:
                                        return s
                                    truncated = encoded[:max_bytes].decode('utf-8', 'ignore')
                                    return truncated + "..."

                                # Before sending
                                message_to_send = message
                                while len(message_to_send.encode('utf-8')) > MAX_PAYLOAD_BYTES:
                                    # Truncate aggressively (for example, by removing the last 1000 characters each time)
                                    message_to_send = truncate_to_bytes(message_to_send, len(message_to_send.encode('utf-8')) - 5000)
                                    # Optionally, log or notify user
                                    print(f"Truncated message to {len(message_to_send.encode('utf-8'))} bytes")

                                message = message_to_send
                                

                                
                                await session.send(input=to_json_serializable(message), end_of_turn=True)
                                all_responses2 = await handle_response_refresh(session, tool_impl=execute_query)
                                if all_responses2 is None:  # or whatever "bad" value you chose
                                    print(f"Agent failed on attempt {attempt_T}, retrying...")
                                    st.toast("I ran into trouble...RESTARTING", icon="❌")
                                    await asyncio.sleep(1)
                                    continue
                                else:                                          
                                    agent_result2 = "\n".join(msg.text for msg in all_responses2 if msg.text)  
                                    formatted_state = format_global_state(global_state)
                                    
                                    # # After agent_result2 (Agent 2)
                                    # st.session_state.agent_repo["agents"]["agent2_news"] = {
                                    #     "result": agent_result2,
                                    #     "timestamp": datetime.now().isoformat(),
                                    #     "sources": source_str  # From your user's source selection
                                    # }
                                    # st.session_state.agent_repo["execution_order"].append("agent2_news")
                                    # After agent_result2 (Agent 2)
                                    add_agent_result("agent2_news", {
                                        "result": agent_result2,
                                        "timestamp": datetime.now().isoformat(),
                                        "sources": source_str  # From your user's source selection
                                    })
                                    st.session_state.agent_progress["agent2_news"] = True
                                    agent1_toast.toast("AGENT 1...ZOLTAR DATABASE", icon="✅")
                                    agent2_toast.toast("AGENT 2...NEWS ARTICLES", icon="✅")
                            except Exception as e:
                                st.toast("I ran into trouble...RESTARTING", icon="❌")
                                # st.session_state.agent_progress["agent2_news"] = False
                                return                        
                        else:
                            agent_result2 = st.session_state.agent_repo["agents"].get("agent2_news", {}).get("result", None)

                        if not st.session_state.agent_progress.get("agent3_plots"):
                            try:
                                agent3_toast = st.toast("AGENT 3...OVERVIEW PLOTS", icon="⏳")
        
                                message = f"""Use the result of the first agent findings: {agent_result}. ** end of first agent result ** 
                                     Your task is to create a seaborn plot.  After completing the plot, you should analyze data used for plotting and and create a section "References to visualization", the discussion of the new visualization.
        
                                     You should familarize yourself with contents of Zoltar sqlite3 database to interact with it for Stock trading education app using [execute_query_tool_def.to_json_dict()] tool and should become an expert on the contents of the database and the formats of all variables; and you have access to results found by prior Agent (initial Agent findings: section below) 
                                    Use daily data unless specified otherwise (not 'all_' - since that one which contains intraday data).
                                    Once you have the information you need, you will generate and run some code to get data for the  plot from Zoltar Database tables on the stocks found by Agent #1 as a python seaborn chart, preferrably over time, 
                                    Then generate the plot:
                                    all plot components need to fit horizontally in one frame/image - an informative chart with 2 or 3 or 4 equal horizontally aligned sections:
                                    {viz_section}
                                    Turn x-axis labels -45 degrees.
                             
                           
                                    AND THIS IS ABSOLUTELY CRUCIAL: limit Date ranges to less than 3 months, use complex and nested query logic to FILTER UPFRONT and use aggregation logic in queries when possible.
                                    to get data from db in every SQL query and communication instead of transmitting actual data, or everything will crash.  Estimate size of output using Zoltar database tables detail and expected query output. (be cautious not to hit the total limit of 808576 bytes) 
                                    and don't use textblob. use integers instead of string for indicies. in the past, this has been the issue and helped fix: the structure of the output now. It's a dictionary with a "result" key, whose value is a string containing a JSON-like structure. Inside that string, there's a "results" key containing a list of lists , where each inner list represents a row of data.
                                    high_risk_data['result']  and low_risk_data['result'] are strings, not dictionaries. use the json.loads() function to parse the strings.
                                    If plotting fails more than 2 times, simplify significantly and send only 1 month of data to reduce transmitted payload.
                                    Generate Python code and execute to create a matplotlib/seaborn plot. Make sure to save it with this exact name: stock_price_plot.png
                                    here's an example of how to extract data and use it:
                                        import pandas as pd
                                        import json
                                        
                                        symbols = ['STO1', 'STO2', 'STO3', 'STO4', 'STO5']
                                        sql_returns = f" - tripple quotes here
                                        SELECT Symbol, Score, Score_HoldPeriod, Date
                                        FROM high_risk
                                        WHERE Symbol IN ('"','".join(symbols)') wrong syntax here
                                        AND Date = (SELECT MAX(Date) FROM high_risk WHERE Symbol IN (.join(symbols)')) wrong syntax here
                                        " - tripple quote here
                                        returns_data = default_api.execute_query(sql=sql_returns)
                                        
        
                                    """
                     
                                print(f"> {message}\n")

                                def truncate_to_bytes(s, max_bytes):
                                    encoded = s.encode('utf-8')
                                    if len(encoded) <= max_bytes:
                                        return s
                                    truncated = encoded[:max_bytes].decode('utf-8', 'ignore')
                                    return truncated + "..."

                                # Before sending
                                message_to_send = message
                                while len(message_to_send.encode('utf-8')) > MAX_PAYLOAD_BYTES:
                                    # Truncate aggressively (for example, by removing the last 1000 characters each time)
                                    message_to_send = truncate_to_bytes(message_to_send, len(message_to_send.encode('utf-8')) - 1000)
                                    # Optionally, log or notify user
                                    print(f"Truncated message to {len(message_to_send.encode('utf-8'))} bytes")

                                message = message_to_send
                                
                                await session.send(input=to_json_serializable(message), end_of_turn=True)
                                all_responses2b = await handle_response_refresh(session, tool_impl=execute_query)
                                agent_result2b = "\n".join(msg.text for msg in all_responses2b if msg.text)  
                                if all_responses2b is None:  # or whatever "bad" value you chose
                                    print(f"Agent failed on attempt {attempt_T}, retrying...")
                                    st.toast("I ran into trouble...RESTARTING", icon="❌")
                                    await asyncio.sleep(1)
                                    continue
                                else:                                          
                
                                    # After agent_result2b (Agent 3)
                                    # st.session_state.agent_repo["agents"]["agent3_plots"] = {
                                    #     "result": agent_result2b,
                                    #     "timestamp": datetime.now().isoformat(),
                                    #     "visualizations": viz_section
                                    # }
                                    # st.session_state.agent_repo["execution_order"].append("agent3_plots")
                                    # After agent_result2b (Agent 3)
                                    add_agent_result("agent3_plots", {
                                        "result": agent_result2b,
                                        "timestamp": datetime.now().isoformat(),
                                        "visualizations": viz_section
                                    })
                                    
                                    agent1_toast.toast("AGENT 1...ZOLTAR DATABASE", icon="✅")
                                    agent2_toast.toast("AGENT 2...NEWS ARTICLES", icon="✅")
                                    agent3_toast.toast("AGENT 3...OVERVIEW PLOTS", icon="✅")
                                    st.session_state.agent_progress["agent3_plots"] = True
                            except Exception as e:
                                st.toast("I ran into trouble...RESTARTING", icon="❌")
                                # st.session_state.agent_progress["agent3_plots"] = False
                                return                            
                        else:
                            agent_result2b = st.session_state.agent_repo["agents"].get("agent3_plots", {}).get("result", None)

                        #formatted_state = format_global_state(global_state)
        
                        # while (
                        #     not global_state["images"] or
                        #     all(
                        #         (img is None) or
                        #         (isinstance(img, str) and not img.strip()) or
                        #         (isinstance(img, (bytes, bytearray)) and len(img) == 0)
                        #         for img in global_state["images"]
                        #     )
                        # ):
                            
                        max_tries = 3
                        tries = 0
                        agent4_toasts = []
                        try:
        
                            # while (
                            #     (tries < max_tries) and (
                            #         not global_state["images"] or
                            #         all(
                            #             (img is None)
                            #             or (isinstance(img, str) and not img.strip())
                            #             or (isinstance(img, (bytes, bytearray)) and (len(img) == 0 or is_blank_png(img)))
                            #             for img in global_state["images"]
                            #         )                    
                            #     )
                            # ):
            
        
                                
                            while (
                                (tries < max_tries) and (
                                    (not st.session_state.image) or
                                    is_blank_png(st.session_state.image)                   
                                ) and (Pie_chart or Return_hold  or low_ranks_trend or recommendations_table)
                            ):
                                # print("BEFORE FALLBACK: tries", tries, "image", st.session_state.image, "is_blank", is_blank_png(st.session_state.image))                                
                                tries += 1
                                
                                toast_msg = f"AGENT 4...FALLBACK PLOTS (TRY #{tries})"
                                agent1_toast.toast("AGENT 1...ZOLTAR DATABASE", icon="✅")
                                agent2_toast.toast("AGENT 2...NEWS ARTICLES", icon="✅")
                                agent3_toast.toast("AGENT 3...OVERVIEW PLOTS", icon="✅")
                                agent4_toast = st.toast(toast_msg, icon="⏳")
                                agent4_toasts.append(agent4_toast)
                            
                                def truncate_to_bytes(s, max_bytes):
                                    encoded = s.encode('utf-8')
                                    if len(encoded) <= max_bytes:
                                        return s
                                    truncated = encoded[:max_bytes].decode('utf-8', 'ignore')
                                    return truncated + "..."  # Add ellipsis to indicate truncation
                            
                                # Truncate agent_result for tries > 1
                                if tries == 1:
                                    agent_result_to_use = agent_result
                                else:
                                    agent_result_to_use = truncate_to_bytes(agent_result, len(agent_result) - tries * 1000)
                                    
                                message = f"""Use the result of the first agent findings: {agent_result_to_use}. ** end of first agent result ** 
                                     Your task is to create a plot. This is attempt number {tries}.  After completing the plot, yoou should analyze data used for plotting and and create a section "References to visualization", the discussion of the new visualization.                                        
                                     You can interact with Zoltar SQL database for Stock trading education app using [execute_query_tool_def.to_json_dict()] tool and should become an expert on the contents of the database and the formats of all variables; and you have access to results found by prior Agent (initial Agent findings: section below) 
                                    Use daily data unless specified otherwise (not 'all_' - since that one which contains intraday data).
                                    can interact with an SQL database for Stock trading education app. You will take the users' questions and turn them into SQL
                                    queries using the tools available. Once you have the information you need, you will generate and run some code to plot data from Zoltar Database tables on the stocks found by Agent #1 as a python seaborn chart, preferrably over time, 
                                    Then generate the plot with only two horizontally lined up sections from the requested vizualizations below, which need to fit in one landscape positioned frame/image - an informative chart with the following sections:
                                    {viz_section}
                                    Turn x-axis labels -45 degrees.                     
                                    
                            
                                    AND THIS IS ABSOLUTELY CRUCIAL: The prior attempt to generate the plot failed due to exceeding payload limit and being careless, even after taking this into account.. limit Date ranges to less than 3 months, use complex and nested query logic to FILTER UPFRONT and use aggregating functions in queries when possible
                                    to get data from db in every SQL query and communication instead of transmitting actual data, or everything will crash.  Estimate size of output using Zoltar database tables detail and expected query output. (be cautious not to hit the total limit of 808576 bytes) 
                                    and don't use textblob.  in the past, this has been the issue and helped fix: the structure of the output now. It's a dictionary with a "result" key, whose value is a string containing a JSON-like structure. Inside that string, there's a "results" key containing a list of lists , where each inner list represents a row of data.
                                    high_risk_data['result']  and low_risk_data['result'] are strings, not dictionaries. use the json.loads() function to parse the strings.
                                    If plotting fails more than 2 times, simplify significantly and send only 1 month of data. 
                                    Generate Python code and execute to create matplotlib/seaborn plot.
                                    here's an example of how to extract data and use it:
                                        import pandas as pd
                                        import json
                                        
                                        symbols = ['STO1', 'STO2', 'STO3', 'STO4', 'STO5']
                                        sql_returns = f" - tripple quotes here
                                        SELECT Symbol, Score, Score_HoldPeriod, Date
                                        FROM high_risk
                                        WHERE Symbol IN ('"','".join(symbols)') wrong syntax here
                                        AND Date = (SELECT MAX(Date) FROM high_risk WHERE Symbol IN (.join(symbols)')) wrong syntax here
                                        " - tripple quote here
                                        returns_data = default_api.execute_query(sql=sql_returns)           
                                    
                                    """
                    
                                print(f"> {message}\n")
                                # Before sending:
                                # debug_payload(message)
                                try:
                                    await session.send(input=to_json_serializable(message), end_of_turn=True)
                                    all_responses2c = await handle_response_refresh(session, tool_impl=execute_query)
                                    # Defensive: handle None
                                    if all_responses2c is None:
                                        all_responses2c = []
                                    agent_result2c = "\n".join(msg.text for msg in all_responses2c if msg and hasattr(msg, 'text'))
                                    agent1_toast.toast("AGENT 1...ZOLTAR DATABASE", icon="✅")
                                    agent2_toast.toast("AGENT 2...NEWS ARTICLES", icon="✅")
                                    agent3_toast.toast("AGENT 3...OVERVIEW PLOTS", icon="✅")
                                    agent4_toast.toast(toast_msg, icon="✅")
                                    # If successful, break out of the loop
                                    break
                                except Exception as e:
                                    error_placeholder = st.empty()
                                    error_placeholder.error(f"Plotting attempt {tries} failed: {e}")
                                    agent4_toast.toast(f"AGENT 4 failed on attempt {tries}: {e}", icon="❌")
                                    # Optionally, wait before next try
                                    await asyncio.sleep(1)
                                    error_placeholder.empty()
                                    continue  # Go to next try
                                #formatted_state = format_global_state(global_state)    
                        except RuntimeError as e:
                            st.error(f"Stage 2c failed: {e}")
                            agent_result2c = "Stage 2c failed. No plot generated due to exceeding payload limit."
                            st.toast("AGENT 4 failed: Could not generate plot.", icon="❌")                    
                            # Optionally: continue to next stage    
                        # st.toast("AGENT 5...COMPILE REPORT", icon="⏳")  # Shows a floating toast message
                        agent1_toast.toast("AGENT 1...ZOLTAR DATABASE", icon="✅")
                        agent2_toast.toast("AGENT 2...NEWS ARTICLES", icon="✅")
                        agent3_toast.toast("AGENT 3+4...OVERVIEW PLOTS", icon="✅")
                        agent5_toast = st.toast("AGENT 5...SHAP ANALYSIS", icon="⏳")

                        # sleep(30)


# another agent to handle SHAP analysis

                        message = f"""Use the result of the first agent findings: {agent_result}. ** end of first agent result ** 
                              Your task is to generate SHAP analysis section for the final reoprt on these stocks.
                            You should always attempt to create a SHAP table for every stock found, and print all of the ones found in final response - the records in SHAP tables may not exist for every stock - check them every time.  
                              You should familarize yourself with contents of Zoltar sqlite3 database, specifically the 3 SHAP tables and the code below to create a meaningful table, to interact with it using [execute_query_tool_def.to_json_dict()] tool  for Stock trading education app using [execute_query_tool_def.to_json_dict()] tool and should become an expert on the contents of the database and the formats of all variables; and you have access to results found by prior Agent (initial Agent findings: section below) 
                            Use daily data unless specified otherwise (not 'all_' - since that one which contains intraday data).

                            here's an example of how to extract data and use it:
                                import pandas as pd
                                import json
                                
                                symbols = ['STO1', 'STO2', 'STO3', 'STO4', 'STO5']
                                sql_returns = f" - tripple quotes here
                                SELECT Symbol, Score, Score_HoldPeriod, Date
                                FROM high_risk
                                WHERE Symbol IN ('"','".join(symbols)') wrong syntax here
                                AND Date = (SELECT MAX(Date) FROM high_risk WHERE Symbol IN (.join(symbols)')) wrong syntax here
                                " - tripple quote here
                                returns_data = default_api.execute_query(sql=sql_returns)

                        """
                        message +="""
                            This is the exact function (with notes for places to replace with tripple quotes. Convert the code to use for default_api.execute_query tool and use it to create the table of Features and corresponding SHAP Values for each symbol (but you have to check all 3 SHAP files - Small, Mid and Large):  
                            def create_shap_table(symbols, db_path='zoltar_database.sqlite3'):
                                conn = sqlite3.connect(db_path)
                                all_shap = []
                                
                                # Get list of SHAP tables
                                shap_tables = pd.read_sql( ***tripple quote here***
                                    SELECT name 
                                    FROM sqlite_master 
                                    WHERE type='table' 
                                    AND name LIKE 'shap_summary_%'
                                ***tripple quote here***, db_conn)['name'].tolist()
                                
                                for symbol in symbols:
                                    symbol_data = []
                                    
                                    for table in shap_tables:
                                        # Get most recent data for symbol
                                        query = f***tripple quote here***
                                            SELECT * 
                                            FROM {table} 
                                            WHERE Symbol = '{symbol}'
                                            LIMIT 1
                                        ***tripple quote here***
                                        df = pd.read_sql(query, db_conn)
                                        
                                        if not df.empty:
                                            # Process SHAP values
                                            numeric_cols = df.select_dtypes(include='number').columns
                                            for col in numeric_cols:
                                                value = df[col].values[0]
                                                if pd.notnull(value) and value != 0:
                                                    symbol_data.append({
                                                        'Symbol': symbol,
                                                        'SHAP Table': table,
                                                        'Feature': col,
                                                        'SHAP Value': f"{value:.9f}",
                                                        'Impact': "Increasing" if value > 0 else "Decreasing"
                                                    })
                                    
                                    if not symbol_data:
                                        all_shap.append(pd.DataFrame({
                                            'Symbol': [symbol],
                                            'Status': ['No SHAP data found']
                                        }))
                                    else:
                                        all_shap.append(pd.DataFrame(symbol_data))
                                
                                conn.close()
                                return pd.concat(all_shap).reset_index(drop=True)
                            
                            # Test with your symbols
                            symbols = ['F']  - this is an example - substitute the symbol of interest here
                            shap_results = create_shap_table(symbols)

                            Iterate through the stock symbols and create a table of top 5 features for each. The table should be printed as is from the above function. If symbol is not on any of the SHAP tables, mark it missing.
                            To get feature names use SHAP database table column names  (alphanumeric) - this is important to present in the table.

                            """
#                         message+="""This is the working function in my app that uses dfs as inputs - you have these tables in sqlite database. Convert the code to create the table of Features and corresponding SHAP Values for each symbol after going through the logic below to extract usable tables:
#                             def load_shap_summaries():
#                                 cap_sizes = ['Large', 'Mid', 'Small']
#                                 combined_summary_df = pd.DataFrame()
                            
#                                 for cap_size in cap_sizes:
#                                     latest_file = find_most_recent_file(data_dir, f'combined_SHAP_summary_{cap_size}_')
#                                     if latest_file:
#                                         df = pd.read_pickle(latest_file)
#                                         combined_summary_df = pd.concat([combined_summary_df, df])
#                                     else:
#                                         print(f"No SHAP summary file found for {cap_size} cap size.")
                            
#                                 return combined_summary_df  
#                             combined_summary_df = load_shap_summaries()


#                             def create_shap_table(combined_summary_df, symbol):
#                                 if symbol not in combined_summary_df.index:
#                                     return None
                                
#                                 stock_data = combined_summary_df.loc[symbol]
#                                 numeric_data = stock_data[pd.to_numeric(stock_data, errors='coerce').notnull()]
#                                 top_features = numeric_data.abs().sort_values(ascending=False).head(5)
#                                 shap_table = []
                                
#                                 for feature in top_features.index:
#                                     value = numeric_data[feature]
#                                     if pd.notnull(value) and value != 0:
#                                         direction = "Increasing" if value > 0 else "Decreasing"
#                                         shap_table.append({
#                                             "Feature": feature,
#                                             "SHAP Value": f"{value:.9f}",
#                                             "Impact": direction
#                                         })
                                
#                                 return pd.DataFrame(shap_table)

#                             shap_df = create_shap_table(combined_summary_df, symbol)
#                             if shap_df is not None:
#                                 st.table(shap_df)
#                             else:
#                                 st.write("No SHAP data available for this stock.")
                                
                                
#                             Convert the code to use for default_api.execute_query tool and using SHAP sqlite tables and iterate through the stock symbols and create a table of top 5 features for each. The table should be printed as is from the above function. If symbol is not on any of the SHAP tables, mark it missing.
#                             Use SHAP table columns names (not sequential numbering but alphanumeric column names that I have in my original dfs/sql tables) for feature names use.  They are named with alphanumeric names that need to be presented in the table.

# """
                        print(f"> {message}\n")
                        await session.send(input=to_json_serializable(message), end_of_turn=True)
                        all_responses4 = await handle_response_refresh(session, tool_impl=execute_query)
                        agent_result4 = "\n".join(msg.text for msg in all_responses4 if msg.text)  
                        
                        add_agent_result("agent_result4", {
                            "result": agent_result4,
                            "timestamp": datetime.now().isoformat()
                        })                        
                        
                        agent1_toast.toast("AGENT 1...ZOLTAR DATABASE", icon="✅")
                        agent2_toast.toast("AGENT 2...NEWS ARTICLES", icon="✅")
                        agent3_toast.toast("AGENT 3+4...OVERVIEW PLOTS", icon="✅")
                        agent5_toast.toast("AGENT 5...SHAP ANALYSIS", icon="✅")
                        agent6_toast = st.toast("AGENT 6...COMPILE REPORT", icon="⏳")
            
                        #message = f"Generate and run some code to pull necessary data from Zoltar Ranks Database for stocks found by prior agent. Plot the Price and Zoltar Ranks over time as a python seaborn chart. Return base64-encoded images.  Here is the result of the first agent findings: {agent_result2}. ***IMPORTANT*** there is a limit of 4000 characters on output so use efficient sub-queries to filter and limit timeframe to 30 days."
                        message = f"""Combine the results of prior agants into a comprehensive report, and make sure to use all information synthesized by prior agents to answer this original query: {user_query}. ** End of User Query ** 
                        Here is the result of the first agent findings: {agent_result}. ***End of AGENT 1 results***
                        Here is the result of the second agent findings: {agent_result2}. ***End of AGENT 2 results**** 
                        And this is commentary of the supporting plots: {agent_result2b} *** End of Agent 3 Results *** 
                        And this is the SHAP section: {agent_result4}  *** End of Agent 4 Results *** 
                        The final report needs to have an executive structure, containing 
                            1. Summary section with a sentence caputuring the essense of the report and table of Fundamentals/About Information and overall recommendation column (Buy, Mixed, Sell), 
                            2. News and Ratings section with Summary table for News and for Analyst Ratings with columns: Analyst Consensus,Blogger Sentiment,	Crowd Wisdom,	News Sentiment; 
                                Make sure to include the links section for each stock listed (from agent 2 results) below the summary table
                            3. Quant Section with Zoltar Ranks, their direction and SHAP discussion; 
                            4. Conclusion based on contents of prior section. 
                        Return just the Final Executive Report and nothing else."""
                        print(f"> {message}\n")
                        # debug_payload(message)
                        await session.send(input=to_json_serializable(message), end_of_turn=True)
                        all_responses3 = await handle_response_refresh(session, tool_impl=execute_query)
                        agent_result3 = "\n".join(msg.text for msg in all_responses3 if msg.text)  
                        st.session_state.final_agent_result = agent_result3
    
                        # After agent_result3 (Final report)
                        add_agent_result("agent5_final_report", {
                            "result": agent_result3,
                            "timestamp": datetime.now().isoformat(),
                            "source": "Final Executive Report"
                        })                    
                        agent1_toast.toast("AGENT 1...ZOLTAR DATABASE", icon="✅")
                        agent2_toast.toast("AGENT 2...NEWS ARTICLES", icon="✅")
                        agent3_toast.toast("AGENT 3+4...OVERVIEW PLOTS", icon="✅")
                        agent5_toast.toast("AGENT 5...SHAP ANALYSIS", icon="✅")
                        agent6_toast.toast("AGENT 6...COMPILE REPORT", icon="✅")
                        st.toast("Final report completed!", icon="✅")
                        st.balloons()
                        break
                        #formatted_state = format_global_state(global_state)
                except Exception as e:
                    # Show error for 1 second, then clear
                    error_placeholder = st.empty()
                    error_placeholder.error(f"Connection failed (attempt {attempt_T}/{max_attempts_T}): {e}")
                    st.toast("I ran into trouble...RESTARTING", icon="❌")
                    # await asyncio.sleep(1)
                    error_placeholder.empty()
                    if attempt_T == max_attempts_T:
                        st.error("All attempts to connect failed. Please try again with less complex settings.")
                        return
        # if not success_T:
        #     st.error("All attempts to connect failed. Please try again later.")    
        # with col2:
        # Run the async code
        asyncio.run(main(user_query))





# OLDER VERSION (2.3)










# Email section 
    # if st.session_state.final_agent_result:
    agent_keys = st.session_state.agent_repo["execution_order"]
    has_valid_result = any(
    st.session_state.agent_repo["agents"][key].get("result")
    for key in agent_keys
    )
    if has_valid_result:
# show all agent results (ala carte)
        # if st.checkbox("Show Agent Repository"):
        #     st.subheader("Agent Execution History")
        with st.expander("Tracking of Agent interactions throughout the run:"):
            # Show execution order
            st.write("### Execution Sequence")
            for idx, agent_key in enumerate(st.session_state.agent_repo["execution_order"], 1):
                agent_data = st.session_state.agent_repo["agents"][agent_key]
                st.write(f"{idx}. **{agent_key}** ({agent_data['timestamp']})")
        
            # Create a tab for each agent, always displaying all results
            agent_keys = st.session_state.agent_repo["execution_order"]
            if agent_keys:
                agent_tabs = st.tabs([f"{key}" for key in agent_keys])
            
                for tab, agent_key in zip(agent_tabs, agent_keys):
                    agent_data = st.session_state.agent_repo["agents"][agent_key]
                    with tab:
                        st.markdown(f"#### Agent: `{agent_key}`")
                        st.write(f"**Timestamp:** {agent_data['timestamp']}")
                        st.write("**Raw Result:**")
                        st.code(agent_data["result"], language="text")
                        st.write("**Metadata:**")
                        st.json({k: v for k, v in agent_data.items() if k != "result"})
            
                # Save to JSON file
                with open("agent_repo.json", "w") as f:
                    json.dump(st.session_state.agent_repo, f)
                
            # Load from JSON file
            if st.button("Load Previous Agent Repository"):
                if os.path.exists("agent_repo_t.json"):
                    with open("agent_repo_t.json", "r") as f:
                        st.session_state.agent_repo = json.load(f)
                        st.toast("Loaded Previous Agent Repo", icon="✅")
                        st.rerun()
                elif os.path.exists("agent_repo.json"):
                    with open("agent_repo.json", "r") as f:
                        st.session_state.agent_repo = json.load(f)
                        st.toast("No Previous Repo, Loaded Current", icon="✅")
                        # st.rerun()

        
        with st.popover("✅ Ready to share the results?"):   
     # still continuing with col2
           
            ## 5.24.25: new section to email results
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            import smtplib
            from email.mime.multipart import MIMEMultipart
            from email.mime.text import MIMEText
            from email.mime.base import MIMEBase
            from email.mime.image import MIMEImage
            from email import encoders
            
            def add_bold_runs(paragraph, text):
                import re
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = paragraph.add_run(part[2:-2])
                        run.bold = True
                    else:
                        paragraph.add_run(part)
            
            def save_to_docx(content, filename="Bot_Output.docx", image_path="stock_price_plot.png"):
                doc = Document()
                doc.add_heading('Your Zoltar Financial Research', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if os.path.exists(image_path):
                    doc.add_picture(image_path, width=Inches(6))
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    if line.startswith('## '):
                        header_text = line[3:].strip()
                        p = doc.add_heading(header_text, level=2)
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    elif line.startswith('* '):
                        bullet_text = line[2:].strip()
                        p = doc.add_paragraph(style='List Bullet')
                        add_bold_runs(p, bullet_text)
                    else:
                        p = doc.add_paragraph()
                        add_bold_runs(p, line)
                doc.save(filename)
                return filename
            
            def send_email(sender, password, recipient, doc_path):
                msg = MIMEMultipart()
                msg['From'] = f"Zoltar Financial <{sender}>"
                msg['To'] = recipient
                msg['Subject'] = "Your Zoltar Research Report"
                body = "Thank you for using Zoltar Financial Research Assistant. Please find attached the generated report."
                msg.attach(MIMEText(body, 'plain'))
                with open(doc_path, "rb") as attachment:
                    part = MIMEBase("application", "octet-stream")
                    part.set_payload(attachment.read())
                encoders.encode_base64(part)
                part.add_header("Content-Disposition", f"attachment; filename={os.path.basename(doc_path)}")
                msg.attach(part)
                try:
                    server = smtplib.SMTP_SSL('smtp.gmail.com', 465)
                    server.login(sender, password)
                    server.send_message(msg)
                    server.close()
                    return True
                except Exception as e:
                    st.error(f"Failed to send email: {e}")
                    return False
            def get_image_base64():
                import requests, base64
                image_url = 'https://github.com/apod-1/ZoltarFinancial/raw/main/docs/ZoltarSurf2.png'
                response = requests.get(image_url)
                if response.status_code == 200:
                    img_data = response.content
                    img_base64 = base64.b64encode(img_data).decode('utf-8')
                    return img_base64
                else:
                    print(f"Failed to fetch image. Status code: {response.status_code}")
                    return None
            img_b64 = get_image_base64()
            
            def send_email(sender, password, recipient, doc_path):
                msg = MIMEMultipart()
                msg['From'] = f"Zoltar Financial <{sender}>"
                msg['To'] = recipient
                msg['Subject'] = "Your Zoltar Research Report"
            
                # Email body with inline image reference
                # html_body = """
                #     <html>
                #         <body>
                #             <h2>Stock Price Plot</h2>
                #             <img src="cid:stock_price_plot">
                #             <p>Thank you for using Zoltar Financial Research Assistant. Please find attached the generated report.</p>
                #             <p>May the riches be with you..</p>
                #         </body>
                #     </html>
                # """
                # msg.attach(MIMEText(html_body, 'html'))
                html_body = f"""
                    <html>
                        <body>
                            <h2>Your Stock Plots</h2>
                            <img src="cid:stock_price_plot">
                            <p>Thank you for using Zoltar Financial Research Assistant. Please find attached the generated report.  Pardon our mess - we are working on improving the user experience.  This is not an investment advice.</p>
                            <p><img src="data:image/png;base64,{img_b64}" alt="ZoltarSurf" style="max-width: 600px; width: 30%; height: auto;"></p>
                            <p>May the riches be with you..</p>
                        </body>
                    </html>
                """
                msg.attach(MIMEText(html_body, 'html'))            
                # Attach the plot image inline
                try:
                    with open("stock_price_plot.png", "rb") as img_file:
                        img = MIMEImage(img_file.read())
                        img.add_header('Content-ID', '<stock_price_plot>')
                        img.add_header('Content-Disposition', 'inline', filename="stock_price_plot.png")
                        msg.attach(img)
                except Exception as e:
                    # Optionally handle missing image
                    pass
            
                # Attach the report
                with open(doc_path, "rb") as attachment:
                    part = MIMEBase("application", "octet-stream")
                    part.set_payload(attachment.read())
                encoders.encode_base64(part)
                part.add_header("Content-Disposition", f"attachment; filename={os.path.basename(doc_path)}")
                msg.attach(part)
            
                try:
                    server = smtplib.SMTP_SSL('smtp.gmail.com', 465)
                    server.login(sender, password)
                    server.send_message(msg)
                    server.close()
                    return True
                except Exception as e:
                    st.error(f"Failed to send email: {e}")
                    return False
            
            # --- Streamlit App ---
            
            st.header("Share your research results", help="Save this report as a .docx file and email it to yourself.")
        
        
            
            # Example: get your research content and image path
            # Replace this with your actual result variable
            content = st.session_state.get("final_agent_result", "The results are empty!")
            image_path = "stock_price_plot.png"
            
            try:
                sender = st.secrets["GMAIL"]["GMAIL_ACCT"]
                password = st.secrets["GMAIL"]["GMAIL_PASS"]
            except:
                # If Streamlit secrets are not available, use environment variables
                sender = os.getenv('GMAIL_ACCT')
                password = os.getenv('GMAIL_PASS') 
                # st.error("Gmail credentials not found in secrets. Please check your configuration.")
            
            with st.form("email_form"):
                recipient = st.text_input("Recipient email address")
                # sender = st.text_input("Sender Gmail address")
                # password = st.text_input("Sender Gmail password (use App Password)", type="password")
                submitted = st.form_submit_button("Send Report")
                if submitted:
                    if not recipient or not sender or not password:
                        st.error("Please fill in all fields.")
                    else:
                        current_date = datetime.now().strftime('%m%d%y')
                        doc_path = save_to_docx(content, filename=f"zoltar_financial_research_report_{current_date}.docx", image_path=image_path)
                        st.success(f"Document saved as {doc_path}")
                        if send_email(sender, password, recipient, doc_path):
                            st.success(f"Report sent successfully to {recipient}!")